"""Generate the CardioKB System Design Word document.

This is a high-level planning doc for collaborators to divvy up work,
not a basic technical specification.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()
    return table


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        p.add_run(f'{bold_prefix}: ').bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        p.add_run(f'{bold_prefix}: ').bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def build_doc():
    doc = Document()

    # -- Title Page --
    doc.add_paragraph()
    title = doc.add_heading('CardioKB: System Design & Work Division', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        'A planning document for collaborators building a biomedical knowledge graph '
        'for hypothesis generation and precision medicine'
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Version 2.0  |  March 2026\n').font.size = Pt(11)
    meta.add_run('12-Week Rotation Project (January \u2013 April 2026)').font.size = Pt(11)

    doc.add_page_break()

    # -- Table of Contents --
    add_heading(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. What This Document Is For',
        '2. What Is Built and Working',
        '   2.1. Data Pipeline (36 Sources)',
        '   2.2. AI Agents',
        '   2.3. Web Interface',
        '   2.4. Post-Processing & Scoring',
        '3. What Is Planned or Partially Implemented',
        '4. What Is Explicitly Out of Scope',
        '5. Proposed Work Division: Agent Layer vs. Data Layer',
        '   5.1. Data/Parser Layer',
        '   5.2. Agent/Orchestration Layer',
        '   5.3. ML/Analysis Layer',
        '   5.4. Interface Layer',
        '   5.5. Dependency Map',
        '6. Schema Decisions & Edge Merge Recommendations',
        '   6.1. Current Schema Overview',
        '   6.2. Edges That Should Be Merged',
        '   6.3. Edges That Should Stay Separate',
        '   6.4. Naming Convention Cleanup',
        '7. Knowledge Graph to ML: Hypothesis Generation Pipeline',
        '   7.1. Current State (Working)',
        '   7.2. Planned: Embedding-Based Approaches',
        '   7.3. Planned: Feature Engineering for Classifiers',
        '   7.4. Integration Points',
        '8. Current Stats (Snapshot)',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

    doc.add_page_break()

    # =================================================================
    # 1. What This Document Is For
    # =================================================================
    add_heading(doc, '1. What This Document Is For', level=1)
    doc.add_paragraph(
        'This is a planning document, not a technical specification. Its purpose is to help '
        'collaborators understand what exists, what remains to be built, and how to divide '
        'the work. Use it to decide who owns what and to identify dependencies between pieces.'
    )
    doc.add_paragraph(
        'CardioKB is a disease-agnostic biomedical knowledge graph currently scoped to cardiovascular '
        'disease through a modular filter. It integrates 36 data sources into Neo4j with ~4.9M nodes '
        'and ~26.3M relationships. Two Claude-powered AI agents handle on-demand data ingestion. '
        'A link prediction notebook demonstrates that the graph is ML-ready. The system is adapted '
        'from AlzKB (Alzheimer\'s Knowledge Base) but replaces the RDF pipeline with direct Cypher '
        'loading, adds 28+ sources, and introduces AI-assisted parser generation.'
    )

    doc.add_page_break()

    # =================================================================
    # 2. What Is Built and Working
    # =================================================================
    add_heading(doc, '2. What Is Built and Working', level=1)
    doc.add_paragraph(
        'Everything in this section is tested, runs end-to-end, and produces correct output. '
        'These are not aspirational \u2014 they work today.'
    )

    # 2.1 Data Pipeline
    add_heading(doc, '2.1. Data Pipeline (36 Sources)', level=2)
    doc.add_paragraph(
        'The full batch pipeline (python src/main.py) downloads, parses, exports TSVs, and loads '
        'into Neo4j in a single run. All 36 parsers inherit from BaseParser and produce '
        'Dict[str, DataFrame]. 85 ontology configs in src/ontology_configs.py declaratively '
        'map those DataFrames to Neo4j nodes and relationships. The loader uses MERGE for '
        'idempotent loading. Every relationship carries a source property for provenance.'
    )

    doc.add_paragraph('Phase 1 \u2014 Core Parsers (8):')
    add_table(doc,
              ['Source', 'What It Provides', 'Notable Details'],
              [
                  ['ClinicalTrials.gov', '82K trials, 674 condition + 18K intervention edges',
                   'API v2, JSON caching per disease term'],
                  ['ClinPGx (PharmGKB successor)', '1.1K variant, 506 + 360 + 304 pgx edges',
                   'Public REST API, drug labels + clinical annotations'],
                  ['NCBI Gene', '194K gene nodes', 'Public FTP, gene metadata only (no edges)'],
                  ['DoRothEA (OmniPath)', '15K TF-gene interactions',
                   'morScore + confidence (A\u2013D) properties on edges'],
                  ['OMIM', '7.3K gene-disease edges', 'Requires OMIM_API_KEY'],
                  ['DisGeNET', '20K gene-disease edges', 'Requires DISGENET_API_KEY, CVD-filtered'],
                  ['DrugBank', '41K drugs, 19K drug-target edges',
                   'Streaming XML parser for 1.8GB file'],
                  ['AOP-DB', '18.5K gene-pathway edges (bidirectional)',
                   'SQL dump or MySQL, adverse outcome pathways'],
              ])

    doc.add_paragraph('Phase 2 \u2014 Hetionet Component Parsers (21):')
    add_table(doc,
              ['Source', 'Key Contribution', 'Edge Count'],
              [
                  ['Disease Ontology', '19K disease nodes + ontology structure', 'Nodes only'],
                  ['Gene Ontology', 'BP + MF + CC annotations', '322K edges'],
                  ['Uberon', '14.9K anatomy nodes', 'Nodes only'],
                  ['MeSH', '966 symptom nodes', 'Nodes only'],
                  ['SIDER', 'Drug side effects', '148K edges'],
                  ['LINCS L1000', 'Compound regulation + gene regulation', '16.7K edges (w/ zScore)'],
                  ['MEDLINE', 'Literature cooccurrence (disease-symptom, disease-anatomy)', '1.3K edges'],
                  ['DrugCentral', 'Drug-disease treatment + palliation + pharmacologic classes', '18K edges'],
                  ['GWAS Catalog', 'Genome-wide association gene-disease links', '45.5K edges'],
                  ['BindingDB', 'Chemical-gene binding', '4.2K edges'],
                  ['PubTator Central', 'Literature-mined gene-disease + disease-disease', '3.4M edges'],
                  ['CTD', 'Chemical increases/decreases gene expression', '431K edges'],
                  ['Bgee', 'Tissue over/underexpression', '5.3M edges (w/ expressionScore)'],
                  ['Hetionet (precomputed)', 'Side effects, gene interactions, covariance', '143K edges'],
                  ['Jensen DISEASES', 'Text-mined gene-disease', '20.5K edges'],
                  ['Jensen TISSUES', 'Text-mined gene-tissue expression', '982K edges'],
                  ['HPO', '19K phenotypes, gene-phenotype associations', '30.5K edges'],
                  ['Reactome', 'Curated pathway membership', '32.6K edges (bidirectional)'],
                  ['WikiPathways', 'Community-curated pathway membership', '17.1K edges (bidirectional)'],
                  ['STRING', 'Protein-protein interactions (confidence > 700)', '229K edges'],
                  ['OpenTargets', 'Evidence-scored gene-disease associations', '2.4M edges'],
              ])

    doc.add_paragraph('Phase 3 \u2014 Agent-Generated Parsers (7):')
    doc.add_paragraph(
        'These were generated entirely by the DatabaseAgent from just a name + URL: '
        'HGNC (gene enrichment), HGNC Families (34K edges), ClinVar (4.5M variants, 12.6M edges), '
        'DrugAge (866 aging edges), CellAge (senescence genes), AnAge (4.6K species), GenAge (aging genes).'
    )

    doc.add_paragraph('ID Mapping & Cross-Reference Resolution:')
    doc.add_paragraph(
        'src/id_mapping.py handles the fundamental challenge of data integration: different sources '
        'use different identifiers. It registers 20 ID systems (NCBI Gene, Ensembl, gene symbols, '
        'DOID, UMLS CUI, MeSH, DrugBank IDs, etc.) and provides validate_mapping(), '
        'suggest_mapping(), and create_missing_nodes() for gap repair. Two critical post-processing '
        'remaps run before Neo4j loading: PubTator MeSH\u2192DOID (via Disease Ontology xrefs) and '
        'GWAS trait\u2192DOID (3-level strategy: name match, curated trait match, ontology URI extraction).'
    )

    # 2.2 AI Agents
    add_heading(doc, '2.2. AI Agents', level=2)

    doc.add_paragraph('DatabaseAgent (src/database_agent.py):')
    doc.add_paragraph(
        'Autonomously generates new parsers from just a name and URL. Downloads first 64KB to detect '
        'format, sends sample + BaseParser source to Claude, generates parser + ontology configs, '
        'validates column references, checks for performance anti-patterns, registers in the pipeline, '
        'runs the parser, validates ID mappings, and loads into Neo4j. 7 production parsers were '
        'created this way. Uses Claude Haiku 4.5 by default.'
    )

    doc.add_paragraph('DiseaseQueryAgent (src/disease_agent.py):')
    doc.add_paragraph(
        'On-demand disease enrichment via the web UI. Standardizes disease names via Claude (handles '
        'abbreviations like "AF" \u2192 "Atrial Fibrillation"), checks Neo4j cache, fetches DisGeNET '
        'gene-disease associations + ClinicalTrials.gov trials, loads into Neo4j, caches results. '
        'Capped at 200 disease IDs and 500 trials per query for performance.'
    )

    # 2.3 Web Interface
    add_heading(doc, '2.3. Web Interface', level=2)
    doc.add_paragraph(
        'Single-page app (interface/index.html) backed by Flask (src/api.py, port 5050) '
        'with 15 REST + SSE endpoints:'
    )
    add_bullet(doc, 'Vis.js graph visualization with Core/Discovery layer toggle, '
               'node-type filter chips, specificity ranking, CSV/JSON export',
               'Explore tab')
    add_bullet(doc, 'Neo4j Browser-style multi-panel Cypher interface with query templates, '
               'read-only enforcement, table + graph result views',
               'Query tab')
    add_bullet(doc, 'AI-driven disease enrichment (DisGeNET + ClinicalTrials.gov) with SSE progress',
               'Build Knowledge Graph (sidebar)')
    add_bullet(doc, 'Variable-hop (1\u20133) subgraph extraction with stats and JSON/CSV export',
               'Extract Disease Subgraph (sidebar)')
    add_bullet(doc, 'Parser status, health check charts, ID mapping report, pipeline execution',
               'Admin panel (password-protected)')
    add_bullet(doc, 'Add New Database (DatabaseAgent UI, password-protected)',
               'Advanced')

    # 2.4 Post-Processing
    add_heading(doc, '2.4. Post-Processing & Scoring', level=2)
    add_bullet(doc, '1.0 / (distinct Disease neighbor count). Stored as node property. '
               'Genes connected to 5 diseases score 0.2; genes connected to 20K score 0.00005. '
               'Disease nodes get 0.0. Computed in batches of 50K to respect Neo4j memory limits.',
               'specificityScore')
    add_bullet(doc, 'All 26.3M edges now carry a normalized weight \u2208 [0, 1]. Seven sources '
               'use real score normalization (STRING combinedScore/1000, OpenTargets score, '
               'Jensen min-max, Bgee expressionScore min-max, LINCS abs(zScore) min-max, '
               'DoRothEA confidence letter \u00d7 morScore). Binary edges get 0.9 (DrugBank) '
               'or 0.8 (all others).',
               'Edge weights')
    add_bullet(doc, 'Pipeline health check (src/orchestrator.py) dynamically detects parser status '
               'from Neo4j source labels. ID mapping validation report cached at '
               'reports/id_mapping_report.json.',
               'Monitoring')

    doc.add_page_break()

    # =================================================================
    # 3. What Is Planned or Partially Implemented
    # =================================================================
    add_heading(doc, '3. What Is Planned or Partially Implemented', level=1)

    add_heading(doc, 'Working but needs extension', level=2)
    add_bullet(doc, 'Link prediction (Common Neighbors + Weighted Adamic-Adar) is demonstrated '
               'in notebook_class/cardiokb_link_prediction.ipynb for SCN5A and atrial fibrillation. '
               'The notebook normalizes all edge weights and produces ranked gene-disease predictions. '
               'Not yet automated or integrated into the web UI.')
    add_bullet(doc, 'Disease scoping exists for 5 disease areas (CVD, Alzheimer\'s, cancer, asthma, '
               'diabetes) but only CVD has been used in production pipeline runs.')
    add_bullet(doc, 'The DiseaseQueryAgent enriches for any disease interactively, but only sources '
               'DisGeNET + ClinicalTrials.gov. It cannot trigger parser re-runs for the other 34 sources.')

    add_heading(doc, 'Designed but not started', level=2)
    add_bullet(doc, 'Graph embeddings (node2vec, TransE) for downstream ML \u2014 no code exists yet. '
               'The edge weight infrastructure is in place.')
    add_bullet(doc, 'Supervised link prediction classifiers using graph features + embeddings as input. '
               'The notebook demonstrates topological features (CN, Adamic-Adar) but no classifier is trained.')
    add_bullet(doc, 'Drug repurposing pipeline: use link prediction scores to rank novel drug-disease '
               'associations and validate against known treatments. Schema supports this (drugTreatsDisease '
               'from DrugCentral provides ground truth).')
    add_bullet(doc, 'Automated pipeline scheduling with change detection.')
    add_bullet(doc, 'Graph versioning and diff tracking between runs.')
    add_bullet(doc, 'Multi-user web deployment with authentication (currently local-only).')

    add_heading(doc, 'Known gaps in data coverage', level=2)
    add_bullet(doc, 'UniProt protein data \u2014 would connect genes to protein structure and function.')
    add_bullet(doc, 'IntAct \u2014 curated protein interactions, would complement STRING\'s computational scores.')
    add_bullet(doc, 'ChEMBL \u2014 bioactivity data, would strengthen drug-target evidence.')
    add_bullet(doc, 'FDA FAERS \u2014 adverse event reports, would complement SIDER\'s side effect data.')
    add_bullet(doc, 'CTD chemical-to-drug matching is partial (~41% match rate due to MeSH ID coverage). '
               'Improving this would recover ~250K additional drug-gene expression edges.')

    doc.add_page_break()

    # =================================================================
    # 4. What Is Explicitly Out of Scope
    # =================================================================
    add_heading(doc, '4. What Is Explicitly Out of Scope', level=1)
    doc.add_paragraph(
        'These decisions are intentional and should not be revisited without a compelling reason:'
    )
    add_bullet(doc, 'The graph is stored in Neo4j, not RDF/SPARQL. The AlzKB predecessor used '
               'RDF but was migrated to Cypher for performance and simpler tooling. No SPARQL endpoint '
               'is planned.')
    add_bullet(doc, 'The system builds a general-purpose biomedical graph, not a CVD-specific one. '
               'Disease filtering is done at query time (specificity scores, subgraph extraction) '
               'and by 3 disease-aware parsers. The remaining 33 parsers load everything.')
    add_bullet(doc, 'Real-time streaming ingestion. This is a batch system; the pipeline runs to '
               'completion and the graph is queried after. The DiseaseQueryAgent provides near-real-time '
               'enrichment for individual diseases, but it\'s additive, not streaming.')
    add_bullet(doc, 'Production multi-tenant deployment. The web dashboard is single-user, local. '
               'Admin operations (pipeline runs, database agent) are password-gated but there\'s '
               'no user management.')
    add_bullet(doc, 'Natural language querying. The Query tab accepts Cypher directly. '
               'A Claude-powered NL-to-Cypher translator could be built but is not planned for '
               'the 12-week rotation.')
    add_bullet(doc, 'Graph neural networks (GNNs). The ML strategy focuses on topological features '
               'and shallow embeddings (node2vec), not deep learning on graphs. The graph is too large '
               'for full-graph GNN training without sampling infrastructure.')

    doc.add_page_break()

    # =================================================================
    # 5. Proposed Work Division
    # =================================================================
    add_heading(doc, '5. Proposed Work Division: Agent Layer vs. Data Layer', level=1)
    doc.add_paragraph(
        'The system naturally divides into four workstreams that can be developed semi-independently. '
        'This section defines ownership boundaries, interfaces between layers, and what each person '
        'needs to know.'
    )

    # 5.1 Data/Parser Layer
    add_heading(doc, '5.1. Data/Parser Layer', level=2)
    p = doc.add_paragraph()
    p.add_run('Owns: ').bold = True
    p.add_run('src/parsers/, src/ontology_configs.py, src/neo4j_loader.py, src/id_mapping.py, '
              'src/main.py, data/, ontology/')

    p = doc.add_paragraph()
    p.add_run('Responsibilities:').bold = True

    add_bullet(doc, 'Adding new data source parsers (extend BaseParser, add ontology config)')
    add_bullet(doc, 'Fixing ID mapping gaps (improving match rates, adding remap strategies)')
    add_bullet(doc, 'Maintaining the ontology config schema (column mappings, source labels)')
    add_bullet(doc, 'Schema evolution (adding new node/relationship types)')
    add_bullet(doc, 'Expanding disease filter files (new disease areas beyond the 5 existing)')
    add_bullet(doc, 'Pipeline reliability (error handling, partial failure recovery)')

    p = doc.add_paragraph()
    p.add_run('Interface contract: ').bold = True
    p.add_run('Parsers produce Dict[str, DataFrame]. Config keys must match parser output keys. '
              'Every relationship config must include source_label. Node types and relationship types '
              'are defined here and consumed by every other layer.')

    p = doc.add_paragraph()
    p.add_run('What you need to know: ').bold = True
    p.add_run('pandas, Neo4j Cypher basics, the ontology config format (read src/ontology_configs.py '
              'header comments), and the ID mapping system.')

    # 5.2 Agent/Orchestration Layer
    add_heading(doc, '5.2. Agent/Orchestration Layer', level=2)
    p = doc.add_paragraph()
    p.add_run('Owns: ').bold = True
    p.add_run('src/database_agent.py, src/disease_agent.py, src/agent.py, src/orchestrator.py')

    p = doc.add_paragraph()
    p.add_run('Responsibilities:').bold = True

    add_bullet(doc, 'DatabaseAgent improvements (better column detection, handling XML/JSON sources, '
               'multi-file datasets)')
    add_bullet(doc, 'DiseaseQueryAgent expansion (new data sources beyond DisGeNET + ClinicalTrials.gov)')
    add_bullet(doc, 'Pipeline health monitoring and automated recovery')
    add_bullet(doc, 'Prompt engineering for Claude calls (standardization, code generation)')
    add_bullet(doc, 'Caching strategy (DiseaseCache nodes, file-level caching)')

    p = doc.add_paragraph()
    p.add_run('Interface contract: ').bold = True
    p.add_run('Agents produce parsers that conform to the BaseParser interface. DiseaseQueryAgent '
              'loads data using the same Neo4jLoader and MERGE semantics as the batch pipeline. '
              'SSE events follow the format {type, message, data?} for the API layer.')

    p = doc.add_paragraph()
    p.add_run('What you need to know: ').bold = True
    p.add_run('Claude API (tool use, structured outputs), the BaseParser interface, '
              'Neo4j MERGE semantics, SSE streaming.')

    # 5.3 ML/Analysis Layer
    add_heading(doc, '5.3. ML/Analysis Layer', level=2)
    p = doc.add_paragraph()
    p.add_run('Owns: ').bold = True
    p.add_run('notebook_class/, models/, scripts/compute_specificity.py')

    p = doc.add_paragraph()
    p.add_run('Responsibilities:').bold = True

    add_bullet(doc, 'Graph embedding generation (node2vec, TransE, or similar)')
    add_bullet(doc, 'Link prediction model development (beyond the current notebook proof-of-concept)')
    add_bullet(doc, 'Drug repurposing scoring and validation')
    add_bullet(doc, 'Feature engineering from graph topology + edge weights')
    add_bullet(doc, 'Evaluation framework (hold-out edges, time-split validation)')

    p = doc.add_paragraph()
    p.add_run('Interface contract: ').bold = True
    p.add_run('ML code reads from Neo4j via the Python driver. Edge weights are pre-normalized '
              'to [0, 1] and stored as r.weight. Specificity scores are on nodes as n.specificityScore. '
              'Predictions should produce ranked (gene, disease, score) triples that can be loaded '
              'back into Neo4j as a new relationship type or exposed via the API.')

    p = doc.add_paragraph()
    p.add_run('What you need to know: ').bold = True
    p.add_run('Neo4j Python driver, the edge weight normalization scheme (see notebook section 3), '
              'which relationship types are curated vs. text-mined (important for training/validation splits).')

    # 5.4 Interface Layer
    add_heading(doc, '5.4. Interface Layer', level=2)
    p = doc.add_paragraph()
    p.add_run('Owns: ').bold = True
    p.add_run('interface/index.html, src/api.py')

    p = doc.add_paragraph()
    p.add_run('Responsibilities:').bold = True

    add_bullet(doc, 'New API endpoints for ML predictions (e.g., /api/predictions?gene=SCN5A)')
    add_bullet(doc, 'Visualization of prediction results in the Explore tab')
    add_bullet(doc, 'Query template expansion')
    add_bullet(doc, 'Export improvements (subgraph formats for ML tools)')

    p = doc.add_paragraph()
    p.add_run('Interface contract: ').bold = True
    p.add_run('Flask endpoints return JSON. SSE streams use EventSource. The frontend is a '
              'single HTML file with inline JS/CSS. vis.js DataSets must use clear()/add(), '
              'never destroy()/new Network() for refreshes.')

    # 5.5 Dependency Map
    add_heading(doc, '5.5. Dependency Map', level=2)
    doc.add_paragraph('Who blocks whom:')
    add_table(doc,
              ['Task', 'Depends On', 'Blocks'],
              [
                  ['New parser (Data layer)', 'Nothing', 'ML layer (more edges = better predictions)'],
                  ['Schema changes (Data layer)', 'Nothing',
                   'All layers (node/rel types are the shared vocabulary)'],
                  ['ID mapping improvements (Data layer)', 'Nothing',
                   'ML layer (higher match rates = more training data)'],
                  ['Graph embeddings (ML layer)', 'Stable schema + edge weights',
                   'Prediction models, drug repurposing'],
                  ['Link prediction model (ML layer)', 'Embeddings OR topological features',
                   'API predictions endpoint, drug repurposing'],
                  ['Drug repurposing (ML layer)', 'Link prediction model + drugTreatsDisease ground truth',
                   'Nothing (end goal)'],
                  ['DatabaseAgent improvements (Agent layer)', 'BaseParser interface (stable)',
                   'Faster source onboarding'],
                  ['DiseaseQueryAgent expansion (Agent layer)', 'New parsers for target sources',
                   'Interactive enrichment for more data types'],
                  ['Prediction API (Interface layer)', 'Link prediction model output',
                   'Nothing (end-user feature)'],
              ])

    doc.add_page_break()

    # =================================================================
    # 6. Schema Decisions & Edge Merge Recommendations
    # =================================================================
    add_heading(doc, '6. Schema Decisions & Edge Merge Recommendations', level=1)
    doc.add_paragraph(
        'The graph has 20 node types and 43 relationship types. Some edge types overlap semantically '
        'and should be evaluated for merging. Others look similar but carry different semantics that '
        'matter for ML. This section flags each case with a recommendation.'
    )

    # 6.1 Current Schema Overview
    add_heading(doc, '6.1. Current Schema Overview', level=2)
    doc.add_paragraph('20 node types: Gene (194K), Variant (4.5M), Disease (19K), Drug (41K), '
                      'ClinicalTrial (82K), BiologicalProcess (24K), Phenotype (19K), BodyPart (14K), '
                      'MolecularFunction (10K), Pathway (6.5K), SideEffect (5.7K), Species (4.6K), '
                      'CellularComponent (4K), GeneFamily (1.9K), PharmacologicClass (1.6K), '
                      'Symptom (966), DrugLabel (378), TranscriptionFactor (367), AgeingProperty (3), '
                      '_Metadata (1).')

    doc.add_paragraph('43 relationship types across 28 source labels. The top 5 by edge count are: '
                      'bodyPartUnderexpressesGene (5.3M, Bgee), variantInGene/hasVariant (4.4M each, ClinVar), '
                      'geneAssociatesWithDisease (3.7M, 6 sources), diseaseAssociatesWithDisease (2.1M, PubTator).')

    # 6.2 Edges to Merge
    add_heading(doc, '6.2. Edges That Should Be Merged', level=2)

    # drugCausesSideEffect vs compoundCausesSideEffect
    p = doc.add_paragraph()
    p.add_run('MERGE: drugCausesSideEffect + compoundCausesSideEffect \u2192 drugCausesSideEffect').bold = True
    doc.add_paragraph(
        'drugCausesSideEffect (138K edges, Hetionet) and compoundCausesSideEffect (148K edges, SIDER) '
        'describe the exact same relationship: a drug/compound causes a side effect. The only reason '
        'they have different names is that they were loaded from different sources at different times. '
        'Recommendation: Rename compoundCausesSideEffect to drugCausesSideEffect in the SIDER config. '
        'The r.source property already distinguishes Hetionet vs. SIDER provenance. This is the '
        'clearest merge candidate in the schema.'
    )

    # chemicalBindsGene vs drugBindsGene
    p = doc.add_paragraph()
    p.add_run('MERGE: chemicalBindsGene + drugBindsGene \u2192 drugBindsGene').bold = True
    doc.add_paragraph(
        'chemicalBindsGene (4.2K edges, BindingDB) and drugBindsGene (19K edges, DrugBank) '
        'both represent drug/chemical-to-gene binding events. Both connect Drug nodes to Gene nodes. '
        'BindingDB uses "chemical" in its name because its identifiers include non-drug chemicals, '
        'but in CardioKB these are all mapped to Drug nodes. Recommendation: Merge into drugBindsGene. '
        'r.source distinguishes BindingDB vs. DrugBank.'
    )

    # chemicalIncreases/DecreasesExpression vs compoundUp/DownregulatesGene
    p = doc.add_paragraph()
    p.add_run('CONSIDER MERGING: chemical{Increases,Decreases}Expression + '
              'compound{Up,Down}regulatesGene').bold = True
    doc.add_paragraph(
        'CTD contributes chemicalIncreasesExpression (218K) and chemicalDecreasesExpression (213K). '
        'LINCS L1000 contributes compoundUpregulatesGene (4.7K) and compoundDownregulatesGene (5.8K). '
        'These describe the same biology: a chemical/drug changes gene expression levels. The naming '
        'differs because CTD uses "expression" language while LINCS uses "regulation" language. '
        'Recommendation: Merge each pair (increases\u2192up, decreases\u2192down) into a single '
        'relationship type: drugUpregulatesGene and drugDownregulatesGene. LINCS edges carry zScore '
        'properties that CTD edges lack, which is fine \u2014 not all edges need the same properties. '
        'This reduces 4 relationship types to 2.'
    )

    # 6.3 Edges to Keep Separate
    add_heading(doc, '6.3. Edges That Should Stay Separate', level=2)

    p = doc.add_paragraph()
    p.add_run('KEEP SEPARATE: drugTreatsDisease vs. drugPalliatesDisease').bold = True
    doc.add_paragraph(
        'Both come from DrugCentral. drugTreatsDisease (1.3K edges) represents curated treatment '
        'relationships. drugPalliatesDisease (292 edges) represents symptom alleviation without '
        'treating the underlying disease. This distinction matters for drug repurposing: a palliative '
        'drug is not a treatment candidate. Keep both; they serve different ML targets.'
    )

    p = doc.add_paragraph()
    p.add_run('KEEP SEPARATE: geneInteractsWithGene vs. geneRegulatesGene vs. geneCovariesWithGene').bold = True
    doc.add_paragraph(
        'geneInteractsWithGene (234K, STRING + Hetionet) represents physical protein-protein interactions. '
        'geneRegulatesGene (6.3K, LINCS L1000) represents directional regulatory relationships. '
        'geneCovariesWithGene (127, Hetionet) represents coexpression patterns. These are biologically '
        'distinct: interaction \u2260 regulation \u2260 coexpression. Keep all three.'
    )

    p = doc.add_paragraph()
    p.add_run('KEEP SEPARATE: geneAssociatesWithDisease (6 sources)').bold = True
    doc.add_paragraph(
        'This single relationship type aggregates 3.7M edges from PubTator (literature-mined), '
        'OpenTargets (evidence-scored), GWAS Catalog (statistical association), DisGeNET (curated + '
        'text-mined), OMIM (Mendelian), and Jensen DISEASES (text-mined). These should NOT be split '
        'into separate relationship types \u2014 they represent the same semantic relationship from '
        'different evidence sources. The r.source property and per-source scores/weights provide '
        'evidence stratification. For ML, use r.source as a feature or filter by evidence type.'
    )

    p = doc.add_paragraph()
    p.add_run('KEEP SEPARATE: diseaseAssociatesWithDisease vs. diseaseResemblesDisease').bold = True
    doc.add_paragraph(
        'diseaseAssociatesWithDisease (2.1M, PubTator) is literature cooccurrence: "these diseases '
        'appear in the same papers." diseaseResemblesDisease (109, MEDLINE) is semantic similarity. '
        'Different signals, keep both.'
    )

    # 6.4 Naming Cleanup
    add_heading(doc, '6.4. Naming Convention Cleanup', level=2)
    doc.add_paragraph(
        'The schema mixes naming conventions from different eras of development. Some relationships '
        'use SCREAMING_SNAKE (STUDIES_CONDITION, TESTS_INTERVENTION, VARIANT_IN, AFFECTS_RESPONSE_TO) '
        'while most use camelCase. This happened because the ClinicalTrials.gov and ClinPGx parsers '
        'were added later with a different convention.'
    )
    doc.add_paragraph('Recommendation: Rename for consistency (requires config + data migration):')
    add_table(doc,
              ['Current', 'Proposed', 'Source'],
              [
                  ['STUDIES_CONDITION', 'trialStudiesCondition', 'ClinicalTrials.gov'],
                  ['TESTS_INTERVENTION', 'trialTestsIntervention', 'ClinicalTrials.gov'],
                  ['VARIANT_IN', 'variantInGene', 'ClinPGx (note: ClinVar already uses variantInGene)'],
                  ['AFFECTS_RESPONSE_TO', 'geneAffectsResponseTo', 'ClinPGx'],
              ])
    doc.add_paragraph(
        'Priority: Low. These work fine as-is and renaming requires a full pipeline re-run. '
        'But if you\'re doing a schema migration for the merges above, batch these in.'
    )

    doc.add_page_break()

    # =================================================================
    # 7. Knowledge Graph to ML
    # =================================================================
    add_heading(doc, '7. Knowledge Graph \u2192 ML: Hypothesis Generation Pipeline', level=1)
    doc.add_paragraph(
        'The end goal of CardioKB is not the graph itself but what it enables: computational '
        'hypothesis generation for drug repurposing, disease mechanism discovery, and biomarker '
        'identification. This section describes the pipeline from graph to ML predictions.'
    )

    # 7.1 Current State
    add_heading(doc, '7.1. Current State (Working)', level=2)
    doc.add_paragraph(
        'The link prediction notebook (notebook_class/cardiokb_link_prediction.ipynb) demonstrates '
        'the full workflow:'
    )
    add_numbered(doc, 'All 26.3M edges normalized to weight \u2208 [0, 1] using source-specific rules',
                 'Edge weight normalization')
    add_numbered(doc, 'Count shared neighbors between query node and target. '
                 'Applied to SCN5A \u2192 CVD diseases; top hit: cerebrovascular disease (22 shared neighbors)',
                 'Common Neighbors')
    add_numbered(doc, 'Weight shared neighbors inversely by degree and scale by edge weights. '
                 'Applied to SCN5A \u2192 arrhythmia diseases; identifies arrhythmogenic cardiomyopathies',
                 'Weighted Adamic-Adar')
    add_numbered(doc, 'Flipped perspective: given atrial fibrillation, predict missing gene associations. '
                 'Top predictions: FAU, PRKACG, PRKACB, calmodulin-like genes (CALML3/4/5/6)',
                 'Disease-anchored gene prediction')
    add_numbered(doc, 'Confirmed Long QT Syndrome \u2192 SCN5A is already directly linked (14 edges '
                 'from OMIM + OpenTargets), validating the graph captures known biology',
                 'Known-relationship validation')

    # 7.2 Embeddings
    add_heading(doc, '7.2. Planned: Embedding-Based Approaches', level=2)
    doc.add_paragraph(
        'The next step is generating node embeddings that capture graph structure in dense vectors. '
        'These serve as features for any downstream ML task.'
    )
    add_bullet(doc, 'Random walks on the weighted graph produce node sequences; '
               'Word2Vec trains on sequences to produce d-dimensional node vectors. '
               'Captures structural similarity (nodes with similar neighborhoods get similar vectors). '
               'The edge weights are already normalized and ready for biased walks.',
               'node2vec')
    add_bullet(doc, 'Knowledge graph embedding that represents entities and relations as vectors. '
               'Learns h + r \u2248 t for each (head, relation, tail) triple. '
               'Better at capturing relation-specific patterns than node2vec. '
               'The 43 relationship types provide rich relational signal.',
               'TransE / RotatE')
    add_bullet(doc, 'models/ directory (currently empty) should store trained embedding files, '
               'model checkpoints, and evaluation metrics.',
               'Storage')

    # 7.3 Feature Engineering
    add_heading(doc, '7.3. Planned: Feature Engineering for Classifiers', level=2)
    doc.add_paragraph(
        'For a supervised link prediction classifier (e.g., predict whether a gene-disease link exists), '
        'combine topological features with embeddings:'
    )
    add_table(doc,
              ['Feature Type', 'Features', 'Source'],
              [
                  ['Topological', 'Common Neighbors, Adamic-Adar, Jaccard coefficient, '
                   'preferential attachment, Katz index', 'Graph structure'],
                  ['Node properties', 'specificityScore, node degree, degree by relationship type',
                   'Neo4j node properties'],
                  ['Edge properties', 'r.weight, r.source counts, max/mean weight by source',
                   'Neo4j edge properties'],
                  ['Embeddings', 'node2vec vectors, TransE entity vectors',
                   'Trained embedding models'],
                  ['Source diversity', 'Number of distinct r.source labels on existing edges, '
                   'coverage across evidence types', 'r.source aggregation'],
              ])
    doc.add_paragraph(
        'Training data: use existing geneAssociatesWithDisease edges as positives, '
        'random (gene, disease) pairs with no edge as negatives. For drug repurposing: '
        'drugTreatsDisease edges as positives. Use time-split validation if edge timestamps '
        'are available (OpenTargets and GWAS provide date fields).'
    )

    # 7.4 Integration Points
    add_heading(doc, '7.4. Integration Points', level=2)
    doc.add_paragraph(
        'Where the ML layer connects to the rest of the system:'
    )
    add_table(doc,
              ['Integration Point', 'Direction', 'Details'],
              [
                  ['Neo4j \u2192 ML', 'Read', 'ML reads graph via Python driver. '
                   'All edges have r.weight (float, 0\u20131) and r.source (string). '
                   'All nodes have n.specificityScore (float, 0\u20131).'],
                  ['ML \u2192 Neo4j', 'Write', 'Prediction scores can be loaded as new relationship '
                   'properties or a new relationship type (e.g., predictedAssociation with score property).'],
                  ['ML \u2192 API', 'Via Neo4j', 'Once predictions are in the graph, existing API endpoints '
                   '(/api/graph, /api/query) can serve them. New endpoint /api/predictions could wrap '
                   'a Cypher query that filters by prediction score.'],
                  ['ML \u2192 Web UI', 'Via API', 'Predicted edges could appear in the Explore tab '
                   'with a distinctive color/style. Discovery layer already supports 2-hop exploration.'],
                  ['Edge weights \u2192 ML', 'Read', 'Normalization scheme is documented in notebook '
                   'section 3. Seven sources have real scores; others get flat weights.'],
                  ['Drug repurposing ground truth', 'Read', 'drugTreatsDisease (1.3K edges, DrugCentral) '
                   'and drugPalliatesDisease (292 edges) provide labeled examples for evaluation.'],
              ])

    doc.add_page_break()

    # =================================================================
    # 8. Current Stats
    # =================================================================
    add_heading(doc, '8. Current Stats (Snapshot)', level=1)
    doc.add_paragraph(
        'These numbers are from the most recent pipeline run. For live counts, query Neo4j directly '
        'or use GET /api/graph-stats.'
    )
    add_table(doc,
              ['Metric', 'Value'],
              [
                  ['Total nodes', '4,921,062'],
                  ['Total relationships', '26,344,399'],
                  ['Node types', '20'],
                  ['Relationship types', '43'],
                  ['Data sources (parsers)', '36'],
                  ['Relationship source labels', '28'],
                  ['Ontology configs', '85'],
                  ['Disease filter terms (CVD)', '90'],
                  ['Agent-generated parsers', '7'],
                  ['Edge weight coverage', '100% (all edges have r.weight)'],
              ])

    doc.add_paragraph('Top 10 relationship types by count:')
    add_table(doc,
              ['Relationship Type', 'Count', 'Source(s)'],
              [
                  ['bodyPartUnderexpressesGene', '5,334,316', 'Bgee'],
                  ['variantInGene', '4,439,480', 'ClinVar'],
                  ['hasVariant', '4,439,480', 'ClinVar'],
                  ['geneAssociatesWithDisease', '3,706,670', '6 sources'],
                  ['diseaseAssociatesWithDisease', '2,138,895', 'PubTator'],
                  ['variantAssociatedWithDisease', '1,862,448', 'ClinVar'],
                  ['associatedWithVariant', '1,862,448', 'ClinVar'],
                  ['geneExpressedInBodyPart', '982,039', 'Jensen TISSUES'],
                  ['geneInteractsWithGene', '234,533', 'STRING, Hetionet'],
                  ['chemicalIncreasesExpression', '218,140', 'CTD'],
              ])

    # Save
    output_path = Path('/Users/nawaza/Desktop/Cardio-KB/docs/CardioKB_System_Design.docx')
    doc.save(str(output_path))
    print(f"Saved: {output_path}")
    return output_path


if __name__ == '__main__':
    build_doc()
