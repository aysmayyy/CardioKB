"""Generate the CardioKB Redundancy Removal & CVD Pivot changelog Word document."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()
    return table


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        p.add_run(f'{bold_prefix}: ').bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def build_doc():
    doc = Document()

    # Title
    doc.add_paragraph()
    title = doc.add_heading(
        'CardioKB: Redundancy Removal & CVD Pivot Changelog', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        'Version control document tracking source deduplication '
        'and the transition from disease-agnostic to CVD-focused knowledge graph'
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('April 2026').font.size = Pt(11)

    doc.add_page_break()

    # ─── 1. Summary ───
    add_heading(doc, '1. Summary of Changes', level=1)
    doc.add_paragraph(
        'This document records two related changes made to the CardioKB schema and data pipeline:'
    )
    add_bullet(doc, 'Removed 10 redundant data sources (including Hetionet precomputed dump). '
               'Every node type and edge type is now served by exactly one authoritative database. '
               'Source count reduced from 36 to 26.',
               'Redundancy removal')
    add_bullet(doc, 'Officially scoped the knowledge graph to cardiovascular disease. '
               'Created CVD-specific ontology files for genes (3,984 symbols from OMIM + DisGeNET, cleaned), '
               'node types (19), edge types (41), and diseases (184 terms, expanded).',
               'CVD pivot')
    add_bullet(doc, 'Identified 3 stale sources (SIDER, LINCS L1000, MEDLINE) pinned to archived '
               'GitHub commits. Proposed live replacements: DrugBank adverse reactions, clue.io API, '
               'PubTator cooccurrence expansion.',
               'Stale source analysis')

    # ─── 2. Sources Removed ───
    add_heading(doc, '2. Sources Removed (10)', level=1)

    add_heading(doc, '2.1. Removed Entirely (10 sources)', level=2)
    add_table(doc,
              ['Source', 'Was Providing', 'Replaced By', 'Rationale'],
              [
                  ['DisGeNET', 'Disease nodes + geneAssociatesWithDisease (20K edges)',
                   'Disease Ontology (nodes), OpenTargets (edges)',
                   'OpenTargets has 2.4M edges with evidence scores; Disease Ontology is the canonical disease ontology'],
                  ['GWAS Catalog', 'geneAssociatesWithDisease (45K edges)',
                   'OpenTargets',
                   'OpenTargets already ingests GWAS Catalog directly'],
                  ['Jensen DISEASES', 'geneAssociatesWithDisease (20.5K edges)',
                   'OpenTargets',
                   'OpenTargets covers text-mining with better scoring'],
                  ['OMIM', 'geneAssociatesWithDisease (7.3K edges)',
                   'OpenTargets',
                   'OpenTargets includes genetic evidence; OMIM gene data retained in CVD gene ontology file'],
                  ['WikiPathways', 'Pathway nodes + geneInPathway (8.6K edges)',
                   'Reactome',
                   'Reactome is gold-standard curated; many WikiPathways imported from Reactome'],
                  ['AOP-DB', 'Pathway nodes + geneInPathway (18.5K edges)',
                   'Reactome',
                   'AOP-DB focuses on toxicology pathways, less relevant for CVD; Reactome has broader coverage'],
                  ['HGNC', 'Gene nodes (enrichment)',
                   'NCBI Gene',
                   'NCBI Gene is the primary gene reference with daily updates'],
                  ['CellAge', 'Gene nodes (senescence genes)',
                   'NCBI Gene',
                   'Node-only source, senescence genes already in NCBI Gene'],
                  ['GenAge', 'Gene nodes (aging genes)',
                   'NCBI Gene',
                   'Node-only source, aging genes already in NCBI Gene'],
                  ['Hetionet (precomputed)', 'drugCausesSideEffect (138K) + geneInteractsWithGene (5.1K) + geneCovariesWithGene (127)',
                   'SIDER (side effects), STRING (PPI)',
                   'drugCausesSideEffect redundant with SIDER compoundCausesSideEffect; PPI redundant with STRING; geneCovariesWithGene only 127 edges, dropped'],
              ])

    # ─── 3. Sources Modified ───
    add_heading(doc, '3. Sources Modified (2)', level=1)
    add_table(doc,
              ['Source', 'Change', 'Reason'],
              [
                  ['PubTator Central', 'Removed geneAssociatesWithDisease (1.2M edges). '
                   'Kept diseaseAssociatesWithDisease (2.1M edges, unique).',
                   'OpenTargets covers gene-disease; PubTator uniquely provides disease-disease cooccurrence'],
                  ['ClinPGx', 'Removed Variant from node contribution. '
                   'Kept DrugLabel nodes and all 4 unique edge types.',
                   'ClinVar is primary Variant source (4.5M vs 1.1K). ClinPGx pharmacogenomics edges are unique'],
              ])

    # ─── 4. Final Source Assignments ───
    add_heading(doc, '4. Final Source Assignments (26 sources)', level=1)
    doc.add_paragraph(
        'After deduplication, each node type and edge type has exactly one authoritative source:'
    )

    add_heading(doc, '4.1. Node Type Assignments', level=2)
    add_table(doc,
              ['Node Type', 'Authoritative Source', 'Count'],
              [
                  ['Gene', 'NCBI Gene', '194,553'],
                  ['Disease', 'Disease Ontology', '12,012'],
                  ['Drug', 'DrugBank + CTD', '24,414'],
                  ['Variant', 'ClinVar', '4,488,042'],
                  ['ClinicalTrial', 'ClinicalTrials.gov', '85,677'],
                  ['Pathway', 'Reactome', '2,806'],
                  ['BiologicalProcess', 'Gene Ontology', '24,547'],
                  ['MolecularFunction', 'Gene Ontology', '10,123'],
                  ['CellularComponent', 'Gene Ontology', '4,069'],
                  ['BodyPart', 'Uberon', '14,937'],
                  ['Phenotype', 'HPO', '19,389'],
                  ['Symptom', 'NCBI MeSH', '966'],
                  ['SideEffect', 'SIDER', '5,734'],
                  ['TranscriptionFactor', 'DoRothEA', '367'],
                  ['PharmacologicClass', 'DrugCentral', '1,646'],
                  ['GeneFamily', 'HGNC Families', '1,934'],
                  ['DrugLabel', 'ClinPGx', '378'],
                  ['AgeingProperty', 'DrugAge', '3'],
                  ['Species', 'AnAge', '4,645'],
              ])

    add_heading(doc, '4.2. Edge Type Assignments', level=2)
    add_table(doc,
              ['Edge Type', 'Authoritative Source', 'Previous Sources'],
              [
                  ['geneAssociatesWithDisease', 'OpenTargets',
                   'Was 6 sources: OpenTargets, DisGeNET, GWAS, PubTator, Jensen, OMIM'],
                  ['geneInteractsWithGene', 'STRING',
                   'Was 2: STRING, Hetionet. Hetionet fully removed.'],
                  ['geneInPathway / pathwayContainsGene', 'Reactome',
                   'Was 3: Reactome, WikiPathways, AOP-DB'],
                  ['diseaseAssociatesWithDisease', 'PubTator', 'No change (was sole source)'],
                  ['compoundCausesSideEffect', 'SIDER',
                   'drugCausesSideEffect (Hetionet) removed; SIDER is sole side-effect source'],
                  ['geneCovariesWithGene', 'REMOVED',
                   'Was Hetionet only (127 edges); dropped with Hetionet removal'],
                  ['All other edge types', '(unchanged)', 'Each already had a single source'],
              ])

    doc.add_page_break()

    # ─── 5. Impact ───
    add_heading(doc, '5. Impact Assessment', level=1)

    add_heading(doc, '5.1. Edge Count Changes', level=2)
    add_table(doc,
              ['Edge Type', 'Before (all sources)', 'After (single source)', 'Edges Removed'],
              [
                  ['geneAssociatesWithDisease', '3,706,670', '2,364,224 (OpenTargets)', '1,342,446'],
                  ['geneInteractsWithGene', '234,533', '229,433 (STRING)', '5,100'],
                  ['geneInPathway', '43,383', '16,317 (Reactome)', '27,066'],
                  ['pathwayContainsGene', '43,383', '16,317 (Reactome)', '27,066'],
              ])

    doc.add_paragraph(
        'Total edges removed from redundancy: ~1.4M. These were duplicate assertions of the same '
        'biological relationships from lower-priority sources. The r.source property on remaining '
        'edges unambiguously identifies the authoritative database.'
    )

    add_heading(doc, '5.2. Node Count Changes', level=2)
    doc.add_paragraph(
        'Pathway nodes decreased from 6,469 to 2,806 (Reactome only, losing WikiPathways and '
        'AOP-DB pathway definitions). Disease nodes: 12,012 (Disease Ontology only, no more '
        'DrugCentral CUI orphan nodes). Drug nodes: 24,414 (19,842 DrugBank + 4,572 CTD unique). '
        'Variant count unchanged at 4,488,042 (ClinVar). '
        'Note: edge counts above are pre-CVD-filtering; final loaded counts are lower after '
        'CVD gene filtering (see edge_types.txt for current counts).'
    )

    # ─── 6. Stale Source Analysis ───
    add_heading(doc, '6. Stale Source Analysis', level=1)
    doc.add_paragraph(
        'Three remaining sources are pinned to archived GitHub commits and are no longer updated '
        'by their maintainers. They remain in the pipeline but are flagged for replacement:'
    )
    add_table(doc,
              ['Source', 'Issue', 'Replacement Strategy', 'Status'],
              [
                  ['SIDER', 'Pinned to 2015 GitHub commit (dhimmel/SIDER4). '
                   'SIDER website discontinued.',
                   'Extend DrugBank parser to extract adverse reactions from XML',
                   'Planned'],
                  ['LINCS L1000', 'Pinned to 2020 GitHub commit (dhimmel/lincs). '
                   'Original L1000 data is from 2017.',
                   'Replace with clue.io REST API for live L1000 data',
                   'Planned'],
                  ['MEDLINE cooccurrence', 'Pinned to GitHub commit (dhimmel/medline). '
                   'Literature cooccurrence from static dump.',
                   'Expand PubTator Central parser to include disease-symptom '
                   'and disease-anatomy cooccurrence',
                   'Planned'],
              ])

    # ─── 7. CVD Ontology Files ───
    add_heading(doc, '7. New CVD Ontology Files', level=1)
    doc.add_paragraph(
        'The following ontology files were created to formalize the CVD scope:'
    )
    add_table(doc,
              ['File', 'Contents', 'Count'],
              [
                  ['ontology/schema/node_types.txt',
                   'All 19 node types with primary keys, source databases, descriptions', '19 types'],
                  ['ontology/schema/edge_types.txt',
                   'All 43 edge types with source/target node types, source databases, edge counts',
                   '43 types'],
                  ['ontology/genes/cvd.txt',
                   'CVD-associated gene symbols merged from OMIM + DisGeNET, cleaned of LOC* and OMIM phenotype symbols',
                   '3,984 genes'],
                  ['ontology/diseases/cvd.txt',
                   'CVD disease terms for filtering (expanded with congenital, pulmonary vascular, and more)',
                   '184 terms'],
              ])

    # ─── 8. Visualization ───
    add_heading(doc, '8. Updated Database Visualization', level=1)
    doc.add_paragraph(
        'The database_visualization/ directory was updated to reflect CardioKB instead of AlzKB:'
    )
    add_bullet(doc, 'alzkb_databases.csv replaced with cardiokb_databases.csv (26 sources)',
               'CSV')
    add_bullet(doc, 'alzkb_source_schema_template.html replaced with cardiokb_source_schema_template.html '
               '(19 node types, 43 edge types, 3 integration categories: Direct/Hetionet/Agent)',
               'Template')
    add_bullet(doc, 'build_latest_schema.py updated with CardioKB file paths and all 41 edge + 19 node mappings',
               'Build script')
    add_bullet(doc, 'cardiokb_source_schema_latest.html generated (open in browser to view interactive D3 graph)',
               'Output')

    # ─── 9. Migration Notes ───
    add_heading(doc, '9. Migration Notes', level=1)
    doc.add_paragraph(
        'The following steps are required to apply these changes to the live Neo4j database:'
    )
    p = doc.add_paragraph(style='List Number')
    p.add_run('Remove redundant parsers from src/main.py: ').bold = True
    p.add_run('DisGeNETParser, GWASParser, JensenLabParser, OMIMParser (edge loading only), '
              'WikiPathwaysParser, AOPDBParser, HGNCParser, CellAgeParser, GenAgeParser, '
              'HetionetPrecomputedParser')

    p = doc.add_paragraph(style='List Number')
    p.add_run('Update src/ontology_configs.py: ').bold = True
    p.add_run('Set skip=True or remove configs for the 10 dropped sources. '
              'Remove geneAssociatesWithDisease from PubTator config.')

    p = doc.add_paragraph(style='List Number')
    p.add_run('Clean Neo4j: ').bold = True
    p.add_run('Delete edges from removed sources: '
              'MATCH ()-[r]->() WHERE r.source IN ["DisGeNET","GWAS Catalog","Jensen DISEASES",'
              '"OMIM","WikiPathways","AOP-DB","Hetionet"] DELETE r')

    p = doc.add_paragraph(style='List Number')
    p.add_run('Re-run pipeline: ').bold = True
    p.add_run('python src/main.py --skip-download to rebuild with only the 26 authoritative sources')

    p = doc.add_paragraph(style='List Number')
    p.add_run('Verify: ').bold = True
    p.add_run('Run health check to confirm all 26 sources are loaded with expected edge counts')

    # Save
    output_path = Path('/Users/nawaza/Desktop/Cardio-KB/docs/CardioKB_Redundancy_Changelog.docx')
    doc.save(str(output_path))
    print(f"Saved: {output_path}")
    return output_path


if __name__ == '__main__':
    build_doc()
