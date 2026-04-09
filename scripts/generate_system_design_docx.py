"""Generate CardioKB_System_Design.docx from the markdown content."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

# ============================================================
# CONTENT
# ============================================================

doc.add_heading('CardioKB System Design', level=0)

doc.add_heading('1. System Overview', level=1)
doc.add_paragraph(
    'CardioKB is a CVD-focused biomedical knowledge graph integrating 26 deduplicated '
    'data sources into Memgraph. The system consists of four main components:'
)
items = [
    'ETL Pipeline \u2014 Downloads, parses, and loads biomedical data into the graph',
    'Graph Database \u2014 Memgraph instance storing 4.9M nodes and 7.7M relationships',
    'Web Interface \u2014 Flask backend + vis.js frontend for exploration and querying',
    'AI Agents \u2014 DatabaseAgent (parser generation) and DiseaseQueryAgent (on-demand enrichment)',
]
for item in items:
    doc.add_paragraph(item, style='List Number')

# Architecture diagram as monospace
p = doc.add_paragraph()
p.style = doc.styles['Normal']
arch = (
    "26 Data Sources (APIs, FTP, XML, TSV)\n"
    "        \u2193\n"
    "ETL Pipeline (main.py) \u2192 Parsers \u2192 TSV \u2192 Loader\n"
    "        \u2193\n"
    "Memgraph (Bolt:7687) \u2190 DatabaseAgent (Claude API)\n"
    "        \u2193\n"
    "Flask API (:5050) \u2190 DiseaseQueryAgent (ClinicalTrials.gov API v2)\n"
    "        \u2193\n"
    "Web Dashboard (index.html, vis.js)"
)
run = p.add_run(arch)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# -- 2. Graph Statistics --
doc.add_heading('2. Graph Statistics', level=1)
add_table(
    ['Metric', 'Value'],
    [
        ['Total nodes', '4,896,258'],
        ['Total relationships', '7,683,150'],
        ['Node types', '19'],
        ['Relationship types', '43'],
        ['Data sources', '26'],
        ['Source labels on edges', '23'],
        ['Ontology configs', '86'],
    ]
)

# -- 3. Source-to-Schema Mapping --
doc.add_heading('3. Source-to-Schema Mapping', level=1)

doc.add_heading('3.1 Node Type Ownership', level=2)
doc.add_paragraph(
    'Each node type has exactly one authoritative source. All nodes carry a '
    'specificityScore property (1.0 / count(Disease neighbors); Disease nodes = 0.0).'
)
add_table(
    ['Node Type', 'Count', 'Authoritative Source', 'Key Properties'],
    [
        ['Variant', '4,488,042', 'ClinVar', 'variantId, commonName, chromosome, position, changeClassification, gene'],
        ['Gene', '194,553', 'NCBI Gene', 'xrefNcbiGene, geneSymbol, commonName, chromosome, typeOfGene, xrefEnsembl, xrefHGNC, xrefOMIM'],
        ['ClinicalTrial', '85,691', 'ClinicalTrials.gov', 'trialId, commonName, condition, interventionName, phase, status'],
        ['BiologicalProcess', '24,547', 'Gene Ontology', 'geneOntologyId, commonName, definition'],
        ['Drug', '24,414', 'DrugBank + CTD', 'xrefDrugbank, commonName, xrefCasRN'],
        ['Phenotype', '19,389', 'HPO', 'xrefHPO, commonName, definition, synonyms'],
        ['BodyPart', '14,937', 'Uberon', 'xrefUberon, commonName, definition'],
        ['Disease', '12,012', 'Disease Ontology', 'xrefDiseaseOntology, commonName, definition'],
        ['MolecularFunction', '10,123', 'Gene Ontology', 'geneOntologyId, commonName, definition'],
        ['SideEffect', '5,734', 'SIDER', 'xrefUmlsCUI, commonName'],
        ['Species', '4,645', 'AnAge', 'speciesName, commonName, maximumLifespan, sampleSize'],
        ['CellularComponent', '4,069', 'Gene Ontology', 'geneOntologyId, commonName, definition'],
        ['Pathway', '2,806', 'Reactome', 'pathwayName'],
        ['GeneFamily', '1,934', 'HGNC Families', 'familyId, familyName'],
        ['PharmacologicClass', '1,646', 'DrugCentral', 'classId, classType, commonName'],
        ['Symptom', '966', 'NCBI MeSH', 'xrefMeSH, commonName, meshTreeNumber'],
        ['DrugLabel', '378', 'ClinPGx', 'labelId, commonName, drug, gene, regulatorySource, testing'],
        ['TranscriptionFactor', '367', 'DoRothEA', 'TF'],
        ['AgeingProperty', '3', 'DrugAge', 'propertyName'],
    ]
)

doc.add_heading('3.2 Edge Type Ownership', level=2)
doc.add_paragraph(
    'Each row shows the source database, the edge it contributes, the node types '
    'it connects, and the current count in the graph.'
)
add_table(
    ['Source', 'Relationship', 'From \u2192 To', 'Count', 'Edge Properties'],
    [
        ['ClinVar', 'hasVariant', 'Gene \u2192 Variant', '2,267,095', '\u2014'],
        ['ClinVar', 'variantInGene', 'Variant \u2192 Gene', '2,267,095', '\u2014'],
        ['ClinVar', 'variantAssociatedWithDisease', 'Variant \u2192 Disease', '99,707', '\u2014'],
        ['ClinVar', 'associatedWithVariant', 'Disease \u2192 Variant', '99,707', '\u2014'],
        ['Bgee', 'bodyPartUnderexpressesGene', 'BodyPart \u2192 Gene', '784,026', 'expressionScore'],
        ['Bgee', 'bodyPartOverexpressesGene', 'BodyPart \u2192 Gene', '1,872', 'expressionScore'],
        ['OpenTargets', 'geneAssociatesWithDisease', 'Gene \u2192 Disease', '103,879', '\u2014'],
        ['PubTator', 'geneAssociatesWithDisease', 'Gene \u2192 Disease', '673,374', '\u2014'],
        ['PubTator', 'diseaseAssociatesWithDisease', 'Disease \u2192 Disease', '4,320', '\u2014'],
        ['Jensen TISSUES', 'geneExpressedInBodyPart', 'Gene \u2192 BodyPart', '215,235', '\u2014'],
        ['HPO', 'geneAssociatesWithPhenotype', 'Gene \u2192 Phenotype', '162,994', '\u2014'],
        ['LINCS L1000', 'geneRegulatesGene', 'Gene \u2192 Gene', '150,540', 'zScore'],
        ['LINCS L1000', 'compoundUpregulatesGene', 'Drug \u2192 Gene', '10,278', 'zScore'],
        ['LINCS L1000', 'compoundDownregulatesGene', 'Drug \u2192 Gene', '10,218', 'zScore'],
        ['SIDER', 'compoundCausesSideEffect', 'Drug \u2192 SideEffect', '148,518', '\u2014'],
        ['STRING', 'geneInteractsWithGene', 'Gene \u2192 Gene', '121,170', 'confidence'],
        ['CTD', 'chemicalIncreasesExpression', 'Drug \u2192 Gene', '116,451', '\u2014'],
        ['CTD', 'chemicalDecreasesExpression', 'Drug \u2192 Gene', '97,951', '\u2014'],
        ['Gene Ontology', 'geneParticipatesInBiologicalProcess', 'Gene \u2192 BP', '50,350', '\u2014'],
        ['Gene Ontology', 'geneHasMolecularFunction', 'Gene \u2192 MF', '26,935', '\u2014'],
        ['Gene Ontology', 'geneAssociatedWithCellularComponent', 'Gene \u2192 CC', '25,794', '\u2014'],
        ['Reactome', 'geneInPathway', 'Gene \u2192 Pathway', '44,979', '\u2014'],
        ['Reactome', 'pathwayContainsGene', 'Pathway \u2192 Gene', '44,979', '\u2014'],
        ['ClinicalTrials.gov', 'STUDIES_CONDITION', 'ClinicalTrial \u2192 Disease', '27,866', '\u2014'],
        ['ClinicalTrials.gov', 'TESTS_INTERVENTION', 'ClinicalTrial \u2192 Drug', '17,492', '\u2014'],
        ['NCBI Gene', 'geneInSpecies', 'Gene \u2192 Species', '26,417', '\u2014'],
        ['DrugCentral', 'pharmacologicClassIncludesCompound', 'PharmClass \u2192 Drug', '16,403', '\u2014'],
        ['DrugCentral', 'compoundInPharmacologicClass', 'Drug \u2192 PharmClass', '16,403', '\u2014'],
        ['DrugCentral', 'drugTreatsDisease', 'Drug \u2192 Disease', '245', '\u2014'],
        ['DrugCentral', 'drugPalliatesDisease', 'Drug \u2192 Disease', '96', '\u2014'],
        ['DoRothEA', 'transcriptionFactorInteractsWithGene', 'TF \u2192 Gene', '12,985', 'morScore, confidence'],
        ['BindingDB', 'chemicalBindsGene', 'Drug \u2192 Gene', '12,250', '\u2014'],
        ['DrugBank', 'drugBindsGene', 'Drug \u2192 Gene', '12,089', '\u2014'],
        ['HGNC', 'geneInFamily', 'Gene \u2192 GeneFamily', '5,123', '\u2014'],
        ['HGNC', 'familyContainsGene', 'GeneFamily \u2192 Gene', '5,123', '\u2014'],
        ['ClinPGx', 'VARIANT_IN', 'Variant \u2192 Gene', '1,091', '\u2014'],
        ['ClinPGx', 'drugLabelAnnotatesGene', 'DrugLabel \u2192 Gene', '503', '\u2014'],
        ['ClinPGx', 'drugLabelDescribesDrug', 'DrugLabel \u2192 Drug', '345', '\u2014'],
        ['ClinPGx', 'AFFECTS_RESPONSE_TO', 'Gene \u2192 Drug/PharmClass', '243', '\u2014'],
        ['DrugAge', 'associatedWithAging', 'Gene \u2192 AgeingProperty', '386', '\u2014'],
        ['Disease Ontology', 'diseaseIsSubtypeOf', 'Disease \u2192 Disease', '258', '\u2014'],
        ['MEDLINE', 'diseaseLocalizesToAnatomy', 'Disease \u2192 BodyPart', '244', '\u2014'],
        ['MEDLINE', 'diseasePresentsSymptom', 'Disease \u2192 Symptom', '117', '\u2014'],
        ['MEDLINE', 'diseaseResemblesDisease', 'Disease \u2192 Disease', '4', '\u2014'],
    ]
)

doc.add_heading('3.3 Per-Source Summary', level=2)
add_table(
    ['#', 'Source', 'Parser', 'Access', 'Nodes Contributed', 'Edges Contributed', 'Total Edges'],
    [
        ['1', 'ClinicalTrials.gov', 'ClinicalTrialsParser', 'Public API v2', 'ClinicalTrial (85,691)', 'STUDIES_CONDITION, TESTS_INTERVENTION', '45,358'],
        ['2', 'ClinPGx', 'ClinPGxParser', 'Public API', 'DrugLabel (378)', '4 edge types', '2,182'],
        ['3', 'NCBI Gene', 'NCBIGeneParser', 'Public FTP', 'Gene (194,553)', 'geneInSpecies', '26,417'],
        ['4', 'DoRothEA', 'DoRothEAParser', 'Public API', 'TranscriptionFactor (367)', 'transcriptionFactorInteractsWithGene', '12,985'],
        ['5', 'DrugBank', 'DrugBankParser', 'XML file', 'Drug (19,842)', 'drugBindsGene', '12,089'],
        ['6', 'Disease Ontology', 'DiseaseOntologyParser', 'Public', 'Disease (12,012)', 'diseaseIsSubtypeOf', '258'],
        ['7', 'Gene Ontology', 'GeneOntologyParser', 'Public', 'BP (24,547), MF (10,123), CC (4,069)', '3 edge types', '103,079'],
        ['8', 'Uberon', 'UberonParser', 'Public', 'BodyPart (14,937)', '\u2014', '0'],
        ['9', 'NCBI MeSH', 'MeSHParser', 'Public', 'Symptom (966)', '\u2014', '0'],
        ['10', 'SIDER', 'SIDERParser', 'Public', 'SideEffect (5,734)', 'compoundCausesSideEffect', '148,518'],
        ['11', 'LINCS L1000', 'LINCS1000Parser', 'Public', '\u2014', '3 edge types', '171,036'],
        ['12', 'MEDLINE', 'MEDLINECooccurrenceParser', 'Public', '\u2014', '3 edge types', '365'],
        ['13', 'DrugCentral', 'DrugCentralParser', 'Public', 'PharmacologicClass (1,646)', '4 edge types', '33,147'],
        ['14', 'BindingDB', 'BindingDBParser', 'Public', '\u2014', 'chemicalBindsGene', '12,250'],
        ['15', 'PubTator', 'PubTatorParser', 'Public FTP', '\u2014', '2 edge types', '677,694'],
        ['16', 'CTD', 'CTDParser', 'Public', 'Drug (4,572 unique)', '2 edge types', '214,402'],
        ['17', 'Bgee', 'BgeeParser', 'Public FTP', '\u2014', '2 edge types', '785,898'],
        ['18', 'Jensen TISSUES', 'JensenTissuesParser', 'Public', '\u2014', 'geneExpressedInBodyPart', '215,235'],
        ['19', 'HPO', 'HPOParser', 'Public', 'Phenotype (19,389)', 'geneAssociatesWithPhenotype', '162,994'],
        ['20', 'Reactome', 'ReactomeParser', 'Public', 'Pathway (2,806)', '2 edge types', '89,958'],
        ['21', 'STRING', 'STRINGParser', 'Public', '\u2014', 'geneInteractsWithGene', '121,170'],
        ['22', 'OpenTargets', 'OpenTargetsParser', 'Public', '\u2014', 'geneAssociatesWithDisease', '103,879'],
        ['23', 'HGNC Families', 'HGNCFamiliesParser', 'Public', 'GeneFamily (1,934)', '2 edge types', '10,246'],
        ['24', 'ClinVar', 'ClinVarParser', 'Public FTP', 'Variant (4,488,042)', '4 edge types', '4,733,604'],
        ['25', 'DrugAge', 'DrugAgeParser', 'Public', 'AgeingProperty (3)', 'associatedWithAging', '386'],
        ['26', 'AnAge', 'AnAgeParser', 'Public', 'Species (4,645)', '\u2014', '0'],
    ]
)

# -- 4. ETL Pipeline Architecture --
doc.add_heading('4. ETL Pipeline Architecture', level=1)

doc.add_heading('4.1 Pipeline Flow', level=2)
flow = (
    "main.py [--skip-download] [--skip-neo4j]\n"
    "  Phase 1: Download (skippable) \u2192 data/raw/<source>/\n"
    "  Phase 2: Parse \u2192 pandas DataFrames\n"
    "  Phase 3: TSV Export \u2192 data/processed/<source>/*.tsv\n"
    "  Phase 4: Graph Load (skippable) \u2192 UNWIND Cypher batching, MERGE, r.source\n"
    "  Phase 5: Post-processing \u2192 compute_specificity.py"
)
p = doc.add_paragraph()
run = p.add_run(flow)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('4.2 Parser Architecture', level=2)
doc.add_paragraph(
    'All 26 parsers inherit from BaseParser (src/parsers/base_parser.py) with three methods: '
    'download(), parse(), and export_tsv().'
)
doc.add_paragraph('Direct (5): Custom parsers hitting live APIs/files', style='List Bullet')
doc.add_paragraph('Hetionet-derived (17): Parse from component files or original source data', style='List Bullet')
doc.add_paragraph('Agent-generated (4): Created by DatabaseAgent (HGNC Families, ClinVar, DrugAge, AnAge)', style='List Bullet')

doc.add_heading('4.3 Ontology Configs', level=2)
doc.add_paragraph(
    '86 entries in src/ontology_configs.py map TSV files to the graph schema. '
    'Node configs define label, file pattern, ID field, and properties. '
    'Relationship configs define rel_type, source/target node labels, ID fields, '
    'properties, and a source_label field that the loader sets as r.source on every relationship.'
)

doc.add_heading('4.4 ID Harmonization', level=2)
add_table(
    ['Mapping', 'Purpose'],
    [
        ['MeSH \u2192 DOID', 'PubTator disease IDs to Disease Ontology'],
        ['EFO \u2192 DOID', 'OpenTargets disease IDs to Disease Ontology'],
        ['CUI \u2192 DOID', 'DrugCentral disease IDs to Disease Ontology'],
        ['DrugBank ID merging', 'CTD chemicals matched to existing DrugBank Drug nodes'],
    ]
)

# -- 5. CVD Disease Scoping --
doc.add_heading('5. CVD Disease Scoping', level=1)
doc.add_paragraph(
    'ontology/disease_filter.txt is a symlink to ontology/diseases/cvd.txt (184 CVD terms). '
    'The CVD AND-filter applies strict disease scoping with word-boundary matching to:'
)
doc.add_paragraph('OpenTargets: EFO-to-DOID mapped, filtered to CVD \u2192 103,879 edges', style='List Bullet')
doc.add_paragraph('PubTator: Literature-mined, filtered to CVD scope \u2192 677,694 edges', style='List Bullet')
doc.add_paragraph('ClinVar: Variant-disease associations, CVD filtered \u2192 199,414 edges', style='List Bullet')
doc.add_paragraph('ClinicalTrials.gov: Queries per CVD disease term \u2192 45,358 edges', style='List Bullet')

doc.add_heading('Available Disease Filters', level=2)
add_table(
    ['File', 'Terms', 'Disease Area'],
    [
        ['cvd.txt', '184', 'Cardiovascular disease (active default)'],
        ['alzheimers.txt', '35', "Alzheimer's & related dementias"],
        ['cancer.txt', '70', 'Cancer / oncology'],
        ['asthma.txt', '48', 'Asthma & respiratory diseases'],
        ['diabetes.txt', '52', 'Diabetes & metabolic diseases'],
    ]
)
doc.add_paragraph(
    'CVD gene list: ontology/genes/cvd.txt contains 3,984 CVD gene symbols '
    'sourced from OMIM + DisGeNET, cleaned of LOC* loci and OMIM phenotype symbols.'
)

# -- 6. Web Interface --
doc.add_heading('6. Web Interface', level=1)

doc.add_heading('6.1 Backend (Flask, port 5050)', level=2)
add_table(
    ['Endpoint', 'Method', 'Purpose'],
    [
        ['/api/graph-stats', 'GET', 'Live node/relationship counts from Memgraph'],
        ['/api/query', 'POST', 'Execute arbitrary Cypher queries'],
        ['/api/explore', 'GET', 'Disease subgraph exploration with specificity ranking'],
        ['/api/subgraph', 'POST', 'N-hop disease subgraph extraction (JSON/CSV export)'],
        ['/api/agent/build', 'POST', 'DiseaseQueryAgent: enrich graph for a disease'],
        ['/api/agent/build-disease-graph', 'POST', 'SSE-streamed disease graph building'],
        ['/api/specificity-info', 'GET', 'Specificity score metadata'],
    ]
)

doc.add_heading('6.2 Frontend (interface/index.html)', level=2)
doc.add_paragraph('Explore tab: vis.js force-directed graph with DataSet-based rendering, node type filtering, specificity-ranked results, click-to-inspect detail panels, CSV/JSON export', style='List Bullet')
doc.add_paragraph('Query tab: Neo4j Browser-style multi-panel results; each query appends a new panel with table/graph tabs', style='List Bullet')
doc.add_paragraph('Build Knowledge Graph (sidebar): Claude API standardizes disease name \u2192 ClinicalTrials.gov API v2 fetch \u2192 Memgraph load \u2192 auto-explore', style='List Bullet')
doc.add_paragraph('Extract Disease Subgraph (sidebar): 1-3 hop extraction with JSON/CSV export', style='List Bullet')
doc.add_paragraph('Dashboard: Live stats from /api/graph-stats', style='List Bullet')
doc.add_paragraph('Admin: Parser status, health checks, full pipeline trigger', style='List Bullet')

doc.add_heading('6.3 Specificity Scoring', level=2)
doc.add_paragraph(
    'Pre-computed as n.specificityScore on every node. Formula: 1.0 / count(Disease neighbors). '
    'Disease nodes get 0.0; nodes with no Disease connections get 1.0. '
    'Script: scripts/compute_specificity.py (auto-runs at end of pipeline). '
    'Metadata stored in _Metadata node with timestamp.'
)

# -- 7. AI Agents --
doc.add_heading('7. AI Agents', level=1)

doc.add_heading('7.1 DatabaseAgent (src/database_agent.py)', level=2)
doc.add_paragraph('Autonomously generates new parsers from a database name + download URL:')
doc.add_paragraph('Downloads first 64KB to detect format (TSV, CSV, JSON, XML)', style='List Number')
doc.add_paragraph('Sends sample + BaseParser source to Claude API \u2192 generates parser + ontology configs', style='List Number')
doc.add_paragraph('Saves parser file, registers configs, integrates into pipeline', style='List Number')
doc.add_paragraph('Executes parser, validates output, loads into Memgraph', style='List Number')
doc.add_paragraph('4 parsers in production: HGNC Families, ClinVar, DrugAge, AnAge')

doc.add_heading('7.2 DiseaseQueryAgent (src/disease_agent.py)', level=2)
doc.add_paragraph('On-demand disease enrichment via web interface:')
doc.add_paragraph('User enters disease name in "Build Knowledge Graph" sidebar', style='List Number')
doc.add_paragraph('Claude API standardizes the disease name', style='List Number')
doc.add_paragraph('Queries ClinicalTrials.gov API v2 for matching trials', style='List Number')
doc.add_paragraph('Loads results into Memgraph (ClinicalTrial nodes + edges)', style='List Number')
doc.add_paragraph('Caches results in DiseaseCache node (same disease returns instantly)', style='List Number')
doc.add_paragraph('SSE-streamed progress to frontend', style='List Number')

# -- 8. Legacy Sources --
doc.add_heading('8. Legacy Sources', level=1)
doc.add_paragraph('Three sources use archived/pinned data with no live API replacement:')
add_table(
    ['Source', 'Data Vintage', 'Edges', 'Why Retained'],
    [
        ['SIDER', '2015 GitHub commit', '148,518', 'Only source for drug \u2192 side effect relationships'],
        ['LINCS L1000', '2020 GitHub commit', '171,036', 'Gene regulation + drug expression; clue.io requires institutional access'],
        ['MEDLINE', 'Pinned GitHub commit', '365', 'Unique disease \u2192 anatomy/symptom cooccurrence not in PubTator'],
    ]
)

# -- 9. Deduplication Principles --
doc.add_heading('9. Deduplication Principles', level=1)
doc.add_paragraph(
    'One authoritative source per edge type \u2014 no two databases contribute the same '
    'relationship type, with the exception of geneAssociatesWithDisease (OpenTargets curated '
    '+ PubTator literature-mined = complementary evidence).', style='List Number'
)
doc.add_paragraph(
    '10 sources removed during systematic dedup audit: DisGeNET, GWAS Catalog, Jensen DISEASES, '
    'OMIM, WikiPathways, AOP-DB, HGNC base, CellAge, GenAge, Hetionet precomputed.', style='List Number'
)
doc.add_paragraph(
    'Full rationale documented in docs/CardioKB_Redundancy_Changelog.docx.', style='List Number'
)

# Save
output_path = '/Users/nawaza/Desktop/Cardio-KB/docs/CardioKB_System_Design.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
