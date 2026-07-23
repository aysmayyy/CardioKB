"""Generate CardioKB paper draft as Word document."""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def para(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    return p


def labeled_para(label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(11)
    return p


def italic_para(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    return p


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Shading'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
    return t


# === TITLE ===
title = doc.add_heading(
    'CardioKB: A Cardiovascular Disease Knowledge Graph for '
    'Multi-Source Integration and Drug Repurposing via Link Prediction', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Asma Nawaz')
r.bold = True
r.font.name = 'Times New Roman'
r = p.add_run('1,*, Jay Moran1, Binglan Li2, Jason H. Moore1')
r.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    '1Department of Computational Biomedicine, Cedars-Sinai Medical Center, '
    'Los Angeles, CA, USA\n'
    '2Department of Biomedical Data Science, Stanford University, Stanford, CA, USA\n'
    '*To whom correspondence should be addressed: asma.nawaz@cshs.org')
r.font.size = Pt(10)
r.font.name = 'Times New Roman'

doc.add_page_break()

# === ABSTRACT ===
heading('Abstract', 1)

labeled_para('Motivation: ',
    'Cardiovascular disease (CVD) research requires integrating genetic, '
    'molecular, pharmacological, and clinical evidence scattered across '
    'dozens of specialized databases. Existing biomedical knowledge graphs '
    'cover all human diseases at insufficient CVD depth or rely on static '
    'data snapshots. No publicly available knowledge graph provides '
    'CVD-focused, multi-source integration with an interactive query '
    'interface and embedded drug repurposing predictions.')

labeled_para('Results: ',
    'We present CardioKB, a CVD-focused knowledge graph integrating 23 '
    'curated data sources into 453,037 nodes and 5,461,783 relationships '
    'across 17 node types and 28 relationship types. CardioKB was '
    'constructed using BaseAgent, a multi-agent AI framework, and includes '
    'a systematic entity-resolution audit that merged 6,055 duplicate drug '
    'nodes. We implement a drug repurposing pipeline using CompGCN and '
    'RotatE knowledge graph embeddings decoded by XGBoost, achieving AUROC '
    'of 0.9865, MRR of 0.2284, and Hits@10 of 38.3% on a held-out test '
    'set. The 14,435 predicted drug-disease edges are stored in the graph '
    'and exposed through an interactive web interface with provenance metadata.')

labeled_para('Availability and Implementation: ',
    'CardioKB is freely accessible at [URL] with no login required. '
    'Source code: https://github.com/aysmayyy/CardioKB. '
    'Deployed via Docker Compose (Memgraph + Flask), ~16 GB RAM.')

labeled_para('Contact: ', 'asma.nawaz@cshs.org')

# === 1. INTRODUCTION ===
heading('1 Introduction', 1)

para('Cardiovascular disease (CVD) is one of the leading causes of '
     'mortality worldwide and yet relevant genetic, molecular, and '
     'pharmacological information is fragmented across multiple databases '
     'that operate separately. The evidence needed for a researcher to ask '
     'even a single question—such as whether an existing drug could be '
     'repurposed to treat a given cardiovascular condition—requires '
     'exploration of gene-disease association resources (OpenTargets, '
     'PubTator), drug and pharmacology databases (DrugBank, DrugCentral), '
     'clinical trial registries (ClinicalTrials.gov), pathway and ontology '
     'resources (Reactome, Gene Ontology, Disease Ontology), and variant '
     'databases (ClinVar). These resources use different identifiers, '
     'scopes, and update cadences. A researcher who wants to reason across '
     'these layers must manually explore sources, reconcile identifiers, '
     'and compile information—a process that is slow, error-prone, and '
     'poorly reproducible.')

para('Knowledge graphs have emerged as resources to integrate multi-source '
     'biomedical evidence. Existing knowledge graphs include Hetionet [1], '
     'DRKG [2], PrimeKG [3], and PharmKG [4]. While these resources have '
     'enabled link prediction and drug repurposing, they cover all human '
     'diseases rather than focusing on a specific field such as CVD. '
     'Hetionet, for example, was built from source snapshots assembled '
     'around 2017 and has not been updated since. Given the pace of '
     'biomedical research, CVD-relevant information such as recent clinical '
     'trial data, updated drug indications, and current variant-disease '
     'associations is not fully represented at the depth that CVD research '
     'requires.')

para('AlzKB [5], an Alzheimer disease (AD) focused knowledge graph, '
     'demonstrated the value of narrowing disease scope: it integrates 22 '
     'data sources and applies graph data science and machine learning to '
     'propose therapeutic targets and repurpose drugs. AlzKB was constructed '
     'manually. CardioKB adapts this approach for CVD using BaseAgent [6], '
     'a multi-agent AI system that autonomously constructs disease-specific '
     'knowledge graphs given a disease scope and template codebase. '
     'BaseAgent coordinates specialist agents for ontology configuration, '
     'parser engineering, schema mapping, and evaluation, and was configured '
     'to emulate AlzKB\'s infrastructure while substituting a CVD-specific '
     'source list, schema, and evaluation protocol.')

para('An issue that arises when aggregating multiple data sources is entity '
     'resolution. Each source uses unique identifiers for drugs, genes, and '
     'diseases. Naïve integration produces duplicate nodes that are not '
     'only cosmetically problematic but dangerous for downstream machine '
     'learning, as they can leak test-set information into training data '
     'and inflate reported performance in ways that are difficult to detect '
     'without thorough auditing.')

para('In this paper we present CardioKB, a CVD-focused knowledge graph '
     'integrating 23 curated data sources into 453,037 nodes and 5,461,783 '
     'relationships spanning 17 node types and 28 relationship types. '
     'CardioKB is deployed via Docker Compose with an interactive web '
     'interface including graph visualization, disease subgraph extraction, '
     'and natural-language-to-Cypher querying via Eng2Cypher [9]. We '
     'implement a drug repurposing pipeline using CompGCN [7] and RotatE '
     '[8] knowledge graph embeddings and document a systematic '
     'entity-resolution audit. Our contributions are: (1) a CVD-focused '
     'knowledge graph deduplicated across 23 sources with an auditable '
     'entity-resolution methodology; (2) a validated drug repurposing '
     'pipeline benchmarked with ranking metrics directly comparable to '
     'AlzKB; and (3) a deployed, freely accessible instance with an '
     'interactive web interface for immediate research use.')

# === 2. SYSTEM AND METHODS ===
heading('2 System and Methods', 1)

heading('2.1 Data source selection', 2)
para('CardioKB integrates 23 data sources selected for: (1) authoritative '
     'coverage of a specific entity type, (2) non-redundancy, and (3) '
     'public accessibility. An initial set of 37 candidate sources was '
     'reduced to 23 through a deduplication audit that removed 14 redundant '
     'sources (Supplementary Table S1). Five direct parsers query live APIs '
     'or structured files: ClinicalTrials.gov (API v2), ClinPGx, NCBI Gene, '
     'DoRothEA/OmniPath, and DrugBank. Fifteen component parsers target '
     'original upstream databases for relationship types previously accessed '
     'through Hetionet [1]: Disease Ontology, Gene Ontology, Uberon, MeSH, '
     'SIDER, LINCS L1000, DrugCentral, BindingDB, PubTator, CTD, Bgee, '
     'HPO, Reactome, STRING, and OpenTargets. Three additional sources '
     'complete the graph: HGNC Gene Families, ClinVar, and text-mined '
     'DrugBank indication edges. Two legacy sources (SIDER, LINCS L1000) '
     'are retained from static files as no live API exists.')

heading('2.2 Graph schema and construction', 2)
para('The schema defines 17 node types and 28 relationship types (Figure '
     '1). Gene nodes (193,795; 42.8%) serve as the central hub. Every '
     'relationship carries a source property identifying the originating '
     'database. Seven relationship types carry quantitative properties '
     '(STRING combined scores, Bgee expression scores, DoRothEA '
     'mode-of-regulation, Reactome evidence codes, OpenTargets association '
     'scores, DrugBank interaction types, ClinVar clinical significance). '
     'CardioKB was constructed using BaseAgent [6], configured with a '
     'CVD-specific ontology (184 disease terms, 3,984 gene symbols). Each '
     'parser extends a common BaseParser class. The full pipeline rebuilds '
     'in approximately 5 minutes.')

heading('2.3 Drug-disease treatment edges', 2)
para('The drugTreatsDisease relationship (4,852 edges) aggregates four '
     'sources: CTD curated therapeutic associations (3,099), text-mined '
     'DrugBank indication fields (1,449), DrugCentral FDA-approved '
     'indications with CUI-to-DOID mapping (157), and ClinicalTrials.gov '
     'Phase 3/4 trials filtered by four criteria—primaryPurpose=TREATMENT, '
     'EXPERIMENTAL arm, first-listed condition, trialCount deduplication—'
     'reducing inferred edges from 868 to 147 (83.1% reduction).')

heading('2.4 Entity resolution', 2)
para('The multi-source pipeline created duplicate Drug nodes when sources '
     'loaded the same compound under different internal identifiers sharing '
     'the same canonical DrugBank cross-reference. Post-hoc entity '
     'resolution identified 5,611 duplicate groups (2×–8× duplication), '
     'removed 6,055 nodes, transferred 474,641 edges, and deduplicated '
     '9,094 redundant edges, reducing Drug nodes from 32,849 to 26,794. '
     'This eliminated ~1.5% spurious ML predictions (known treatments '
     'leaking via duplicates) and ~8.7% wasted prediction slots. A '
     'subsequent stale-identifier bug—where Memgraph recycled deleted '
     'node IDs, causing predictions to land on wrong node types—was '
     'detected when the UI displayed lab reagents as predicted treatments '
     'and fixed by matching nodes by name from the live graph. All models '
     'were retrained on the corrected graph.')

# === 3. ALGORITHM ===
heading('3 Algorithm', 1)

heading('3.1 Problem formulation and data preparation', 2)
para('Drug repurposing is formulated as link prediction: predict which '
     'Drug-Disease pairs are likely to have treatment relationships. The '
     'candidate space is 10,310 therapeutic drugs (filtered by requiring '
     'at least one edge in {drugBindsGene, compoundInPharmacologicClass, '
     'compoundCausesSideEffect, drugTreatsDisease, AFFECTS_RESPONSE_TO, '
     'TESTS_INTERVENTION}) crossed with all Disease nodes. The 4,469 '
     'drugTreatsDisease edges in the ML export were split 80/10/10 into '
     'train, validation, and test sets with stratification by disease (seed=42). '
     'Negative samples were generated at 1:1 ratio, excluding all positive '
     'edges from the negative pool.')

heading('3.2 Embedding methods', 2)
para('CompGCN [7] performs relation-aware message passing with subtraction '
     'composition, jointly learning node and relation embeddings across all '
     '28 relationship types. Architecture: 2 GNN layers, 128-dimensional '
     'embeddings, 32M parameters, DistMult scoring, BCE loss, Adam '
     'optimizer (lr=1e-3), dropout 0.3, gradient clipping (max_norm=1.0), '
     'early stopping (patience=20, best at epoch 60). Training: ~7 min on '
     'HPC GPU.')

para('RotatE [8] models relations as rotations in complex space, learning '
     'independent per-entity embeddings. Trained via PyKEEN [11]: 128 '
     'complex dimensions (256 real), 200 epochs, batch size 4,096, '
     'lr=1e-4, 64 negative samples/positive, NSSALoss (margin=9.0, '
     'adversarial_temperature=1.0). Training: ~3.4 hrs on NVIDIA L40S GPU.')

heading('3.3 Feature engineering and decoding', 2)
para('For each drug-disease pair, we compute: Hadamard product, absolute '
     'difference, cosine similarity, and L2 distance of embeddings, '
     'concatenated with six structural features (shared neighbors, Jaccard '
     'coefficient, Adamic-Adar index, preferential attachment, log-degree '
     'of each node). This yields 264-dim (CompGCN) and 520-dim (RotatE) '
     'feature vectors. Three decoders were compared: cosine similarity '
     '(baseline), XGBoost [10] (n_estimators=300, max_depth=6, lr=0.1, '
     'early_stopping=20), and MLP (single hidden layer).')

heading('3.4 Evaluation', 2)
para('Models were evaluated on the held-out test set using AUROC, AUPRC, '
     'Hits@K (K=1, 3, 10), and MRR—the same ranking metrics reported by '
     'AlzKB [5] to enable direct comparison.')

heading('3.5 Results', 2)
italic_para('Table 1. Link prediction performance (XGBoost ranking: CompGCN n=337, RotatE n=316).')

add_table(
    ['Method', 'Decoder', 'AUROC', 'AUPRC', 'Hits@1', 'Hits@3',
     'Hits@10', 'MRR', 'Med.Rank'],
    [
        ['CompGCN', 'Cosine',  '0.310', '0.381', '—', '—',
         '0.2%',  '0.003', '2,230'],
        ['CompGCN', 'XGBoost', '0.987', '0.985', '14.8%', '23.2%',
         '38.3%', '0.228', '22'],
        ['CompGCN', 'MLP',     '0.984', '0.978', '—', '—',
         '29.3%', '0.117', '27.5'],
        ['RotatE',  'Cosine',  '0.781', '0.757', '—', '—',
         '1.7%',  '0.010', '461'],
        ['RotatE',  'XGBoost', '0.983', '0.981', '9.5%', '22.5%',
         '43.7%', '0.205', '15'],
        ['RotatE',  'MLP',     '0.981', '0.979', '—', '—',
         '41.1%', '0.190', '16'],
    ])

para('')
para('XGBoost consistently outperforms cosine and MLP decoders. CompGCN + '
     'XGBoost achieves the highest AUROC (0.9865) and MRR (0.2284). The '
     'cosine baseline performs poorly for CompGCN (AUROC 0.31), indicating '
     'that CompGCN embeddings encode relational structure requiring a '
     'nonlinear decoder, whereas RotatE\'s cosine baseline (0.78) reflects '
     'its optimization for direct geometric comparison.')

italic_para('Table 2. Comparison with AlzKB [5] (best model, ranking metrics).')

add_table(
    ['System', 'Method', 'Decoder', 'MRR', 'Hits@1', 'Hits@3', 'Hits@10'],
    [
        ['AlzKB',   'RotatE',  'Cosine (PyKEEN)', '0.202', '0.126',
         '0.220', '0.358'],
        ['CardioKB', 'CompGCN', 'XGBoost', '0.228', '0.148',
         '0.232', '0.383'],
        ['CardioKB', 'RotatE',  'XGBoost', '0.205', '0.095',
         '0.225', '0.437'],
    ])

para('')
para('CardioKB\'s CompGCN + XGBoost achieves MRR of 0.228 and Hits@10 of '
     '38.3%, outperforming AlzKB\'s best RotatE result (MRR 0.202, Hits@10 '
     '0.358). A key difference is the decoder: AlzKB uses native PyKEEN '
     'cosine scoring, while CardioKB demonstrates that an XGBoost decoder '
     'with structural features substantially improves over cosine baselines '
     '(CompGCN cosine MRR: 0.003 → XGBoost: 0.228).')

heading('3.6 Predictions stored in graph', 2)
para('Predictions exceeding confidence 0.5 were stored as '
     'predictedTreatsDisease edges. CompGCN: 6,607 edges (1,038 drugs '
     '× 37 diseases, confidence 0.989–0.991). RotatE: 7,828 edges '
     '(1,165 drugs × 142 diseases, confidence 0.993–0.997). Total: '
     '14,435 predicted edges. Confidence values reflect embedding geometry '
     'match, not calibrated treatment probabilities.')

# === 4. IMPLEMENTATION ===
heading('4 Implementation', 1)

para('CardioKB is deployed as a Docker Compose stack: Memgraph (in-memory '
     'graph database, Bolt protocol) and Flask (Python 3.11). The graph '
     'export is 304 MB compressed (~14 GB in memory), loaded via binary '
     'volume restore in seconds. The system requires ~16 GB RAM.')

para('The web interface provides three modes without authentication. '
     '(1) Graph exploration: search via autocomplete, interactive '
     'force-directed visualization (vis.js), specificity-ranked nodes, '
     'node/edge type filters, ML predictions as orange dashed edges with '
     'toggle, export as CSV/JSON/PNG/PDF. (2) Natural language querying: '
     'Eng2Cypher [9] translates questions to Cypher with schema-aware '
     'context; results in multi-panel layout with table/graph tabs; direct '
     'Cypher editor with read-only safety. (3) Disease subgraph extraction: '
     'N-hop (1–3) subgraph export as JSON/CSV. Edge provenance panels show '
     'source database, evidence properties, and clinical validation '
     'warnings for ML predictions.')

# === 5. DISCUSSION ===
heading('5 Discussion', 1)

italic_para('Table 3. Comparison with existing biomedical knowledge graphs.')

add_table(
    ['', 'Hetionet [1]', 'DRKG [2]', 'PrimeKG [3]', 'AlzKB [5]',
     'CardioKB'],
    [
        ['Focus',      'All diseases', 'All (COVID)', 'All diseases',
         "Alzheimer's", 'CVD'],
        ['Nodes',      '47,031', '97,238', '129,375', '~118,000',
         '453,037'],
        ['Edges',      '2,250,197', '5,874,261', '8,100,498', '~263,000',
         '5,461,783'],
        ['Node types', '11', '13', '10', '13', '17'],
        ['Edge types', '24', '107', '30', '27', '28'],
        ['Sources',    '29', '6', '20', '22', '23'],
        ['Interactive UI', 'No', 'No', 'No', 'Yes', 'Yes'],
        ['NL querying', 'No', 'No', 'No', 'No', 'Yes'],
        ['Auto build', 'No', 'No', 'No', 'No', 'Yes'],
        ['Updated',    '~2017', '2020', '2023', '2024', '2026'],
    ])

para('')
para('CardioKB combines disease-specific depth with automated construction '
     'and embedded ML predictions. AlzKB [5] is the most direct comparator. '
     'CardioKB adapts this approach for CVD using automated multi-agent '
     'construction, adds natural language querying, and documents an '
     'entity-resolution methodology absent from prior work.')

heading('5.1 Decoder choice and structural features', 2)
para('The choice of decoder has a larger impact than embedding method: '
     'CompGCN\'s AUROC jumps from 0.31 (cosine) to 0.99 (XGBoost), a 0.68 '
     'absolute improvement. This demonstrates that KG embeddings encode '
     'treatment-relevant structure that requires a nonlinear decoder to '
     'extract. Structural graph features (shared neighbors, Jaccard, '
     'Adamic-Adar, preferential attachment) contribute beyond embeddings '
     'alone.')

heading('5.2 Prediction characteristics', 2)
para('CompGCN predictions concentrate on high-degree disease hubs (37 '
     'diseases vs. RotatE\'s 142), reflecting neighborhood aggregation '
     'bias. The narrow confidence bands (0.989–0.997) mean scores '
     'distinguish positive from negative but do not meaningfully rank '
     'within positives. The 1:1 random negative sampling inflates AUROC '
     'relative to harder negatives; ranking metrics (MRR, Hits@K) are '
     'more informative. Among top-30 predictions, several have independent '
     'clinical trial support not in training data: levosimendan→CAD, '
     'atorvastatin→cardiomyopathy, colchicine→hypertension, '
     'clopidogrel→cerebrovascular disease (Supplementary Table S2).')

heading('5.3 Entity resolution as methodology', 2)
para('The 18.4% Drug node duplication rate (6,055/32,849) demonstrates '
     'that entity resolution is essential for multi-source KGs, '
     'particularly before ML. Without it, ~1.5% of predictions represented '
     'leaked known treatments and ~8.7% were redundant. We recommend '
     'entity resolution as a mandatory post-integration step.')

heading('5.4 Limitations', 2)
para('Two sources (SIDER, LINCS L1000) are static legacy datasets. The '
     'CVD disease filter (184 terms) bounds disease-gene association scope. '
     'Full NCBI Gene and Disease Ontology catalogs create expected orphan '
     'rates (~78% Gene, ~76% Disease) that do not affect queries or ML. '
     'NL querying depends on the Anthropic API; graph exploration and '
     'Cypher querying function independently. Confidence scores are not '
     'calibrated probabilities. Systematic clinician evaluation of '
     'predictions is planned as future work.')

# === ACKNOWLEDGEMENTS ===
heading('Acknowledgements', 1)

heading('Funding', 2)
para('This work was supported by [FUNDING AGENCY] [grant number]. '
     '[TODO: fill in]')

heading('AI-use disclosure', 2)
para('The web interface code was developed with assistance from Claude '
     'Code (Anthropic). All scientific analysis, experimental design, '
     'result interpretation, and manuscript preparation were performed '
     'by the authors, in accordance with ISCB/COPE guidelines.')

heading('Conflict of interest', 2)
para('None declared.')

heading('Data availability', 2)
para('Source code: https://github.com/aysmayyy/CardioKB. All 23 data '
     'sources are publicly accessible. Graph export available at '
     '[Zenodo DOI].')

# === REFERENCES ===
heading('References', 1)

refs = [
    '[1] Himmelstein DS et al. Systematic integration of biomedical '
    'knowledge prioritizes drugs for repurposing. eLife 2017;6:e26726.',
    '[2] Ioannidis V et al. DRKG — Drug Repurposing Knowledge Graph '
    'for Covid-19. 2020.',
    '[3] Chandak P, Huang K, Zitnik M. Building a knowledge graph to '
    'enable precision medicine. Sci Data 2023;10(1):67.',
    '[4] Zheng S et al. PharmKG: a dedicated knowledge graph benchmark '
    'for biomedical data mining. Brief Bioinform 2021;22(4):bbaa344.',
    '[5] Romano JD et al. The Alzheimer\'s Knowledge Base: A Knowledge '
    'Graph for Alzheimer Disease Research. J Med Internet Res '
    '2024;26(1):e46777.',
    '[6] Li B. BaseAgent: A multi-agent AI system for autonomous '
    'knowledge graph construction. 2026. '
    'github.com/EpistasisLab/BaseAgent.',
    '[7] Vashishth S et al. Composition-Based Multi-Relational Graph '
    'Convolutional Networks. ICLR 2020.',
    '[8] Sun Z et al. RotatE: Knowledge Graph Embedding by Relational '
    'Rotation in Complex Space. ICLR 2019.',
    '[9] Moran J. Eng2Cypher. 2025. '
    'github.com/CenterAIResearch/Eng2Cypher.',
    '[10] Chen T, Guestrin C. XGBoost: A Scalable Tree Boosting System. '
    'KDD 2016:785–794.',
    '[11] Ali M et al. PyKEEN 1.0: A Python Library for Training and '
    'Evaluating Knowledge Graph Embeddings. JMLR 2021;22(82):1–6.',
]
for ref in refs:
    p = doc.add_paragraph()
    r = p.add_run(ref)
    r.font.size = Pt(9)
    r.font.name = 'Times New Roman'

# === TODOs ===
doc.add_page_break()
heading('TODOs Before Submission', 1)
todos = [
    '[URL]: Insert public CardioKB URL (needs HTTPS)',
    '[HPC]: Run ml/compute_hits_1_3.py on HPC to fill Hits@1 and Hits@3',
    '[LICENSE]: Specify open-source license',
    'Funding statement: fill in agency and grant number',
    'Figure 1: Schema diagram (node types + relationship types)',
    'Figure 2: Pipeline architecture diagram',
    'Figure 3: ROC/PR curves for CompGCN+XGBoost and RotatE+XGBoost',
    'Figure 4: Web interface screenshot(s)',
    'Supplementary Table S1: 14 removed sources with justification',
    'Supplementary Table S2: Top 30 predictions per method '
    '(LaTeX tables in ml/results/explainability/)',
    'Verify AlzKB node/edge counts against Romano et al. 2024',
    'Confirm Eng2Cypher citation URL with Jay Moran',
    'Update BaseAgent citation when Binglan\'s paper publishes',
    'Word count check: target ~5,000 words (7 page limit)',
]
for t in todos:
    doc.add_paragraph(t, style='List Bullet')

# === SAVE ===
out_path = os.path.expanduser('~/Downloads/CardioKB_Paper_Draft.docx')
doc.save(out_path)
print(f'Saved to {out_path}')
