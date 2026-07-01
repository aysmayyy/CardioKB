"""Generate CardioKB Study Guide as a formatted Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1E, 0x40, 0x7C)

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(13)

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def bold_run(para, text):
    r = para.add_run(text)
    r.bold = True
    return r

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

# ══════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('CardioKB Study Guide')
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0x1E, 0x40, 0x7C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Comprehensive Reference for Advisor Discussion and Paper Defense')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = author.add_run('Asma Nawaz')
r.font.size = Pt(13)
r.bold = True

affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = affil.add_run('Ph.D. Project: 2026\nMoore Lab, Cedars-Sinai Medical Center')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. What is CardioKB',
    '2. Data Sources (All 22)',
    '3. Graph Structure',
    '4. BaseAgent Pipeline',
    '5. Key Technical Challenges Solved',
    '6. ML Pipeline: Link Prediction for Drug Repurposing',
    '7. Web Interface Walkthrough',
    '   7.1 Application Layout',
    '   7.2 Dashboard Statistics Bar',
    '   7.3 Explore Tab',
    '   7.4 Query Tab',
    '   7.5 Admin Panel',
    '   7.6 Sidebar Tools',
    '   7.7 Visual Design and Theming',
    '8. Deployment',
    '9. Evaluation',
    '10. Future Directions',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if item.startswith('   '):
        p.paragraph_format.left_indent = Cm(1.27)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. WHAT IS CARDIOKB
# ══════════════════════════════════════════════════════════
doc.add_heading('1. What is CardioKB', level=1)

doc.add_heading('The Problem', level=2)
doc.add_paragraph(
    'Cardiovascular disease (CVD) is the leading cause of death globally, responsible for approximately '
    '17.9 million deaths per year. Researchers studying CVD need to understand complex relationships between '
    'genes, drugs, diseases, clinical trials, protein interactions, gene expression, and phenotypes. This '
    'information is scattered across dozens of separate databases, each with different formats, identifiers, '
    'and access methods. A researcher investigating whether a diabetes drug could be repurposed for heart '
    'failure would need to manually query DrugBank for drug targets, STRING for protein interactions, '
    'ClinicalTrials.gov for ongoing trials, CTD for chemical-gene expression effects, and many more. This '
    'process is slow, error-prone, and makes it nearly impossible to discover non-obvious multi-hop '
    'connections across data sources.'
)

doc.add_heading('What CardioKB Does', level=2)
doc.add_paragraph(
    'CardioKB is a CVD-focused biomedical knowledge graph that integrates 22 deduplicated data sources '
    'into a single queryable graph database containing 459,092 nodes and 5,443,134 relationships across '
    '17 node types and 27 relationship types. It provides:'
)
add_bullet(doc, 'A unified graph where a single Cypher query can traverse from a drug to its gene targets to associated diseases to clinical trials testing those drugs')
add_bullet(doc, 'An automated ETL pipeline that can rebuild the entire graph from source data in approximately 5 minutes using BaseAgent multi-agent orchestration')
add_bullet(doc, 'A web interface for interactive exploration, natural language querying, and disease subgraph extraction')
add_bullet(doc, 'An ML-based drug repurposing pipeline that uses knowledge graph embeddings (RotatE, CompGCN) to predict new drug-disease treatment relationships with AUROC up to 0.9717')

doc.add_heading('Why Cardiovascular Disease', level=2)
add_bullet(doc, 'Highest global mortality of any disease category (17.9 million deaths/year)')
add_bullet(doc, 'Rich existing data ecosystem: extensive clinical trial data, well-characterized gene-disease associations, established drug targets')
add_bullet(doc, 'Strong drug repurposing potential: many CVD drugs have complex polypharmacology, and the graph structure can reveal non-obvious therapeutic connections')
doc.add_paragraph(
    'The knowledge graph architecture is disease-agnostic. Alternative disease filter files exist for '
    "Alzheimer's (35 terms), cancer (70 terms), asthma (48 terms), and diabetes (52 terms). Switching "
    'the disease focus requires only changing a symlink.'
)

doc.add_heading('Target Users', level=2)
add_bullet(doc, 'Biomedical researchers investigating CVD mechanisms and multi-target drug interactions')
add_bullet(doc, 'Computational biologists performing network pharmacology and drug repurposing studies')
add_bullet(doc, 'Clinician-scientists exploring evidence-based treatment relationships and clinical trial landscape')
add_bullet(doc, 'Bioinformaticians needing a pre-integrated CVD knowledge base for downstream analysis')

doc.add_heading('How It Differs from Existing Resources', level=2)
doc.add_paragraph(
    'Several biomedical knowledge graphs exist (Hetionet, DRKG, PrimeKG, PharmKG), but CardioKB differs:'
)
add_table(doc,
    ['Differentiator', 'Description'],
    [
        ['CVD-focused', 'The 184-term CVD disease filter ensures all disease-associated edges are relevant to cardiovascular research, reducing noise from the full human disease space'],
        ['Automated rebuild', 'BaseAgent multi-agent orchestration allows full graph reconstruction from source data in ~5 minutes, vs. manual builds that take weeks or months'],
        ['Integrated ML predictions', '1,500 predicted drug-disease treatment edges (500 per method) stored directly in the graph with confidence scores, queryable alongside curated data'],
        ['Provenance tracking', 'Every relationship carries a source property identifying the originating database (e.g., source: "DrugBank")'],
        ['22 deduplicated sources', 'Systematic redundancy audit removed 12 overlapping sources, ensuring each relationship type comes from exactly one authoritative database'],
    ],
    col_widths=[4, 14]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 2. DATA SOURCES
# ══════════════════════════════════════════════════════════
doc.add_heading('2. Data Sources (All 22)', level=1)
doc.add_paragraph(
    'CardioKB integrates 22 data sources, each contributing unique node types and/or relationship types. '
    'Sources were selected based on three criteria: (1) authoritative coverage of a specific biological '
    'entity type, (2) no redundancy with other included sources, and (3) public accessibility. During a '
    'systematic deduplication audit, 12 sources were removed (DisGeNET, GWAS Catalog, Jensen DISEASES, '
    'OMIM, WikiPathways, AOP-DB, HGNC base, CellAge, GenAge, Hetionet precomputed, DrugAge, AnAge).'
)

doc.add_heading('Direct Parsers (5 sources)', level=2)
direct_sources = [
    ['1', 'ClinicalTrials.gov', 'Public API v2', '85,677 trials; 27,866 STUDIES_CONDITION + 17,492 TESTS_INTERVENTION edges', 'Only source for clinical trial data linking interventions to disease conditions'],
    ['2', 'ClinPGx', 'Public API', '1,091 VARIANT_IN + 503 drug label + 224 AFFECTS_RESPONSE_TO edges', 'Only source for pharmacogenomic drug label annotations'],
    ['3', 'NCBI Gene', 'Public FTP', '193,687 Gene nodes (node-only)', 'Authoritative source for human gene identifiers; provides Gene node backbone'],
    ['4', 'DoRothEA', 'Public API', '12,985 TF-gene edges with morScore + confidence', 'Only source for curated TF-gene regulatory interactions'],
    ['5', 'DrugBank', 'XML file', '19,842 Drug nodes; 12,089 drugBindsGene edges', 'Gold-standard drug database for drug-target binding'],
]
add_table(doc, ['#', 'Source', 'Access', 'Contribution', 'Why Included'], direct_sources, col_widths=[0.8, 3, 2.5, 5, 5])

doc.add_heading('Hetionet-Derived Component Parsers (17 sources)', level=2)
hetio_sources = [
    ['6', 'Disease Ontology', '12,012 diseases; 6,447 diseaseIsSubtypeOf edges', 'Authoritative disease taxonomy with hierarchical subtypes'],
    ['7', 'Gene Ontology', '50,350 BP + 26,935 MF + 25,794 CC edges', 'Only source for functional gene annotations (3 ontology domains)'],
    ['8', 'Uberon', '14,937 BodyPart nodes', 'Anatomy identifiers for Bgee and Jensen TISSUES'],
    ['9', 'MeSH', '966 Symptom nodes', 'Standardized symptom vocabulary'],
    ['10', 'SIDER', '148,518 side effect edges', 'Only source for drug side effects (legacy, no live API)'],
    ['11', 'LINCS L1000', '150,535 gene expression edges with zScore', 'Only source for drug-induced gene expression changes (legacy)'],
    ['12', 'MEDLINE', '365 cooccurrence edges', 'Literature cooccurrence (legacy, pinned GitHub commit)'],
    ['13', 'DrugCentral', '16,403 pharmacologic class + 245 treats edges', 'Only source for pharmacologic classes and FDA-approved indications'],
    ['14', 'BindingDB', '12,250 chemicalBindsGene edges', 'Complements DrugBank with experimentally measured binding data'],
    ['15', 'PubTator Central', '744,427 gene-disease + 4,320 disease-disease edges', 'Largest source of literature-mined gene-disease associations'],
    ['16', 'CTD', '116,451 increases + 97,951 decreases expression edges', 'Only source for directional chemical-gene expression effects'],
    ['17', 'Bgee', '784,026 underexpresses + 1,872 overexpresses edges', 'Only source for tissue-specific gene expression patterns'],
    ['18', 'Jensen TISSUES', '215,235 gene-tissue edges', 'Gene-tissue expression associations'],
    ['19', 'HPO', '19,389 phenotypes; 162,994 gene-phenotype edges', 'Only source for gene-phenotype associations'],
    ['20', 'Reactome', '44,979 geneInPathway edges', 'Authoritative curated biological pathway assignments'],
    ['21', 'STRING', '121,170 PPI edges (confidence > 700)', 'Largest source for protein-protein interaction data'],
    ['22', 'OpenTargets', '32,826 gene-disease edges (CVD-filtered)', 'Curated gene-disease associations from multiple evidence types'],
]
add_table(doc, ['#', 'Source', 'Contribution', 'Why Included'], hetio_sources, col_widths=[0.8, 3, 6, 6])

doc.add_heading('Additional Parsers (2 sources)', level=2)
add_table(doc, ['#', 'Source', 'Contribution', 'Why Included'], [
    ['23', 'HGNC Gene Families', '1,934 GeneFamily nodes; 5,123 geneInFamily edges', 'Only source for gene family membership'],
    ['24', 'ClinVar', '4,488,042 Variant nodes; 2,267,095 hasVariant + variantInGene edges', 'Authoritative source for clinical variant interpretations'],
], col_widths=[0.8, 3.5, 6, 5.5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 3. GRAPH STRUCTURE
# ══════════════════════════════════════════════════════════
doc.add_heading('3. Graph Structure', level=1)

doc.add_heading('Scale', level=2)
add_bullet(doc, '459,092 nodes across 17 distinct node types')
add_bullet(doc, '5,443,134 relationships across 27 relationship types')
add_bullet(doc, '22 data sources (19 with edge source labels + 3 ML prediction sources)')
add_bullet(doc, 'Every relationship carries a source property identifying its originating database')

doc.add_heading('Node Types (17)', level=2)
node_types = [
    ['Gene', '193,795', 'NCBI Gene', 'Human genes with symbols, chromosomal locations, cross-references'],
    ['Variant', '135,555', 'ClinVar', 'Genetic variants with clinical significance annotations'],
    ['Drug', '32,849', 'DrugBank + CTD', 'Drugs/chemicals with DrugBank IDs, names, cross-references'],
    ['BiologicalProcess', '24,428', 'Gene Ontology', 'GO biological process terms'],
    ['ClinicalTrial', '21,578', 'ClinicalTrials.gov', 'Clinical studies with phase, status, sponsor info'],
    ['Phenotype', '19,389', 'HPO', 'Observable clinical features from HPO'],
    ['MolecularFunction', '10,056', 'Gene Ontology', 'GO molecular function terms'],
    ['GeneFamily', '4,257', 'HGNC', 'Gene family groupings'],
    ['CellularComponent', '4,076', 'Gene Ontology', 'GO cellular component terms'],
    ['Disease', '3,442', 'Disease Ontology', 'Diseases with DOID identifiers and definitions'],
    ['Pathway', '2,870', 'Reactome', 'Biological pathways'],
    ['PharmacologicClass', '2,359', 'DrugCentral', 'FDA pharmacologic classes'],
    ['SideEffect', '2,227', 'SIDER', 'Drug side effects with UMLS CUI identifiers'],
    ['BodyPart', '1,400', 'Uberon', 'Anatomical structures'],
    ['Symptom', '415', 'MeSH', 'Clinical symptoms'],
    ['TranscriptionFactor', '367', 'DoRothEA', 'Transcription factors with regulatory targets'],
    ['DrugLabel', '29', 'ClinPGx', 'FDA pharmacogenomic drug label annotations'],
]
add_table(doc, ['Node Type', 'Count', 'Source', 'Description'], node_types, col_widths=[3.5, 2, 3.5, 7])

doc.add_heading('Relationship Types (27)', level=2)
rel_types = [
    ['bodyPartOverexpressesGene', 'BodyPart→Gene', 'Bgee', '2,749,193', 'expressionScore'],
    ['geneAssociatesWithDisease', 'Gene→Disease', 'PubTator+OT', '539,964', 'score (OT)'],
    ['chemicalIncreasesExpression', 'Drug→Gene', 'CTD', '343,823', '—'],
    ['chemicalDecreasesExpression', 'Drug→Gene', 'CTD', '328,726', '—'],
    ['geneAssociatesWithPhenotype', 'Gene→Phenotype', 'HPO', '270,265', '—'],
    ['geneInteractsWithGene', 'Gene→Gene', 'STRING', '229,007', 'combinedScore'],
    ['geneInPathway', 'Gene→Pathway', 'Reactome', '137,116', 'evidenceCode'],
    ['variantInGene', 'Variant→Gene', 'ClinVar', '135,393', '—'],
    ['geneParticipatesInBP', 'Gene→BP', 'Gene Ontology', '122,117', '—'],
    ['geneAssociatedWithCC', 'Gene→CC', 'Gene Ontology', '90,141', '—'],
    ['geneHasMolecularFunction', 'Gene→MF', 'Gene Ontology', '76,612', '—'],
    ['compoundUpregulatesGene', 'Drug→Gene', 'LINCS L1000', '74,854', 'zScore'],
    ['compoundCausesSideEffect', 'Drug→SideEffect', 'SIDER', '67,721', '—'],
    ['compoundDownregulatesGene', 'Drug→Gene', 'LINCS L1000', '64,661', 'zScore'],
    ['variantAssocWithDisease', 'Variant→Disease', 'ClinVar', '51,323', 'clinSignificance'],
    ['drugBindsGene', 'Drug→Gene', 'DrugBank', '29,363', 'interactionType'],
    ['geneInFamily', 'Gene→GeneFamily', 'HGNC', '27,022', '—'],
    ['compoundInPharmClass', 'Drug→PharmClass', 'DrugCentral', '25,687', '—'],
    ['chemicalBindsGene', 'Drug→Gene', 'BindingDB', '22,735', '—'],
    ['STUDIES_CONDITION', 'Trial→Disease', 'ClinicalTrials', '20,667', '—'],
    ['tfInteractsWithGene', 'TF→Gene', 'DoRothEA', '15,082', 'morScore, confidence'],
    ['drugTreatsDisease', 'Drug→Disease', 'CTD+CT+DC', '3,782', '—'],
    ['TESTS_INTERVENTION', 'Trial→Drug', 'ClinicalTrials', '3,180', '—'],
    ['diseaseIsSubtypeOf', 'Disease→Disease', 'Disease Ontology', '2,581', '—'],
    ['predictedTreatsDisease', 'Drug→Disease', 'ML Predictions', '1,500', 'confidence'],
    ['AFFECTS_RESPONSE_TO', 'Gene→Drug', 'ClinPGx', '74', '—'],
]
add_table(doc, ['Relationship', 'Direction', 'Source', 'Count', 'Properties'], rel_types, col_widths=[4.5, 2.5, 2.5, 2, 3])

doc.add_heading('Why Memgraph Over Neo4j', level=2)
add_bullet(doc, 'In-memory architecture: entire graph in RAM for sub-second interactive query response')
add_bullet(doc, 'Cypher compatibility: same query language as Neo4j, requiring only minor syntax adjustments')
add_bullet(doc, 'Docker-friendly: lightweight container image with fast startup')
add_bullet(doc, 'No license restrictions: open-source edition has no node/relationship count limits')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 4. BASEAGENT PIPELINE
# ══════════════════════════════════════════════════════════
doc.add_heading('4. BaseAgent Pipeline', level=1)

doc.add_heading('What is BaseAgent', level=2)
doc.add_paragraph(
    'BaseAgent is a multi-agent orchestration framework built by Binglan Li (Li Lab, Stanford) that '
    'automates the construction of biomedical knowledge graphs. It was originally developed for AlzKB '
    "(Alzheimer's Knowledge Base) and provides a template-based system where each data source is "
    'integrated through a standardized pipeline of AI agents.'
)

doc.add_heading('How It Works', level=2)
doc.add_paragraph('BaseAgent uses four specialized agents working in sequence:')
add_bullet(doc, 'Ontology Agent: Analyzes the target data source and maps its entities to the knowledge graph schema')
add_bullet(doc, 'Engineer Agent: Generates a parser that downloads, processes, and converts raw data into standardized TSV format')
add_bullet(doc, 'Mapping Agent: Resolves identifier conflicts between sources (e.g., MeSH IDs to DOID, DrugBank IDs to CAS numbers)')
add_bullet(doc, 'Evaluator Agent: Validates parsed output against expected schemas, checks data quality, reports yield statistics')

doc.add_paragraph('For CardioKB, the pipeline orchestrator (src/main.py) coordinates five phases:')
add_table(doc, ['Phase', 'Description'], [
    ['1. Download', 'Each parser downloads raw data to data/raw/<source>/'],
    ['2. Parse', 'Each parser extracts nodes and edges into pandas DataFrames'],
    ['3. TSV Export', 'DataFrames written to data/processed/<source>/*.tsv'],
    ['4. Graph Load', 'memgraph_loader.py reads TSVs + ontology_configs.py, uses UNWIND-based Cypher batching (batch_size=1000), MERGE to prevent duplicates, sets r.source on every relationship'],
    ['5. Post-processing', 'compute_specificity.py calculates and stores specificityScore on all nodes'],
], col_widths=[3, 14])

doc.add_heading('Key Fixes for CardioKB', level=2)
add_bullet(doc, 'RDF serialization bypass: Implemented direct TSV-to-Memgraph loading with 86 Python ontology configs, bypassing ista RDF serialization that only serializes schema elements')
add_bullet(doc, 'ID format harmonization: Rewrote TSVMemgraphExporter with comprehensive lookup tables resolving gene symbols, NCBI IDs, DrugBank IDs to canonical node IDs')
add_bullet(doc, 'Batch session handling: New database session per batch (batch size 5000) to prevent transaction timeouts')
add_bullet(doc, 'Source property propagation: Added source_label field to all 86 ontology configs for provenance tracking')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 5. KEY TECHNICAL CHALLENGES
# ══════════════════════════════════════════════════════════
doc.add_heading('5. Key Technical Challenges Solved', level=1)

challenges = [
    ('5.1 FlatFileDatabaseParser to owl2.DataLoader (60,000x Speedup)',
     'The original BaseAgent template used ista\'s FlatFileDatabaseParser to load ontology files. For large OWL2 ontologies (Disease Ontology: 42,000 terms), this took over 10 minutes per file. Switching to ista\'s owl2.DataLoader.execute() method processed the same file in under 10 milliseconds — a 60,000x speedup that reduced the pipeline from hours to minutes.'),
    ('5.2 STUDIES_CONDITION Synonym Resolution (1 to 20,667 edges)',
     'The ClinicalTrials.gov parser initially produced only 1 STUDIES_CONDITION edge due to exact string matching between trial condition names ("Heart Attack") and Disease Ontology formal names ("myocardial infarction"). Implementing synonym resolution using Disease Ontology\'s synonym fields increased matches from 1 to 20,667 edges.'),
    ('5.3 CTD Expression Edge Fix (0 to 525,907 edges)',
     'The initial CTD parser produced 0 expression edges despite CTD containing hundreds of thousands. A column mapping bug caused gene IDs to be read from the organism column. Fixing this yielded 525,907 expression edges — the single largest edge count fix in the project.'),
    ('5.4 MemgraphExporter Streaming Fix (Silent Triple Loss)',
     'The BaseAgent TSVMemgraphExporter silently dropped triples when DataFrames contained null values or mixed-type columns. Added explicit null handling and mixed-type coercion, plus logging for every skipped row, recovering thousands of edges.'),
    ('5.5 Disease CUI Mapping Gap',
     'DrugCentral uses UMLS CUI identifiers but CardioKB uses DOID. Implemented CUI-to-DOID mapping using Disease Ontology xref fields, increasing DrugCentral drugTreatsDisease edges from near-zero to 245.'),
    ('5.6 drugTreatsDisease Edge Aggregation',
     'The critical drugTreatsDisease relationship comes from three complementary sources: CTD (2,757 curated), ClinicalTrials.gov (868 extracted), and DrugCentral (157 FDA-approved), totaling 3,782 deduplicated edges that serve as positive training labels for the ML pipeline.'),
]
for title, desc in challenges:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 6. ML PIPELINE
# ══════════════════════════════════════════════════════════
doc.add_heading('6. ML Pipeline: Link Prediction for Drug Repurposing', level=1)

doc.add_heading('6.1 Why Link Prediction', level=2)
doc.add_paragraph(
    'Drug repurposing (finding new therapeutic uses for existing drugs) is faster and cheaper than de novo '
    'drug development because repurposed drugs have already passed safety testing. A knowledge graph '
    'naturally encodes the relationships that make a drug useful for a disease: if Drug A binds Gene X, '
    'Gene X is associated with Disease Y, and other drugs binding similar genes treat Disease Y, then '
    'Drug A is a plausible repurposing candidate. Link prediction formalizes this intuition.'
)

doc.add_heading('6.2 Pipeline Architecture', level=2)
add_table(doc, ['Step', 'Script', 'Description'], [
    ['1. Export', 'export_edges.py', 'Extracts all 5.4M edges as (source, target, relation_type) triples'],
    ['2. Split', 'split_edges.py', '80/10/10 train/val/test split, stratified by edge type (seed=42)'],
    ['3. Train', 'HPC GPU', 'Each method trains embeddings on train split only (no leakage)'],
    ['4. Evaluate', 'evaluate_*.py', 'XGBoost decoder trained on train, early-stopped on validation, evaluated on test'],
    ['5. Predict', 'link_prediction*.py', 'Score all 9,735 x 457 = 4,448,895 Drug-Disease candidates; keep top 500 (conf >= 0.5)'],
    ['6. Store', 'Memgraph', 'Insert top 500 as predictedTreatsDisease edges per method'],
], col_widths=[1.5, 3.5, 11])

doc.add_heading('6.3 Data Split', level=2)
add_table(doc, ['Split', 'Total Edges', 'drugTreatsDisease', 'Purpose'], [
    ['Train', '4,349,555', '3,026', 'Embedding training only'],
    ['Validation', '543,683', '~310', 'Early stopping, hyperparameter tuning'],
    ['Test', '543,683', '~325', 'Final evaluation (never seen during training)'],
], col_widths=[2.5, 3, 3, 7])

doc.add_heading('6.4 RotatE (Knowledge Graph Embedding)', level=2)
doc.add_paragraph(
    'RotatE models each relation as a rotation in complex embedding space. For a triple (head, relation, tail), '
    'the scoring function is ||h * r - t||, where h and t are complex vectors and r is a unit-modulus complex '
    'vector (a rotation). Each of the 25 relation types learns its own rotation.'
)
add_bullet(doc, 'Implementation: PyKEEN framework. 256 real dimensions (128 complex).')
add_bullet(doc, 'Training: 200 epochs, batch_size=4096, lr=1e-4, 64 negative samples, NSSALoss (margin=9.0). L40S GPU, ~10.3 hours.')
add_bullet(doc, 'Native MRR: 0.1119 (ranking against all 459K entities).')
add_bullet(doc, 'Best result: AUROC 0.9652 with XGBoost decoder.')

doc.add_heading('6.5 CompGCN (Graph Neural Network)', level=2)
doc.add_paragraph(
    'CompGCN is a message-passing GNN that jointly embeds nodes and relations through composition operators '
    'during neighborhood aggregation. For each node, it aggregates neighbor embeddings composed with relation '
    'embeddings using subtraction: phi(e_neighbor, e_relation) = e_neighbor - e_relation.'
)
p = doc.add_paragraph()
bold_run(p, 'CompGCN vs R-GCN: ')
p.add_run(
    'R-GCN learns a separate weight matrix per relation type (25 x 128 x 128 = 409,600 parameters per layer), '
    'risking overfitting with only 3,782 training edges. CompGCN learns one embedding per relation (25 x 128 = '
    '3,200) plus one shared weight matrix (16,384), dramatically fewer parameters while capturing relation semantics.'
)
add_bullet(doc, 'Implementation: Pure PyTorch. 2 layers, 128 hidden dim, subtraction composition, dropout=0.3, lr=1e-3.')
add_bullet(doc, 'Training: 200 max epochs, early stop at 140, best at 120. 32M parameters. HPC GPU, ~46 minutes.')
add_bullet(doc, 'Best result: AUROC 0.9717 with XGBoost decoder (best overall).')

doc.add_heading('6.6 XGBoost Decoder', level=2)
doc.add_paragraph('All methods use the same XGBoost decoder for fair comparison. Feature vector per Drug-Disease pair:')
add_bullet(doc, 'Hadamard product: emb_drug * emb_disease (element-wise dimensional interactions)')
add_bullet(doc, 'Absolute difference: |emb_drug - emb_disease| (dimensional distances)')
add_bullet(doc, 'Cosine similarity: 1 scalar feature')
add_bullet(doc, 'L2 distance: 1 scalar feature')
add_bullet(doc, '6 structural features: shared neighbors, Jaccard coefficient, Adamic-Adar index, log preferential attachment, log degree of drug, log degree of disease')
add_bullet(doc, 'Total: 264 features for 128-dim (CompGCN), 520 for 256-dim (RotatE)')
doc.add_paragraph('XGBoost params: n_estimators=300, max_depth=6, lr=0.1, early_stopping_rounds=20.')

doc.add_heading('6.7 Results Summary', level=2)
add_table(doc, ['Method', 'Dim', 'Decoder', 'AUROC', 'AUPRC', 'Hits@100', 'Hits@200'], [
    ['Node2Vec', '128', 'Cosine', '0.7195', '0.7142', '25.2%', '45.0%'],
    ['Node2Vec', '128', 'XGBoost', '0.9504', '0.9579', '31.1%', '61.8%'],
    ['RotatE', '256', 'Cosine', '0.5299', '0.5401', '19.3%', '32.3%'],
    ['RotatE', '256', 'XGBoost', '0.9652', '0.9655', '31.1%', '60.0%'],
    ['CompGCN', '128', 'Cosine', '0.5058', '0.5041', '16.9%', '30.2%'],
    ['CompGCN', '128', 'XGBoost', '0.9717', '0.9709', '30.5%', '60.6%'],
], col_widths=[2.5, 1.2, 2, 1.8, 1.8, 2, 2])

doc.add_heading('6.8 Metric Interpretation', level=2)
add_bullet(doc, 'AUROC: Probability the model ranks a true treatment above a random non-treatment pair. 0.97 = 97% of the time. Inflated by "easy" random negatives.')
add_bullet(doc, 'Hits@K: Fraction of true treatments in the top K predictions. Maps directly to a practical screening scenario. ~31% of true treatments in top 100 candidates.')
add_bullet(doc, 'For drug repurposing, Hits@K is more meaningful than AUROC because the task is retrieval, not classification.')

doc.add_heading('6.9 Clinical Validation of Top Predictions', level=2)
doc.add_paragraph(
    'For both RotatE and CompGCN, 4 out of 30 top predictions (13.3%) have supporting clinical trial evidence '
    'in CardioKB. This validation rate is notable because these predictions are for Drug-Disease pairs NOT in '
    'the training data.'
)
add_table(doc, ['Method', 'Prediction', 'Clinical Trial'], [
    ['RotatE', 'ACE Inhibitors → heart disease', 'NCT00224809'],
    ['RotatE', 'Atorvastatin → cardiomyopathy', 'NCT00317967'],
    ['RotatE', 'Clopidogrel → cerebrovascular disease', 'NCT01823185'],
    ['RotatE', 'Colchicine → hypertension', 'NCT04916522'],
    ['CompGCN', 'Levosimendan → coronary artery disease', 'NCT00130871'],
    ['CompGCN', 'ACE Inhibitors → heart disease', 'NCT00224809'],
    ['CompGCN', 'Adenosine → myocardial infarction', 'NCT00781404'],
    ['CompGCN', 'Methylprednisolone → coronary artery disease', 'NCT07101367'],
], col_widths=[2.5, 7.5, 4])

doc.add_heading('6.10 Leakage Prevention', level=2)
add_bullet(doc, 'Stratified split: all 25 edge types split proportionally')
add_bullet(doc, 'Negative sampling against full graph (not just train split) to avoid false negatives')
add_bullet(doc, 'Train-only embedding training: val/test edges removed before training')
add_bullet(doc, 'XGBoost early stopping on validation; test set touched only for final evaluation')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 7. WEB INTERFACE WALKTHROUGH
# ══════════════════════════════════════════════════════════
doc.add_heading('7. Web Interface Walkthrough', level=1)
doc.add_paragraph(
    'CardioKB provides a single-page web application (interface/index.html) served by the Flask backend '
    'on port 5050. The interface uses vis.js for network visualization, Chart.js for analytics charts, '
    'and a custom dark/light theme system. This section walks through every feature a user encounters.'
)

doc.add_heading('7.1 Application Layout', level=2)
doc.add_paragraph(
    'The UI is divided into two main regions:'
)
add_bullet(doc, 'Left Sidebar (fixed, 260px wide): Contains the project title, a brief description, Admin tools (System Check button), Tools section (Extract Disease Subgraph), and a light/dark theme toggle.')
add_bullet(doc, 'Main Content Area: Shows the dashboard statistics bar at the top, followed by two tabs (Explore and Query), with a collapsible Admin panel at the bottom.')

doc.add_heading('7.2 Dashboard Statistics Bar', level=2)
doc.add_paragraph(
    'At the top of the main content area, five stat cards display live graph statistics fetched from the '
    '/api/graph-stats endpoint when the page loads:'
)
add_table(doc, ['Card', 'Description', 'Current Value'], [
    ['Total Nodes', 'Count of all entities in the knowledge graph', '459,092'],
    ['Total Relationships', 'Count of all connections between entities', '5,443,134'],
    ['Node Types', 'Distinct entity categories', '17'],
    ['Rel Types', 'Distinct relationship types', '27'],
    ['Data Sources', 'Number of integrated databases', '22'],
], col_widths=[3, 9, 3])
doc.add_paragraph(
    'Each card shows a loading spinner while data is being fetched, then animates to the final count. '
    'Hovering over any card shows a tooltip explaining what the metric represents.'
)

doc.add_heading('7.3 Explore Tab', level=2)
doc.add_paragraph(
    'The Explore tab is the primary graph visualization interface. It lets users search for any disease, '
    'gene, or drug and visualize its immediate neighborhood as an interactive network graph.'
)

doc.add_heading('Search Bar and Autocomplete', level=3)
doc.add_paragraph(
    'The search bar at the top of the Explore tab provides real-time autocomplete as the user types. '
    'The autocomplete dropdown shows matching entities with their type (Disease, Gene, Drug) indicated '
    'by a color-coded badge. Users can:'
)
add_bullet(doc, 'Type a partial name (e.g., "atrial") and select from matching results')
add_bullet(doc, 'Set a Max limit (default 200) to control how many nodes are displayed')
add_bullet(doc, 'Click the "?" button to open a methodology modal explaining how to read the graph')
add_bullet(doc, 'Use quick-explore buttons ("atrial fibrillation", "APOE", "Metoprolol", "coronary artery disease") when the graph is empty')

doc.add_heading('Graph Visualization', level=3)
doc.add_paragraph(
    'The graph is rendered in a large container (70% viewport height, minimum 450px) using vis.js with '
    'force-directed physics layout. Key visual conventions:'
)
add_table(doc, ['Visual Element', 'Meaning'], [
    ['Node shape: Circle', 'All entity types except Disease'],
    ['Node shape: Diamond', 'Disease nodes'],
    ['Node with white border + glow', 'The searched entity (center of the graph)'],
    ['Node color', 'Indicates entity type (blue=Gene, red=Disease, green=Drug, yellow=Pathway, purple=Phenotype, orange=ClinicalTrial, teal=BodyPart, pink=SideEffect, etc.)'],
    ['Node size', 'Proportional to importance/degree in the subgraph'],
    ['Solid gray edge', 'Known relationship from a curated database'],
    ['Dashed cyan edge', 'ML-predicted drug-disease link (hidden by default)'],
    ['Solid orange edge', 'Disease hierarchy (diseaseIsSubtypeOf from Disease Ontology)'],
], col_widths=[4, 12])

doc.add_heading('Node Type Filter Chips', level=3)
doc.add_paragraph(
    'Above the graph, color-coded filter chips allow users to show/hide specific node types. Each chip '
    'corresponds to one of the 17 node types. Chips are active (full opacity, bordered) by default and '
    'can be toggled off (dimmed) to declutter the visualization. This is useful when a disease has many '
    'gene connections and the user wants to focus on drugs or pathways instead.'
)

doc.add_heading('Edge Type Filter Chips', level=3)
doc.add_paragraph(
    'A second row of filter chips below the node type row allows filtering by relationship type. '
    'Each chip represents one of the relationship types present in the current subgraph.'
)

doc.add_heading('ML Predictions Toggle', level=3)
doc.add_paragraph(
    'A labeled checkbox ("Show ML Predictions") controls visibility of the 1,500 predicted drug-disease '
    'edges. When enabled:'
)
add_bullet(doc, 'Cyan dashed edges appear connecting drugs to diseases')
add_bullet(doc, 'A Drug Repurposing Predictions panel expands below the toggle with method filters (RotatE, CompGCN), a sortable predictions table, and CSV/JSON export buttons')
add_bullet(doc, 'Users can filter by embedding method and view confidence scores')
add_bullet(doc, 'An "Experimental" badge and disclaimer remind users these are computational predictions')

doc.add_heading('Detail Panel (Node Inspector)', level=3)
doc.add_paragraph(
    'Clicking any node opens a slide-in panel on the right side of the graph container (320px wide). '
    'The panel displays:'
)
add_bullet(doc, 'Node name and type (with color-coded type badge)')
add_bullet(doc, 'All node properties (e.g., geneSymbol, drugId, diseaseName, specificityScore)')
add_bullet(doc, 'Connected Nodes list showing all neighbors with their types and relationship labels')
add_bullet(doc, 'A close button (x) to dismiss the panel')

doc.add_heading('Edge Provenance Panel', level=3)
doc.add_paragraph(
    'Clicking any edge opens a separate panel titled "Why is this edge here?" showing:'
)
add_bullet(doc, 'The source database (e.g., "Source: DrugBank")')
add_bullet(doc, 'All edge properties (e.g., combinedScore: 850, morScore: 1, confidence: 0.93)')
add_bullet(doc, 'For predicted edges: the prediction method (RotatE/CompGCN), confidence score, and a red "not clinically validated" warning')

doc.add_heading('Node Legend', level=3)
doc.add_paragraph(
    'Below the graph, a legend strip shows all node types with their color-coded dots, so users can '
    'quickly identify what each color represents without hovering over individual nodes.'
)

doc.add_heading('Visual Guide (Collapsible)', level=3)
doc.add_paragraph(
    'A collapsible "Visual Guide" panel between the filter chips and the graph explains the node shapes, '
    'edge styles, and color conventions in a two-column layout (Nodes column and Edges column).'
)

doc.add_heading('About Panel (Collapsible)', level=3)
doc.add_paragraph(
    'Below the graph, an expandable "About: How to Interpret This Graph" panel provides in-depth explanations:'
)
add_bullet(doc, 'What is this graph? — Explanation of knowledge graphs and how CardioKB connects biological entities')
add_bullet(doc, 'Disease-Specificity Score — Formula, interpretation tiers (1.0 = highly specific, 0.02 = broadly connected), and why it matters for drug target identification')
add_bullet(doc, 'ML-Predicted Drug-Disease Links — Full pipeline explanation (embedding → features → XGBoost → prediction), performance table (AUROC, AUPRC, Hits@K), and important limitations disclaimer')
add_bullet(doc, 'Data Sources — Key sources listed in a grid with brief descriptions')

doc.add_heading('Export Options', level=3)
doc.add_paragraph('Below the graph and About panel, an export bar provides four download options:')
add_bullet(doc, 'Export CSV — Download current subgraph as CSV for analysis in Excel, R, or Python')
add_bullet(doc, 'Export JSON — Download as JSON for programmatic use or Cytoscape import')
add_bullet(doc, 'Export PNG — Save current graph visualization as a PNG image')
add_bullet(doc, 'Export PDF — Print current view to PDF via browser print dialog')

doc.add_heading('7.4 Query Tab', level=2)
doc.add_paragraph(
    'The Query tab provides a powerful interface for querying the knowledge graph using either natural '
    'language (AI-powered) or direct Cypher queries.'
)

doc.add_heading('Natural Language Querying (CypherGPT)', level=3)
doc.add_paragraph(
    'At the top of the Query tab, a prominent text input allows users to ask questions in plain English. '
    'The system uses CypherGPT (adapted from Eng2Cypher by Jay Moran, Center for AI Research) to translate '
    'questions into Cypher queries. The pipeline:'
)
add_bullet(doc, 'Introspects the live graph schema (node types, relationship types, properties)')
add_bullet(doc, 'Builds a context-aware prompt with CardioKB-specific instructions and examples')
add_bullet(doc, 'Generates Cypher via Claude API')
add_bullet(doc, 'Auto-validates and corrects the query using fuzzy matching against known labels')
add_bullet(doc, 'Handles cross-node-type medical conditions (Disease/Phenotype/SideEffect) via UNION ALL')

doc.add_paragraph('Example queries users can type:')
add_bullet(doc, '"What drugs target genes associated with heart failure?"')
add_bullet(doc, '"Show me clinical trials studying atrial fibrillation"')
add_bullet(doc, '"Which genes interact with APOE?"')

doc.add_heading('Query Templates', level=3)
doc.add_paragraph(
    'Below the NL input, a set of clickable template buttons provide pre-built Cypher queries for common '
    'tasks. Clicking a template populates the Cypher editor and optionally auto-runs the query.'
)

doc.add_heading('Direct Cypher Editor', level=3)
doc.add_paragraph(
    'Expandable via "Or write Cypher directly," a full-featured textarea allows writing and executing '
    'arbitrary Cypher queries against Memgraph. Features:'
)
add_bullet(doc, 'Monospace font (SF Mono / Fira Code) for code readability')
add_bullet(doc, 'Ctrl+Enter keyboard shortcut to run queries')
add_bullet(doc, 'Read-only mode: no writes allowed (CREATE, DELETE, etc. are blocked)')
add_bullet(doc, 'Save Query button to bookmark frequently used queries')
add_bullet(doc, 'Saved Queries section appears when queries have been bookmarked')

doc.add_heading('Multi-Panel Results (Neo4j Browser Style)', level=3)
doc.add_paragraph(
    'Each query execution appends a new result panel to a vertical stack, rather than replacing previous '
    'results. This allows side-by-side comparison of different queries. Each panel includes:'
)
add_bullet(doc, 'Collapsible header showing the Cypher query text, row count, and execution time')
add_bullet(doc, 'Table view tab with a scrollable results table (sticky headers, sortable columns)')
add_bullet(doc, 'Graph view tab with a vis.js network visualization of returned nodes and relationships')
add_bullet(doc, 'Stats footer showing node count, relationship count, and query duration')
add_bullet(doc, 'Close button (x) to dismiss individual panels')
add_bullet(doc, 'Clear All button to remove all result panels at once')

doc.add_heading('7.5 Admin Panel', level=2)
doc.add_paragraph(
    'At the bottom of the main content area, a collapsible "Admin: Parser Status & Pipeline Log" panel '
    'provides system administration features:'
)

doc.add_heading('Parser Status Table', level=3)
doc.add_paragraph(
    'A two-column layout shows parser status on the left. The table lists every data source parser with '
    'its current status (Working/Failed/Skipped badge) and duration. Status is determined dynamically '
    'by querying the graph for each parser\'s expected nodes and relationships.'
)

doc.add_heading('Health Checks', level=3)
doc.add_paragraph(
    'The right column displays health check results after clicking "System Check" in the sidebar. '
    'Each check shows a green (OK) or red (warning) indicator for:'
)
add_bullet(doc, 'Node count validation per type')
add_bullet(doc, 'Edge count validation per relationship type')
add_bullet(doc, 'Source label coverage (all 19 expected source labels present)')
add_bullet(doc, 'Edge property coverage (combinedScore, morScore, expressionScore, etc.)')
add_bullet(doc, 'Database connectivity')

doc.add_heading('Charts', level=3)
doc.add_paragraph(
    'Two Chart.js bar charts visualize node counts by type and relationship counts by type, providing '
    'an at-a-glance view of graph composition.'
)

doc.add_heading('ID Mapping Report', level=3)
doc.add_paragraph(
    'A detailed table showing match rates for ID mappings between TSV relationship files and existing '
    'graph nodes. Identifies where edges are being lost due to unresolvable identifiers.'
)

doc.add_heading('Pipeline Log', level=3)
doc.add_paragraph(
    'A monospace log panel displays real-time pipeline output with color-coded lines: '
    'green for success, red for errors, blue for informational messages.'
)

doc.add_heading('7.6 Sidebar Tools', level=2)

doc.add_heading('System Check', level=3)
doc.add_paragraph(
    'The "System Check" button in the Admin section triggers a comprehensive health check of all data '
    'sources, database connectivity, and data integrity. A "?" info button next to it explains what the '
    'check covers.'
)

doc.add_heading('Extract Disease Subgraph', level=3)
doc.add_paragraph(
    'The Tools section provides N-hop subgraph extraction with the following controls:'
)
add_bullet(doc, 'Disease name input with placeholder examples (heart failure, atrial fibrillation, coronary artery disease)')
add_bullet(doc, 'Hops slider (1-3): Controls extraction depth. A "?" info button explains each hop level — 1 hop for direct connections, 2 hops for connections-of-connections, 3 hops for broad hypothesis generation')
add_bullet(doc, '"Build Subgraph" button: Queries Memgraph for the complete neighborhood and displays stats')
add_bullet(doc, 'Export JSON / Export CSV buttons: Download the extracted subgraph for downstream analysis in Excel, R, Python, or Cytoscape')
add_bullet(doc, 'An "Instant" badge indicates this queries existing data only — no new data is fetched')

doc.add_heading('Theme Toggle', level=3)
doc.add_paragraph(
    'A "Toggle Light / Dark" button at the bottom of the sidebar switches between a dark theme '
    '(slate blue background, ideal for extended use) and a light theme (white background, better for '
    'screenshots and presentations). The theme is applied via CSS custom properties and persists during '
    'the session.'
)

doc.add_heading('7.7 Visual Design and Theming', level=2)
doc.add_paragraph(
    'The interface uses a modern, minimal design system built entirely with CSS custom properties:'
)
add_table(doc, ['Element', 'Dark Theme', 'Light Theme'], [
    ['Background', '#0f172a (deep navy)', '#f1f5f9 (light gray)'],
    ['Surface (cards, panels)', '#1e293b', '#ffffff'],
    ['Text', '#e2e8f0 (off-white)', '#1e293b (dark gray)'],
    ['Accent color', '#38bdf8 (sky blue)', '#0284c7 (darker blue)'],
    ['Borders', '#475569', '#cbd5e1'],
    ['Success indicators', '#4ade80 (green)', '#16a34a (darker green)'],
    ['Error indicators', '#f87171 (red)', '#dc2626 (darker red)'],
], col_widths=[4, 6, 6])

doc.add_heading('Welcome Tour Modal', level=3)
doc.add_paragraph(
    'On first visit, a three-step welcome tour modal introduces the application:'
)
add_bullet(doc, 'Step 1: Overview of CardioKB (what it is, live stats, disease focus)')
add_bullet(doc, 'Step 2: Explore & Query tabs (how to search, click nodes/edges, use filters)')
add_bullet(doc, 'Step 3: ML Drug Repurposing Predictions (embedding methods, toggle, limitations)')
doc.add_paragraph('Users can navigate with Next/Previous buttons, click dots to jump to a step, or Skip to dismiss.')

doc.add_heading('Tooltip System', level=3)
doc.add_paragraph(
    'A custom JavaScript tooltip engine provides context-sensitive help throughout the interface. '
    'Hovering over any element with a data-tip attribute shows a positioned tooltip bubble. '
    'Additionally, "?" info buttons use a click-triggered popover system with richer HTML content '
    '(multi-paragraph explanations, color-coded tiers, code examples).'
)

doc.add_heading('Responsive Design', level=3)
doc.add_paragraph(
    'The interface adapts to different screen sizes: at viewport widths below 768px, the sidebar '
    'collapses to a full-width horizontal header, the detail panel becomes inline rather than overlaid, '
    'and the two-column admin grid stacks vertically.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 8. DEPLOYMENT
# ══════════════════════════════════════════════════════════
doc.add_heading('8. Deployment', level=1)

doc.add_heading('8.1 Architecture', level=2)
add_table(doc, ['Service', 'Image', 'Port', 'Purpose'], [
    ['memgraph', 'memgraph/memgraph:latest', '7687, 7444', 'Graph database with persistent named volume'],
    ['app', 'Custom (Dockerfile)', '5050', 'Flask web app (Python 3.11-slim)'],
], col_widths=[2.5, 5, 2.5, 6])

doc.add_heading('8.2 Flask API Endpoints', level=2)
add_table(doc, ['Endpoint', 'Method', 'Description'], [
    ['/api/graph-stats', 'GET', 'Live node/relationship counts from Memgraph'],
    ['/api/query', 'POST', 'Execute arbitrary Cypher queries'],
    ['/api/explore', 'GET', 'Disease subgraph exploration with specificity ranking'],
    ['/api/subgraph', 'POST', 'N-hop disease subgraph extraction'],
    ['/api/agent/build-disease-graph', 'POST', 'SSE-streamed disease graph building with Claude API'],
    ['/api/specificity-info', 'GET', 'Specificity score metadata and timestamp'],
    ['/api/nl2cypher', 'POST', 'Natural language to Cypher translation via Claude API'],
], col_widths=[6, 1.5, 8.5])

doc.add_heading('8.3 Graph Data Transfer', level=2)
doc.add_paragraph(
    'Graph data (459K nodes, 5.4M edges, ~14 GB uncompressed) is transferred between machines via '
    'Memgraph volume backups (~1.2 GB compressed). This is faster than Cypher-based import which would '
    'need to replay millions of CREATE/MERGE statements.'
)
add_bullet(doc, 'Export: scripts/export_graph.sh stops Memgraph, tars /var/lib/memgraph from the Docker volume, restarts')
add_bullet(doc, 'Import: scripts/import_graph.sh creates the Docker volume, extracts the archive, starts Memgraph with readiness check')

doc.add_heading('8.4 Environment Variables', level=2)
add_table(doc, ['Variable', 'Purpose'], [
    ['MEMGRAPH_URI, MEMGRAPH_USERNAME, MEMGRAPH_PASSWORD', 'Graph database connection'],
    ['ANTHROPIC_API_KEY', 'Claude API for AI features (NL-to-Cypher, Build KG)'],
    ['ANTHROPIC_FOUNDRY_API_KEY + BASE_URL', 'Azure AI Foundry (preferred over direct API)'],
    ['ADMIN_PASSWORD', 'Admin UI features (pipeline run, add database)'],
    ['DRUGBANK_USERNAME, DRUGBANK_PASSWORD', 'Pipeline only (optional, for DrugBank XML)'],
], col_widths=[7, 9])

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 9. EVALUATION
# ══════════════════════════════════════════════════════════
doc.add_heading('9. Evaluation', level=1)

doc.add_heading('9.1 Automated Health Check System', level=2)
doc.add_paragraph(
    'CardioKB includes an automated evaluation system (src/admin_agent.py) that performs dynamic, '
    'graph-based parser status detection after each pipeline run:'
)
add_bullet(doc, 'Node count validation per type against expected ranges')
add_bullet(doc, 'Edge count validation per relationship type')
add_bullet(doc, 'Source label coverage (all 19 expected labels)')
add_bullet(doc, 'Edge property coverage (combinedScore, morScore, expressionScore, etc.)')
add_bullet(doc, 'Orphan rate calculation per node type')

doc.add_heading('9.2 Orphan Rates (Expected)', level=2)
add_table(doc, ['Node Type', 'Orphan Rate', 'Reason'], [
    ['Symptom', '~100%', 'Only MEDLINE provides symptom edges (skipped). MeSH provides nodes only.'],
    ['DrugLabel', '~100%', 'ClinPGx labels whose target genes/drugs are not in the graph.'],
    ['Disease', '~76%', 'Full Disease Ontology loaded (3,442) but CVD-filtered edges connect only ~813.'],
    ['Gene', '~78%', 'Full NCBI catalog (193K) loaded; CVD-filtered edges connect only ~43K.'],
], col_widths=[3, 2.5, 10.5])
doc.add_paragraph(
    'High orphan rates are a deliberate design choice: loading complete node catalogs ensures no edges '
    'are lost due to missing target nodes.'
)

doc.add_heading('9.3 Known Gaps', level=2)
add_bullet(doc, 'ClinVar variant-disease yield: Only 0.12% of ClinVar associations map to CardioKB diseases (212 CUI overlap)')
add_bullet(doc, 'BindingDB: No binding affinity values loaded as edge properties')
add_bullet(doc, 'CTD: No PubMed IDs loaded as edge properties')
add_bullet(doc, 'MEDLINE: 0 edges (configs skipped in current build)')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 10. FUTURE DIRECTIONS
# ══════════════════════════════════════════════════════════
doc.add_heading('10. Future Directions', level=1)

doc.add_heading('10.1 GNN Next Steps', level=2)
add_bullet(doc, 'Attention-based GNNs (GAT, HGT): Learn attention weights over neighbor types to identify which relationships are most informative for drug repurposing')
add_bullet(doc, 'Deeper models: Current CompGCN uses 2 layers (2-hop). Additional layers could capture longer-range dependencies but risk oversmoothing')
add_bullet(doc, 'Edge-type-specific prediction heads: Separate decoders for drug repurposing, side effect prediction, gene-disease association')
add_bullet(doc, 'Temporal modeling: Incorporate clinical trial phase/status as temporal signals')

doc.add_heading('10.2 CypherGPT Natural Language Layer (Implemented)', level=2)
doc.add_paragraph(
    'The Query tab now includes an AI-powered natural language interface adapted from CypherGPT/Eng2Cypher. '
    'Schema is cached with 24-hour TTL. CardioKB-specific instructions and examples guide the LLM.'
)

doc.add_heading('10.3 drugTreatsDisease Enrichment', level=2)
doc.add_paragraph(
    'DrugBank XML files contain free-text drug descriptions with treatment indications not captured as '
    'structured edges. NLP extraction could increase drugTreatsDisease edges by 2-5x, improve ML training '
    'data diversity, and add indication-level detail as edge properties.'
)

doc.add_heading('10.4 ASAREE/Persona Experiments', level=2)
add_bullet(doc, 'Evidence quality scoring: Automatically assess gene-disease association strength based on supporting edge types')
add_bullet(doc, 'Persona-based querying: Different user personas (clinician, molecular biologist, pharmacologist) see tailored subgraph views')
add_bullet(doc, 'Automated hypothesis generation: Rank novel biological hypotheses based on graph topology patterns')

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'CARDIO_KB_STUDY_GUIDE.docx')
doc.save(output_path)
print(f'Saved to {output_path}')
