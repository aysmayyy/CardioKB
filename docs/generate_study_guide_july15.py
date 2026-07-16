"""Generate CardioKB Complete Study Guide as a formatted Word document.

Final build edition (July 15, 2026): Comprehensive standalone reference covering
the full state of the CardioKB project for advisor discussion and paper defense.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1E, 0x40, 0x7C)

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(13)


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p


def bold_para(bold_text, normal_text):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    p.add_run(normal_text)
    return p


# ══════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('CardioKB Complete Study Guide')
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0x1E, 0x40, 0x7C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Final Build — July 15, 2026\nComprehensive Reference for Advisor Discussion and Paper Defense')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = author.add_run('Asma Nawaz')
r.font.size = Pt(13)
r.bold = True

affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = affil.add_run('Cedars-Sinai Medical Center / Computational Biomedicine')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ('1. Project Overview', 0),
    ('2. Graph Architecture — Node Types', 0),
    ('3. Graph Architecture — Edge Types', 0),
    ('4. Data Sources (All 23)', 0),
    ('4.1 Direct Parsers (5)', 1),
    ('4.2 Hetionet-Derived Component Parsers (15)', 1),
    ('4.3 Additional Sources (3)', 1),
    ('5. Edge Properties', 0),
    ('6. Drug Node Entity Resolution', 0),
    ('7. ML Pipeline — Link Prediction for Drug Repurposing', 0),
    ('7.1 Methodology', 1),
    ('7.2 Full Results Table', 1),
    ('7.3 Training Details', 1),
    ('7.4 Predictions in Graph', 1),
    ('7.5 Why CompGCN Is Primary', 1),
    ('8. Key Metrics Explained', 0),
    ('9. Methodology Notes', 0),
    ('10. Known Limitations', 0),
    ('11. Web Interface Features', 0),
    ('12. Deployment', 0),
    ('13. drugTreatsDisease Sources Breakdown', 0),
]
for item, indent in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(1.27)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════
doc.add_heading('1. Project Overview', level=1)

doc.add_paragraph(
    'CardioKB is a CVD-focused biomedical knowledge graph built as a 12-week rotation project '
    '(January–April 2026). It integrates 23 deduplicated data sources into a single queryable '
    'graph database for disease research, feature selection, and precision medicine.'
)

doc.add_heading('Graph at a Glance', level=2)
add_table(
    ['Metric', 'Value'],
    [
        ['Total Nodes', '453,037'],
        ['Total Relationships', '5,461,783'],
        ['Node Types', '17'],
        ['Relationship Types', '28'],
        ['Data Sources', '23 + 2 ML prediction sources'],
        ['Drug Nodes (deduplicated)', '26,794 (from 32,849 pre-merge)'],
        ['ML Predicted Edges', '14,435 (6,607 CompGCN + 7,828 RotatE)'],
        ['Graph Data Export', '304 MB tar.gz'],
    ],
    col_widths=[5, 11]
)

doc.add_heading('Technology Stack', level=2)
bullet('Database: Memgraph (in-memory graph database, Cypher query language, Bolt protocol)')
bullet('Pipeline: BaseAgent multi-agent orchestration (Ontology, Engineer, Mapping, Evaluator agents)')
bullet('Backend: Flask (Python 3.11), REST API with SSE streaming')
bullet('Frontend: Single-page web UI (vis.js graph visualization, Chart.js analytics)')
bullet('ML: PyKEEN (RotatE), PyTorch (CompGCN), XGBoost (decoder)')
bullet('Deployment: Docker Compose (Memgraph + Flask app)')
bullet('HPC: SLURM job scheduler for GPU-based embedding training')

doc.add_heading('Core Capabilities', level=2)
bullet('Unified graph where a single Cypher query traverses from drugs to gene targets to diseases to clinical trials across all 23 sources')
bullet('Automated ETL pipeline that rebuilds the entire graph from source data in approximately 5 minutes')
bullet('Interactive web interface with AI-powered natural language querying (NL2Cypher via Claude), graph visualization, and disease subgraph extraction')
bullet('ML drug repurposing pipeline using knowledge graph embeddings (CompGCN, RotatE) to predict new drug-disease treatment relationships with AUROC up to 0.9865')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 2. GRAPH ARCHITECTURE — NODE TYPES
# ══════════════════════════════════════════════════════════
doc.add_heading('2. Graph Architecture — Node Types (17)', level=1)

doc.add_paragraph(
    'The graph contains 453,037 nodes across 17 distinct types. Every node carries type-specific '
    'properties and cross-reference identifiers for linking back to source databases.'
)

add_table(
    ['Node Type', 'Count', 'Primary Source'],
    [
        ['Gene', '193,795', 'NCBI Gene'],
        ['Variant', '135,555', 'ClinVar'],
        ['Drug', '26,794', 'DrugBank + CTD (deduplicated)'],
        ['BiologicalProcess', '24,428', 'Gene Ontology'],
        ['ClinicalTrial', '21,578', 'ClinicalTrials.gov'],
        ['Phenotype', '19,389', 'HPO'],
        ['MolecularFunction', '10,056', 'Gene Ontology'],
        ['GeneFamily', '4,257', 'HGNC'],
        ['CellularComponent', '4,076', 'Gene Ontology'],
        ['Disease', '3,442', 'Disease Ontology'],
        ['Pathway', '2,870', 'Reactome'],
        ['PharmacologicClass', '2,359', 'DrugCentral'],
        ['SideEffect', '2,227', 'SIDER'],
        ['BodyPart', '1,400', 'Uberon'],
        ['Symptom', '415', 'MeSH'],
        ['TranscriptionFactor', '367', 'DoRothEA'],
        ['DrugLabel', '29', 'ClinPGx'],
    ],
    col_widths=[3.5, 2, 10.5]
)

doc.add_paragraph(
    'Gene nodes dominate (42.8%), followed by Variant (29.9%) and Drug (5.9%). Gene is the central '
    'hub of the graph — most other entity types connect through genes. The full NCBI Gene and '
    'Disease Ontology catalogs are loaded to ensure no edges are lost due to missing target nodes, '
    'resulting in expected orphan rates (~78% Gene, ~76% Disease).'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 3. GRAPH ARCHITECTURE — EDGE TYPES
# ══════════════════════════════════════════════════════════
doc.add_heading('3. Graph Architecture — Edge Types (28)', level=1)

doc.add_paragraph(
    'The graph contains 5,461,783 relationships across 28 types. Every relationship carries a '
    'source property identifying the originating database (e.g., source: "DrugBank").'
)

add_table(
    ['Relationship Type', 'Count', 'Source(s)'],
    [
        ['bodyPartOverexpressesGene', '2,749,193', 'Bgee'],
        ['geneAssociatesWithDisease', '542,096', 'PubTator + OpenTargets'],
        ['chemicalIncreasesExpression', '343,783', 'CTD'],
        ['chemicalDecreasesExpression', '328,708', 'CTD'],
        ['geneAssociatesWithPhenotype', '270,265', 'HPO'],
        ['geneInteractsWithGene', '229,007', 'STRING'],
        ['geneInPathway', '137,116', 'Reactome'],
        ['variantInGene', '135,393', 'ClinVar'],
        ['geneParticipatesInBiologicalProcess', '122,117', 'Gene Ontology'],
        ['geneAssociatedWithCellularComponent', '90,141', 'Gene Ontology'],
        ['geneHasMolecularFunction', '76,612', 'Gene Ontology'],
        ['compoundUpregulatesGene', '74,854', 'CTD'],
        ['compoundCausesSideEffect', '67,646', 'SIDER'],
        ['compoundDownregulatesGene', '64,661', 'CTD'],
        ['variantAssociatedWithDisease', '51,323', 'ClinVar'],
        ['drugBindsGene', '29,363', 'DrugBank'],
        ['geneInFamily', '27,022', 'HGNC'],
        ['compoundInPharmacologicClass', '24,752', 'DrugCentral'],
        ['chemicalBindsGene', '22,735', 'BindingDB'],
        ['STUDIES_CONDITION', '20,667', 'ClinicalTrials.gov'],
        ['transcriptionFactorInteractsWithGene', '15,082', 'DoRothEA'],
        ['predictedTreatsDisease', '14,435', 'CompGCN + RotatE'],
        ['hasVariant', '8,413', 'ClinVar'],
        ['drugTreatsPhenotype', '5,714', 'DrugBank_Indications'],
        ['drugTreatsDisease', '4,852', 'CTD + DrugBank_Indications + DrugCentral + ClinicalTrials.gov'],
        ['TESTS_INTERVENTION', '3,178', 'ClinicalTrials.gov'],
        ['diseaseIsSubtypeOf', '2,581', 'Disease Ontology'],
        ['AFFECTS_RESPONSE_TO', '74', 'ClinPGx'],
    ],
    col_widths=[4.5, 2, 9.5]
)

doc.add_paragraph(
    'bodyPartOverexpressesGene edges dominate (50.3% of all edges, from Bgee tissue-specific '
    'gene expression data). The drugTreatsDisease edge type (4,852 edges) is the primary ML '
    'training signal for drug repurposing link prediction.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 4. DATA SOURCES
# ══════════════════════════════════════════════════════════
doc.add_heading('4. Data Sources (All 23)', level=1)

doc.add_paragraph(
    'CardioKB integrates 23 data sources, each contributing unique node types and/or relationship '
    'types. Sources were selected based on three criteria: (1) authoritative coverage of a specific '
    'biological entity type, (2) no redundancy with other included sources, and (3) public '
    'accessibility. During a systematic deduplication audit, 14 redundant sources were removed '
    '(DisGeNET, GWAS Catalog, Jensen DISEASES, Jensen TISSUES, MEDLINE, OMIM, WikiPathways, '
    'AOP-DB, HGNC base, CellAge, GenAge, Hetionet precomputed, DrugAge, AnAge).'
)

doc.add_heading('4.1 Direct Parsers (5 sources)', level=2)

add_table(
    ['#', 'Source', 'Access', 'Key Contribution'],
    [
        ['1', 'ClinicalTrials.gov', 'Public API v2', '21,578 ClinicalTrial nodes, 20,667 STUDIES_CONDITION + 3,178 TESTS_INTERVENTION + 147 drugTreatsDisease edges'],
        ['2', 'ClinPGx (PharmGKB successor)', 'Public API', '29 DrugLabel nodes, 1,091 VARIANT_IN, 503 drugLabelAnnotatesGene, 345 drugLabelDescribesDrug, 224 AFFECTS_RESPONSE_TO, 19 AFFECTS_RESPONSE_TO_CLASS edges'],
        ['3', 'NCBI Gene', 'Public FTP', '193,687 Gene nodes (node-only source, provides the Gene node backbone)'],
        ['4', 'DoRothEA (OmniPath)', 'Public API', '367 TranscriptionFactor nodes, 12,985 TF-gene interactions with morScore + confidence'],
        ['5', 'DrugBank', 'XML file', '19,842 Drug nodes, 12,089 drugBindsGene edges'],
    ],
    col_widths=[0.6, 3.5, 2.2, 9.7]
)

doc.add_heading('4.2 Hetionet-Derived Component Parsers (15 sources)', level=2)

doc.add_paragraph(
    'These parsers were originally part of the Hetionet precomputed dataset but have been broken out '
    'into individual component parsers that each query their respective original data source directly.'
)

add_table(
    ['#', 'Source', 'Key Contribution'],
    [
        ['6', 'Disease Ontology (DOID)', '12,012 Disease nodes, 6,447 diseaseIsSubtypeOf edges (2,581 loaded)'],
        ['7', 'Gene Ontology (GO)', '50,350 BP + 26,935 MF + 25,794 CC edges across 3 ontology domains'],
        ['8', 'Uberon (Anatomy)', '14,937 BodyPart nodes (node-only, used by Bgee)'],
        ['9', 'MeSH (Symptoms)', '966 Symptom nodes (node-only, ~100% orphan rate)'],
        ['10', 'SIDER (Side Effects)', '148,518 compoundCausesSideEffect edges (legacy, no live API)'],
        ['11', 'LINCS L1000', '150,535 gene expression edges with zScore (legacy, clue.io requires institutional access)'],
        ['12', 'DrugCentral', '16,403 pharmacologic class + 245 treats + 96 palliates edges (CUI-to-DOID mapped)'],
        ['13', 'BindingDB', '12,250 chemicalBindsGene edges (experimentally measured binding)'],
        ['14', 'PubTator Central', '744,427 geneAssociatesWithDisease + 4,320 diseaseAssociatesWithDisease edges (CVD AND-filter)'],
        ['15', 'CTD', '4,572 unique Drug nodes, 116,451 chemIncreasesExp + 97,951 chemDecreasesExp edges, 3,099 drugTreatsDisease'],
        ['16', 'Bgee (Gene Expression Atlas)', '784,026 underexpresses + 1,872 overexpresses edges with expressionScore'],
        ['17', 'HPO', '19,389 Phenotype nodes, 162,994 geneAssociatesWithPhenotype edges'],
        ['18', 'Reactome', '44,979 geneInPathway + 44,979 pathwayContainsGene edges'],
        ['19', 'STRING', '121,170 geneInteractsWithGene edges (confidence > 700, with combinedScore)'],
        ['20', 'OpenTargets', '32,826 geneAssociatesWithDisease edges (CVD AND-filter, EFO-to-DOID mapped)'],
    ],
    col_widths=[0.6, 3.5, 11.9]
)

doc.add_heading('4.3 Additional Sources (3)', level=2)

add_table(
    ['#', 'Source', 'Key Contribution'],
    [
        ['21', 'HGNC Gene Families', '1,934 GeneFamily nodes, 5,123 geneInFamily + 5,123 familyContainsGene edges'],
        ['22', 'ClinVar', '4,488,042 Variant nodes, 2,267,095 hasVariant + 2,267,095 variantInGene edges with clinicalSignificance'],
        ['23', 'DrugBank_Indications (text-mined)', '1,449 drugTreatsDisease + 5,714 drugTreatsPhenotype edges (text-mined from DrugBank XML indication fields)'],
    ],
    col_widths=[0.6, 4, 11.4]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 5. EDGE PROPERTIES
# ══════════════════════════════════════════════════════════
doc.add_heading('5. Edge Properties', level=1)

doc.add_paragraph(
    'Seven relationship types carry quantitative properties beyond simple edge existence. These '
    'properties enable weighted queries and filtering by evidence strength.'
)

add_table(
    ['Relationship Type', 'Property', 'Source', 'Range / Description'],
    [
        ['geneInteractsWithGene', 'combinedScore', 'STRING', '0–1000 (filtered > 700 = high confidence). Integrates experimental, database, and text-mining evidence.'],
        ['bodyPartOverexpressesGene', 'expressionScore', 'Bgee', 'Quantitative tissue-specific gene expression level.'],
        ['transcriptionFactorInteractsWithGene', 'morScore', 'DoRothEA', '-1 = represses, 0 = unknown, +1 = activates. Mode of regulation.'],
        ['transcriptionFactorInteractsWithGene', 'confidence', 'DoRothEA', 'A (highest) through D (lowest). Evidence confidence level.'],
        ['geneInPathway', 'evidenceCode', 'Reactome', 'Gene Ontology evidence code for pathway membership.'],
        ['geneAssociatesWithDisease', 'score', 'OpenTargets', '0–1 overall association score from genetic, somatic, literature, and drug evidence.'],
        ['drugBindsGene', 'interactionType', 'DrugBank', 'Pharmacological action type (e.g., inhibitor, agonist, antagonist, substrate).'],
        ['variantAssociatedWithDisease', 'clinicalSignificance', 'ClinVar', 'Pathogenic, Likely Pathogenic, Benign, Likely Benign, or VUS (Variant of Uncertain Significance).'],
    ],
    col_widths=[3.8, 2.2, 1.8, 8.2]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 6. DRUG NODE ENTITY RESOLUTION
# ══════════════════════════════════════════════════════════
doc.add_heading('6. Drug Node Entity Resolution', level=1)

doc.add_paragraph(
    'The BaseAgent pipeline created duplicate Drug nodes when multiple sources (DrugBank, CTD, '
    'ClinPGx, DrugCentral) loaded the same compound under different internal drugId values but '
    'shared the same xrefDrugBank canonical identifier. A post-hoc entity resolution step was '
    'required to deduplicate them.'
)

add_table(
    ['Metric', 'Value'],
    [
        ['Original Drug nodes', '32,849'],
        ['Duplicate groups identified', '5,611 (2x–8x duplication, including salt/ester forms)'],
        ['Duplicate nodes removed', '6,055'],
        ['Edges transferred to survivor nodes', '474,641'],
        ['Redundant edges deduplicated', '9,094'],
        ['Final Drug nodes', '26,794'],
        ['Resolution method', 'xrefDrugBank canonical identifier matching'],
        ['Script', 'scripts/merge_duplicate_drugs.py'],
    ],
    col_widths=[5, 11]
)

doc.add_paragraph(
    'Entity resolution eliminated the cross-source Drug node duplication that previously allowed '
    '~1.5% spurious ML predictions (known treatments leaking via duplicate nodes) and ~8.7% wasted '
    'prediction slots (same drug-disease pair predicted from both copies of a drug). All ML models '
    'were retrained on the post-merge graph.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 7. ML PIPELINE
# ══════════════════════════════════════════════════════════
doc.add_heading('7. ML Pipeline — Link Prediction for Drug Repurposing', level=1)

doc.add_heading('7.1 Methodology', level=2)

bold_para('Pipeline: ', 'export_edges.py → split_edges.py → train embeddings (HPC) → link_prediction*.py')

bold_para('Candidate space: ', '10,310 therapeutic drugs × 2,640 diseases (CompGCN) / 2,296 diseases (RotatE)')

bold_para('Positive edges: ', '4,852 drugTreatsDisease edges split into 4,726 train / 485 validation / 484 test (80/10/10 stratified, seed=42)')

bold_para('Feature vector: ', 'Hadamard product + absolute difference + cosine similarity + L2 distance + structural features (shared neighbors, Jaccard coefficient, Adamic-Adar index, preferential attachment, degree). Total: 264 features for 128-dim (CompGCN), 520 features for 256-dim (RotatE).')

bold_para('Negative sampling: ', '1:1 ratio, excluding ALL known Drug-Disease edges (not just train). Seeds: train=42, val=123, test=456.')

bold_para('Therapeutic drug filter: ', 'Drugs must have edges in {drugBindsGene, compoundInPharmacologicClass, compoundCausesSideEffect, drugTreatsDisease, AFFECTS_RESPONSE_TO, TESTS_INTERVENTION}. Reduces 26,794 to 10,310 therapeutic drugs.')

bold_para('MIN_CONFIDENCE threshold: ', '0.5 (only predictions with XGBoost confidence >= 0.5 stored in graph)')

doc.add_heading('7.2 Full Results Table', level=2)

add_table(
    ['Method', 'Decoder', 'AUROC', 'AUPRC', 'Hits@10', 'Hits@50', 'Hits@100', 'Hits@200', 'MRR', 'Med. Rank'],
    [
        ['RotatE (256-dim)', 'Cosine', '0.7807', '0.7569', '1.7%', '9.3%', '13.4%', '27.3%', '0.0095', '461'],
        ['RotatE (256-dim)', 'XGBoost', '0.9828', '0.9812', '39.3%', '72.1%', '84.1%', '92.8%', '0.1890', '17.5'],
        ['RotatE (256-dim)', 'MLP', '0.9810', '0.9786', '41.1%', '71.9%', '83.5%', '91.9%', '0.1898', '16'],
        ['CompGCN (128-dim)', 'Cosine', '0.3100', '0.3810', '0.2%', '0.2%', '0.4%', '0.6%', '0.0027', '2,230'],
        ['CompGCN (128-dim)', 'XGBoost', '0.9865', '0.9854', '36.6%', '68.0%', '82.6%', '93.8%', '0.2141', '22'],
        ['CompGCN (128-dim)', 'MLP', '0.9838', '0.9775', '29.3%', '65.5%', '81.4%', '93.2%', '0.1168', '27.5'],
    ],
    col_widths=[2.5, 1.5, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3]
)

doc.add_paragraph(
    'Best overall: CompGCN + XGBoost (AUROC 0.9865), improving over RotatE + XGBoost by +0.0037. '
    'XGBoost decoder selected for both methods as it consistently outperforms Cosine and MLP decoders.'
)

doc.add_heading('7.3 Training Details', level=2)

bold_para('CompGCN: ', 'Pure PyTorch implementation, 200 epochs (best model at epoch 60), subtraction composition operator, 2 GNN layers, 128 hidden dimensions, 32M parameters. Binary cross-entropy loss with 1:1 negative sampling per epoch. DistMult-style scoring (h * r * t element-wise). Dropout=0.3, learning_rate=1e-3, Adam optimizer, gradient clipping max_norm=1.0. Early stopping patience=20. Trained on HPC GPU in ~7 minutes.')

bold_para('RotatE: ', 'PyKEEN framework. 128 complex dimensions (256 real after concatenating real and imaginary parts). 200 epochs, batch_size=4096, learning_rate=1e-4, 64 negative samples per positive. NSSALoss (Negative Sampling Self-Adversarial Loss, margin=9.0, adversarial_temperature=1.0). Trained on L40S GPU on HPC for ~3.4 hours. Native MRR=0.1890.')

bold_para('XGBoost decoder: ', 'n_estimators=300, max_depth=6, learning_rate=0.1, early_stopping_rounds=20 on validation set. Ensemble of gradient-boosted decision trees that learns the nonlinear relationship between embedding features and treatment probability.')

doc.add_heading('7.4 Predictions in Graph', level=2)

add_table(
    ['Source', 'Edges', 'Unique Drugs', 'Unique Diseases', 'Confidence Range', 'Avg Confidence'],
    [
        ['CompGCN_LinkPrediction', '6,607', '1,038', '37', '0.9887–0.9912', '0.9896'],
        ['RotatE_LinkPrediction', '7,828', '1,165', '142', '0.9932–0.9968', '0.9944'],
        ['Total', '14,435', '—', '—', '—', '—'],
    ],
    col_widths=[3.5, 1.5, 2, 2, 3.5, 2.5]
)

doc.add_paragraph(
    'Predictions are stored as predictedTreatsDisease edges with confidence and source properties. '
    'CompGCN predictions are shown in the Explore tab UI as orange dashed edges; both methods are '
    'fully queryable via the Cypher editor in the Query tab.'
)

doc.add_heading('7.5 Why CompGCN Is Primary', level=2)

bullet('Highest AUROC (0.9865 vs 0.9828 for RotatE)')
bullet('Relation-aware message passing: distinguishes all 27 relationship types during neighborhood aggregation, unlike RotatE which learns independent per-entity embeddings')
bullet('Dramatically faster training (~7 min vs ~3.4 hrs)')
bullet('Higher MRR (0.2141 vs 0.1890), indicating better average ranking of true treatments')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 8. KEY METRICS EXPLAINED
# ══════════════════════════════════════════════════════════
doc.add_heading('8. Key Metrics Explained', level=1)

doc.add_paragraph(
    'This section defines each evaluation metric used in the ML pipeline, written for someone who '
    'needs to understand and defend them in a paper or advisor discussion.'
)

bold_para('AUROC (Area Under ROC Curve): ',
    'The probability that a random positive pair (a drug that does treat a disease) scores higher '
    'than a random negative pair (a drug that does not treat a disease). A value of 0.5 means the '
    'model performs no better than random guessing; 1.0 means perfect discrimination. Our best '
    'AUROC of 0.9865 (CompGCN + XGBoost) means the model correctly ranks a true treatment pair '
    'above a random non-treatment pair 98.7% of the time. Limitation: the 1:1 random negatives '
    'are "easy" — most random Drug-Disease combos are clearly non-therapeutic, which inflates AUROC.')

bold_para('AUPRC (Area Under Precision-Recall Curve): ',
    'Measures the balance between precision (what fraction of predicted positives are actually '
    'positive) and recall (what fraction of actual positives are predicted) across all classification '
    'thresholds. More informative than AUROC for imbalanced datasets because it focuses on the '
    'positive class. A high AUPRC (0.9854) means the model has few false positives among its '
    'top-ranked predictions — when it predicts a treatment, it is usually correct. This is critical '
    'for drug repurposing where false positives waste expensive experimental validation resources.')

bold_para('Hits@K: ',
    'Fraction of test positives where the true disease appears in the top K ranked candidates for '
    'its drug. Hits@100 = 82.6% means that for about 4 out of 5 test drugs, the true disease '
    'appears in the top 100 candidates. This maps directly to a practical screening scenario: if a '
    'researcher examines the top 100 predictions per drug, they will find ~83% of actual treatment '
    'relationships. Hits@200 = 93.8% means examining the top 200 captures ~94%.')

bold_para('MRR (Mean Reciprocal Rank): ',
    'Average of 1/rank across all test positives. If the true disease is ranked 1st, the score is '
    '1.0; ranked 5th, the score is 0.2; ranked 100th, the score is 0.01. Higher is better. '
    'CompGCN\'s MRR of 0.2141 means the true disease is typically ranked around position ~5 '
    '(1/0.2141 ≈ 4.7). This metric rewards models that place the true answer near the very top '
    'of the ranked list.')

bold_para('Median Rank: ',
    'The median position of the true disease among all candidate diseases when ranking predictions '
    'for each test drug. CompGCN\'s median rank of 22 means half of the test drugs have their true '
    'disease ranked in the top 22 candidates. Lower is better. RotatE achieves a slightly better '
    'median rank of 17.5.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 9. METHODOLOGY NOTES
# ══════════════════════════════════════════════════════════
doc.add_heading('9. Methodology Notes', level=1)

doc.add_paragraph(
    'Three important methodology fixes were applied during the project, each significantly '
    'improving data quality and ML prediction reliability.'
)

doc.add_heading('9.1 ClinicalTrials.gov Inference Fix', level=2)

doc.add_paragraph(
    'The original ClinicalTrials.gov parser inferred drugTreatsDisease edges from any Phase 3/4 '
    'trial linking a drug intervention to a disease condition, producing 868 edges. An audit '
    'revealed many were spurious:'
)
bullet('Trials with primaryPurpose of "Prevention" or "Diagnostic" (not treatment)')
bullet('Drugs serving as comparators or placebos rather than experimental interventions')
bullet('Diseases matching secondary conditions rather than the primary condition under study')

doc.add_paragraph('Four filters were applied:')
bullet('primaryPurpose must be "TREATMENT"')
bullet('Drug must be in an EXPERIMENTAL arm (not comparator/placebo)')
bullet('Disease must match the first-listed condition (convention for primary condition)')
bullet('Edges carry a trialCount property counting qualifying trials')

bold_para('Result: ', '868 → 147 edges (83.1% reduction). The filtered edges represent higher-confidence therapeutic associations where the trial was specifically designed to test treatment efficacy.')

doc.add_heading('9.2 Duplicate Drug Node Merge (Entity Resolution)', level=2)

doc.add_paragraph(
    'Multiple sources loaded the same compound under different internal drugId values but shared '
    'the same xrefDrugBank canonical identifier. Entity resolution via scripts/merge_duplicate_drugs.py '
    'identified 5,611 duplicate groups, removed 6,055 nodes, transferred 474,641 edges, and '
    'deduplicated 9,094 redundant edges. Drug nodes went from 32,849 to 26,794. All ML models '
    'were retrained on the post-merge graph with zero cross-split leakage confirmed.'
)

doc.add_heading('9.3 Stale Memgraph-ID Bug Fix', level=2)

doc.add_paragraph(
    'After the drug merge deleted 6,055 nodes, Memgraph recycled their internal IDs. The original '
    'store_predictions.py matched Drug/Disease nodes by memgraph_id from a stale nodes.tsv export, '
    'causing predicted edges to land on wrong node types (Gene, PharmacologicClass, SideEffect) '
    'that had inherited the recycled IDs. This was detected when the web UI showed lab reagents '
    'and food chemicals (e.g., 2-Hydroxyestradiol, Tryptamine, DMSO, pyrachlostrobin) as predicted '
    'treatments for heart disease.'
)

bold_para('Fix: ', 'Rewrote store_predictions.py to match Drug/Disease nodes by name from the live graph instead of stale file-based IDs. All 14,435 predicted edges were verified as Drug → Disease with zero wrong-label contamination.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 10. KNOWN LIMITATIONS
# ══════════════════════════════════════════════════════════
doc.add_heading('10. Known Limitations', level=1)

doc.add_paragraph(
    'The following are known limitations of the current CardioKB build and ML predictions.'
)

bold_para('1. Confidence scores reflect ranking, not calibrated probability. ',
    'The XGBoost decoder outputs values 0–1 representing how closely a drug-disease pair\'s '
    'embedding geometry matches known treatment relationships. CompGCN predictions span '
    '0.989–0.991, RotatE spans 0.993–0.997 — narrow bands indicating uniform model '
    'confidence. These are NOT clinical efficacy estimates. A score of 0.99 does not mean a 99% '
    'chance the drug works; it means the graph structure strongly suggests a therapeutic '
    'relationship exists.')

bold_para('2. Pharmacologic-class parent nodes appear in predictions. ',
    'A subset of top predictions involve PharmacologicClass parent nodes (e.g., "ACE Inhibitors", '
    '"HMG-CoA Reductase Inhibitors") rather than specific drug compounds. These class-level '
    'predictions are structurally valid but less actionable for drug repurposing screening.')

bold_para('3. Hub-driven prediction concentration. ',
    'CompGCN predictions concentrate on high-degree Disease hub nodes: the top 6 diseases (heart '
    'disease, coronary artery disease, hypertension, atherosclerosis, congestive heart failure, '
    'myocardial infarction) account for a disproportionate share. Rarer CVD conditions receive '
    'fewer or no predictions.')

bold_para('4. Limited graph connectivity drugs. ',
    'Some predicted drugs have connectivity limited to a compoundInPharmacologicClass edge, with '
    'predictions driven by the class node\'s embedding rather than the drug\'s own pharmacological '
    'profile.')

bold_para('5. Disease-specific embedding dominance (CompGCN). ',
    'Pulmonary embolism has 6,859 incoming edges but received only 1 CompGCN prediction because '
    '92.9% of its edges are geneAssociatesWithDisease (6,372) versus only 46 drugTreatsDisease. '
    'CompGCN\'s neighborhood-aggregation mechanism causes the embedding to be dominated by '
    'gene-association signal rather than drug-treatment signal. RotatE produced 75 predictions for '
    'pulmonary embolism. Dabigatran case study: the active metabolite was predicted for PE despite '
    'only the prodrug (Dabigatran etexilate) having a curated indication.')

bold_para('6. Two legacy sources retained as-is. ',
    'SIDER (2015 GitHub commit) and LINCS L1000 (2020 GitHub commit) are static datasets with no '
    'live API alternatives available. Their data cannot be updated without institutional access '
    'to successor platforms.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 11. WEB INTERFACE FEATURES
# ══════════════════════════════════════════════════════════
doc.add_heading('11. Web Interface Features', level=1)

doc.add_paragraph(
    'CardioKB provides a single-page web application (interface/index.html) served by the Flask '
    'backend on port 5050. The interface uses vis.js for network visualization, Chart.js for '
    'analytics charts, and a custom dark/light theme system.'
)

doc.add_heading('Explore Tab', level=2)
bullet('Search any gene, drug, or disease with real-time autocomplete (triggers at 2+ characters)')
bullet('Interactive force-directed graph visualization with vis.js (ForceAtlas2Based physics)')
bullet('Specificity-ranked nodes: specificityScore = 1.0 / count(Disease neighbors)')
bullet('Max nodes limit (default 200, max 1000) to control visualization density')
bullet('Node type and edge type filter chips for focused exploration')
bullet('Node detail panel: click any node to see all properties and paginated connections')
bullet('ML predictions displayed as orange dashed edges with "Show ML Predictions" toggle')
bullet('Export: CSV, JSON, PNG, PDF')

doc.add_heading('Query Tab', level=2)
bullet('Natural language to Cypher (NL2Cypher powered by Claude via Anthropic API)')
bullet('Schema-aware system prompt with CardioKB-specific instructions and examples')
bullet('10 pre-built query templates (Disease Subgraph, Gene Neighbors, Drug Targets, etc.)')
bullet('Direct Cypher editor with read-only safety (blocks CREATE, DELETE, MERGE, SET)')
bullet('Multi-panel results (Neo4j Browser style): each query appends a new collapsible panel with Table and Graph view tabs')

doc.add_heading('Extract Disease Subgraph', level=2)
bullet('N-hop subgraph extraction (1–3 hops) for any disease')
bullet('Bulk export as JSON or CSV')
bullet('1 hop = direct connections, 2 hops = shared pathways/comorbidity bridges, 3 hops = broad network for hypothesis generation')

doc.add_heading('Edge Provenance', level=2)
bullet('Click any edge to see the "Why is this edge here?" panel')
bullet('Shows: relationship type, source/target nodes, source database (green badge), all evidence properties')
bullet('For ML predictions: confidence percentage, method name, red "Not clinically validated" warning')

doc.add_heading('Admin Panel', level=2)
bullet('System health check with SSE-streamed results')
bullet('Parser status table with Working/Failed/Skipped badges')
bullet('Node and edge count charts (Chart.js)')
bullet('ID mapping report showing match rates between TSV files and graph nodes')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 12. DEPLOYMENT
# ══════════════════════════════════════════════════════════
doc.add_heading('12. Deployment', level=1)

doc.add_paragraph(
    'CardioKB is deployed using Docker Compose with two services: Memgraph (graph database) and '
    'the Flask web application.'
)

add_table(
    ['Component', 'Details'],
    [
        ['Container orchestration', 'Docker Compose (docker-compose.yml)'],
        ['Database', 'Memgraph (in-memory, Bolt protocol, ports 7687/7444)'],
        ['Web app', 'Flask on Python 3.11-slim (port 5050)'],
        ['Graph data export', '304 MB tar.gz (~14 GB uncompressed in RAM)'],
        ['RAM requirement', '~16 GB for Memgraph to hold the full graph'],
        ['Import method', 'Binary volume restore via scripts/import_graph.sh (seconds, vs minutes for MERGE replay)'],
        ['Export method', 'scripts/export_graph.sh (stops Memgraph, tars volume, restarts)'],
        ['Data persistence', 'Named Docker volume (memgraph-data), survives container restarts'],
    ],
    col_widths=[4, 12]
)

doc.add_heading('Environment Variables', level=2)
add_table(
    ['Variable', 'Required?', 'Purpose'],
    [
        ['MEMGRAPH_PASSWORD', 'Yes', 'Graph database authentication'],
        ['ADMIN_PASSWORD', 'Yes', 'Admin features in web UI'],
        ['ANTHROPIC_API_KEY', 'Optional', 'NL2Cypher in Query tab (UI works fully without it)'],
        ['ANTHROPIC_FOUNDRY_API_KEY + BASE_URL', 'Optional', 'Azure AI Foundry (takes priority when both set)'],
        ['DRUGBANK_USERNAME / PASSWORD', 'No', 'Pipeline only (not needed for deployment)'],
    ],
    col_widths=[5, 2, 9]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 13. drugTreatsDisease SOURCES BREAKDOWN
# ══════════════════════════════════════════════════════════
doc.add_heading('13. drugTreatsDisease Sources Breakdown', level=1)

doc.add_paragraph(
    'The drugTreatsDisease relationship is the most important edge type for the ML drug repurposing '
    'pipeline. It aggregates evidence from four complementary sources, each providing a different '
    'type of evidence for treatment relationships.'
)

add_table(
    ['Source', 'Edges', 'Evidence Type'],
    [
        ['CTD', '3,099', 'Curated chemical-disease therapeutic relationships from scientific literature'],
        ['DrugBank_Indications', '1,449', 'Text-mined from DrugBank XML indication free-text fields (whole-word disease name matching)'],
        ['DrugCentral', '157', 'FDA-approved indications mapped via CUI-to-DOID'],
        ['ClinicalTrials.gov', '147', 'Phase 3/4 trials filtered by 4 criteria (primaryPurpose, EXPERIMENTAL arm, first-listed condition, trialCount)'],
        ['Total (deduplicated)', '4,852', 'Combined, deduplicated by (Drug, Disease) pair after entity resolution merge'],
    ],
    col_widths=[3, 1.5, 11.5]
)

doc.add_paragraph(
    'Training split: 4,726 train / 485 validation / 484 test (80/10/10 stratified).'
)

doc.add_heading('drugTreatsPhenotype', level=2)

doc.add_paragraph(
    'Many clinical conditions (e.g., tachycardia, arrhythmia, edema) exist only as Phenotype nodes '
    '(HPO), not Disease nodes (Disease Ontology). Since drugTreatsDisease connects Drug → Disease, '
    'these conditions had zero treatment edges. The same DrugBank indication text-mining approach '
    'was applied against Phenotype node names, creating 5,714 drugTreatsPhenotype edges '
    '(Drug → Phenotype, source: DrugBank_Indications).'
)

bullet('A blocklist filters out HPO modifier terms (Acute, Chronic, Severe, etc.) to prevent false positive matches')
bullet('NL2Cypher uses UNION ALL across both drugTreatsDisease and drugTreatsPhenotype for treatment queries')
bullet('ML pipeline trains on drugTreatsDisease only (Drug → Disease). drugTreatsPhenotype is a separate relationship type and does not affect ML training data or predictions')

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'CardioKB_Study_Guide_July15.docx')
doc.save(output_path)
print(f'Saved to {output_path}')
