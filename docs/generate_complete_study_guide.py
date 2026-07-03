"""Generate CardioKB Complete Study Guide as a formatted Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1E, 0x40, 0x7C)

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(13)


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p


def bold_para(bold_text, normal_text):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    p.add_run(normal_text)
    return p


# ══════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('CardioKB Complete Study Guide')
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0x1E, 0x40, 0x7C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Comprehensive Reference for Advisor Discussion and Paper Defense')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = author.add_run('Asma Nawaz')
r.font.size = Pt(13)
r.bold = True

affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = affil.add_run('Ph.D. Project: 2026\nMoore Lab, Cedars-Sinai Medical Center')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ('1. What is CardioKB and Why It Exists', 0),
    ('2. Data Sources (All 24)', 0),
    ('3. Graph Structure', 0),
    ('4. BaseAgent Pipeline — How the Graph Was Built', 0),
    ('5. Every Technical Challenge Solved', 0),
    ('6. ML Pipeline — Drug Repurposing Link Prediction', 0),
    ('6.1 Why Link Prediction', 1),
    ('6.2 Data Preparation', 1),
    ('6.3 Node2Vec (Dropped from Paper)', 1),
    ('6.4 RotatE', 1),
    ('6.5 CompGCN', 1),
    ('6.6 Decoders Compared', 1),
    ('6.7 Evaluation Metrics', 1),
    ('6.8 Feature Importance', 1),
    ('6.9 Results and Clinical Validation', 1),
    ('6.10 Future ML Directions', 1),
    ('7. User Interface — Every Feature Explained', 0),
    ('7.1 Explore Tab', 1),
    ('7.2 Query Tab', 1),
    ('7.3 Edge Provenance Panel', 1),
    ('7.4 ML Predicted Edges', 1),
    ('7.5 Sidebar Tools', 1),
    ('7.6 Admin Panel', 1),
    ('7.7 Visual Design', 1),
    ('8. Deployment', 0),
    ('9. Evaluation and Validation', 0),
    ('10. Future Directions', 0),
]
for item, indent in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(1.27)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. WHAT IS CARDIOKB AND WHY IT EXISTS
# ══════════════════════════════════════════════════════════
doc.add_heading('1. What is CardioKB and Why It Exists', level=1)

doc.add_heading('The Problem It Solves', level=2)
doc.add_paragraph(
    'Cardiovascular disease (CVD) is the leading cause of death globally, responsible for approximately '
    '17.9 million deaths per year. Researchers studying CVD need to understand complex relationships between '
    'genes, drugs, diseases, clinical trials, protein interactions, gene expression, and phenotypes. This '
    'information is scattered across dozens of separate databases, each with different formats, identifiers, '
    'and access methods.'
)
doc.add_paragraph(
    'A researcher investigating whether a diabetes drug could be repurposed for heart failure would need to '
    'manually query DrugBank for drug targets, STRING for protein interactions, ClinicalTrials.gov for ongoing '
    'trials, CTD for chemical-gene expression effects, and many more. This process is slow, error-prone, and '
    'makes it nearly impossible to discover non-obvious multi-hop connections across data sources.'
)

doc.add_heading('What CardioKB Does', level=2)
doc.add_paragraph(
    'CardioKB is a CVD-focused biomedical knowledge graph that integrates 22 deduplicated data sources '
    '(24 total parsers, 2 additional) into a single queryable graph database. The current graph contains:'
)
bullet('459,092 nodes across 17 distinct node types')
bullet('5,456,579 relationships across 28 relationship types')
bullet('22 data sources + 3 ML prediction sources (Node2Vec, RotatE, CompGCN)')
bullet('Every relationship carries a source property identifying its originating database')
doc.add_paragraph('It provides four core capabilities:')
bullet('A unified graph where a single Cypher query can traverse from a drug to its gene targets to associated diseases to clinical trials testing those drugs')
bullet('An automated ETL pipeline that can rebuild the entire graph from source data in approximately 5 minutes using BaseAgent multi-agent orchestration')
bullet('A web interface for interactive exploration, AI-powered natural language querying, and disease subgraph extraction with JSON/CSV export')
bullet('An ML-based drug repurposing pipeline that uses knowledge graph embeddings (RotatE, CompGCN) to predict new drug-disease treatment relationships with AUROC up to 0.9717')

doc.add_heading('Why Cardiovascular Disease Specifically', level=2)
bullet('Highest global mortality of any disease category (17.9 million deaths/year)')
bullet('Rich existing data ecosystem: extensive clinical trial data, well-characterized gene-disease associations, established drug targets')
bullet('Strong drug repurposing potential: many CVD drugs have complex polypharmacology, and the graph structure can reveal non-obvious therapeutic connections')
doc.add_paragraph(
    'The knowledge graph architecture is disease-agnostic. Alternative disease filter files exist for '
    "Alzheimer's (35 terms), cancer (70 terms), asthma (48 terms), and diabetes (52 terms). Switching "
    'the disease focus requires only changing a symlink (ontology/disease_filter.txt). The current CVD filter '
    'contains 184 cardiovascular disease terms.'
)

doc.add_heading('Target Users', level=2)
bullet('Biomedical researchers investigating CVD mechanisms and multi-target drug interactions')
bullet('Computational biologists performing network pharmacology and drug repurposing studies')
bullet('Clinician-scientists exploring evidence-based treatment relationships and clinical trial landscape')
bullet('Bioinformaticians needing a pre-integrated CVD knowledge base for downstream analysis')

doc.add_heading('How It Differs from Existing Resources', level=2)
doc.add_paragraph(
    'Several biomedical knowledge graphs exist (Hetionet, DRKG, PrimeKG, PharmKG), but CardioKB differs in key ways:'
)
add_table(
    ['Differentiator', 'CardioKB', 'Existing KGs'],
    [
        ['Disease focus', '184-term CVD filter ensures all disease-associated edges are relevant to cardiovascular research, reducing noise', 'General-purpose (all diseases), no filtering'],
        ['Automated rebuild', 'BaseAgent multi-agent orchestration rebuilds full graph in ~5 min', 'Manual builds take weeks/months'],
        ['Integrated ML', '1,500 predicted drug-disease edges stored in graph with confidence scores', 'ML predictions separate from graph, not queryable inline'],
        ['Provenance tracking', 'Every edge carries source property (e.g., source: "DrugBank")', 'Limited or no per-edge provenance'],
        ['Deduplication', '12 redundant sources removed after systematic audit', 'Often include overlapping sources without audit'],
    ],
    col_widths=[3, 6.5, 6.5]
)

doc.add_heading('The Core Value Proposition', level=2)
doc.add_paragraph(
    'CardioKB\'s value is in the integration. No single database can answer questions like "Which drugs '
    'that target genes associated with atrial fibrillation are currently in Phase 3 clinical trials for '
    'heart failure?" That query requires traversing DrugBank (drug-gene), PubTator/OpenTargets (gene-disease), '
    'and ClinicalTrials.gov (trial-disease) data simultaneously. CardioKB makes this a single Cypher query. '
    'The ML predictions extend this further by suggesting entirely novel drug-disease treatment pairs that '
    'no single database contains.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 2. DATA SOURCES
# ══════════════════════════════════════════════════════════
doc.add_heading('2. Data Sources (All 24)', level=1)
doc.add_paragraph(
    'CardioKB integrates 24 data source parsers (22 core + 2 additional), each contributing unique node '
    'types and/or relationship types. Sources were selected based on three criteria: (1) authoritative '
    'coverage of a specific biological entity type, (2) no redundancy with other included sources, and '
    '(3) public accessibility. During a systematic deduplication audit, 12 sources were removed '
    '(DisGeNET, GWAS Catalog, Jensen DISEASES, OMIM, WikiPathways, AOP-DB, HGNC base, CellAge, GenAge, '
    'Hetionet precomputed, DrugAge, AnAge).'
)

doc.add_heading('Direct Parsers (5 sources)', level=2)

bold_para('1. ClinicalTrials.gov', '')
doc.add_paragraph(
    'Contribution: 85,677 ClinicalTrial nodes, 27,866 STUDIES_CONDITION edges, 17,492 TESTS_INTERVENTION edges. '
    'Why included: Only source for clinical trial data linking interventions to disease conditions. Queries '
    'ClinicalTrials.gov API v2 per CVD disease term (184 terms), caches JSON responses. '
    'Limitation: Drug name matching to DrugBank IDs is inexact (string matching). Some trials study non-drug '
    'interventions (devices, procedures) that do not map to Drug nodes. Phase 3/4 trials with drug interventions '
    'contribute to the drugTreatsDisease edge set (868 edges) as proxy treatment evidence.'
)

bold_para('2. ClinPGx (PharmGKB successor)', '')
doc.add_paragraph(
    'Contribution: 29 DrugLabel nodes, 1,091 VARIANT_IN edges, 503 drugLabelAnnotatesGene edges, '
    '345 drugLabelDescribesDrug edges, 224 AFFECTS_RESPONSE_TO edges, 19 AFFECTS_RESPONSE_TO_CLASS edges. '
    'Why included: Only source for pharmacogenomic drug label annotations linking genetic variants to drug response. '
    'Successor to PharmGKB with cleaner REST API. '
    'Limitation: Small dataset focused on FDA-labeled pharmacogenomic associations. Limited to drugs with '
    'regulatory pharmacogenomic guidance.'
)

bold_para('3. NCBI Gene', '')
doc.add_paragraph(
    'Contribution: 193,687 Gene nodes (node-only source, no edges). '
    'Why included: Authoritative source for human gene identifiers. Provides the Gene node backbone that other '
    'sources connect to via gene symbols or NCBI Gene IDs. '
    'Limitation: Loads entire human gene catalog, so ~78% of Gene nodes are orphans (no CVD-specific edges). '
    'This is by design: the full catalog ensures any gene referenced by other sources can be matched.'
)

bold_para('4. DoRothEA (via OmniPath API)', '')
doc.add_paragraph(
    'Contribution: 367 TranscriptionFactor nodes, 12,985 transcriptionFactorInteractsWithGene edges '
    '(with morScore and confidence properties). '
    'Why included: Only source for curated transcription factor-gene regulatory interactions. Provides mode '
    'of regulation (activating/repressing via morScore: -1 = represses, 0 = unknown, +1 = activates) and '
    'confidence levels (A = highest through E = lowest). '
    'Limitation: Focused on well-characterized TF-gene pairs; coverage of less-studied transcription factors '
    'is incomplete.'
)

bold_para('5. DrugBank', '')
doc.add_paragraph(
    'Contribution: 19,842 Drug nodes, 12,089 drugBindsGene edges. '
    'Why included: Gold-standard drug database with detailed pharmacological data. Provides primary Drug node '
    'identifiers and drug-target binding relationships. '
    'Limitation: XML file access requires registration. Drug-gene edges represent binding targets only, not '
    'all pharmacological mechanisms.'
)

doc.add_heading('Hetionet-Derived Component Parsers (17 sources)', level=2)
doc.add_paragraph(
    'These parsers were originally part of the Hetionet precomputed dataset but have been broken out into '
    'individual component parsers that each query their respective original data source directly.'
)

hetio_sources = [
    ['6', 'Disease Ontology (DOID)', '12,012 Disease nodes; 6,447 diseaseIsSubtypeOf edges', 'Authoritative disease taxonomy with hierarchical subtype relationships. Provides Disease node backbone and DOID identifiers used for cross-source mapping.', 'Not all diseases have UMLS CUI cross-references, creating mapping gaps.'],
    ['7', 'Gene Ontology (GO)', '50,350 BP + 26,935 MF + 25,794 CC edges across 3 ontology domains', 'Only source for functional gene annotations. Essential for understanding biological processes and molecular functions genes participate in.', 'Annotations vary in evidence quality (some electronically inferred).'],
    ['8', 'Uberon (Anatomy)', '14,937 BodyPart nodes (node-only)', 'Provides anatomy identifiers that Bgee and Jensen TISSUES use to link gene expression to anatomical locations.', 'Node-only: depends on Bgee for actual anatomy-gene expression edges.'],
    ['9', 'MeSH (Symptoms)', '966 Symptom nodes (node-only)', 'Provides standardized symptom vocabulary from NCBI Medical Subject Headings.', '~100% orphan rate. No active source provides symptom-disease edges.'],
    ['10', 'SIDER (Side Effects)', '148,518 compoundCausesSideEffect edges', 'Only source for drug side effect data. Essential for safety profiling in drug repurposing.', 'Legacy dataset (2015 GitHub commit). Static, no live API alternative.'],
    ['11', 'LINCS L1000', '150,535 gene expression edges (geneRegulates + up/downreg) with zScore', 'Only source for drug-induced gene expression changes. zScore quantifies expression effect size.', 'Legacy (2020 GitHub commit). clue.io requires institutional access.'],
    ['12', 'MEDLINE', '365 cooccurrence edges (244 anatomy + 117 symptom + 4 disease)', 'Literature cooccurrence from MEDLINE abstracts.', 'Very small dataset from pinned GitHub commit. Legacy source.'],
    ['13', 'DrugCentral', '16,403 pharmacologic class + 245 treats + 96 palliates edges', 'Only source for pharmacologic class assignments and FDA-approved indications. CUI-to-DOID mapped.', 'Not all DrugCentral entries have DrugBank cross-references.'],
    ['14', 'BindingDB', '12,250 chemicalBindsGene edges', 'Complements DrugBank with experimentally measured binding affinities from a different curation source.', 'No binding affinity values loaded as properties (only edge existence).'],
    ['15', 'PubTator Central', '744,427 geneAssociatesWithDisease + 4,320 diseaseAssociatesWithDisease edges (CVD AND-filter)', 'Largest source of literature-mined gene-disease associations.', 'Cooccurrence is not causation. MeSH-to-DOID mapping introduces potential mismatches.'],
    ['16', 'CTD', '4,572 unique Drug nodes; 116,451 chemIncreasesExp + 97,951 chemDecreasesExp edges; 2,757 drugTreatsDisease edges', 'Only source for directional chemical-gene expression effects (increases vs. decreases). Also contributes 73% of drugTreatsDisease edges.', 'Expression relationships curated from literature; may not reflect dose/tissue-specific effects.'],
    ['17', 'Bgee (Gene Expression Atlas)', '784,026 underexpresses + 1,872 overexpresses edges with expressionScore', 'Only source for tissue-specific gene expression patterns with quantitative scores.', 'Heavily skewed toward underexpression (784K vs 1.9K).'],
    ['18', 'Jensen TISSUES', '215,235 gene-tissue edges', 'Gene-tissue expression associations.', 'Uses BTO tissue ontology; mapping to Uberon required BTO-to-Uberon resolution.'],
    ['19', 'HPO', '19,389 Phenotype nodes; 162,994 geneAssociatesWithPhenotype edges', 'Only source for gene-phenotype associations. Connects genetics to observable clinical features.', 'Annotations primarily from rare/Mendelian diseases; may underrepresent complex CVD phenotypes.'],
    ['20', 'Reactome', '44,979 geneInPathway + 44,979 pathwayContainsGene edges', 'Authoritative curated biological pathway assignments.', 'Pathway boundaries are somewhat arbitrary; genes appear in many overlapping pathways.'],
    ['21', 'STRING', '121,170 geneInteractsWithGene edges (confidence > 700, with combinedScore)', 'Largest source for protein-protein interaction data. Confidence > 700 filter ensures high quality.', 'Combined scores integrate heterogeneous evidence; not all are physical binding.'],
    ['22', 'OpenTargets', '32,826 geneAssociatesWithDisease edges (CVD AND-filter, EFO-to-DOID mapped)', 'Curated gene-disease associations from genetic, somatic, literature, and drug evidence.', 'Uses EFO ontology requiring EFO-to-DOID mapping. Some EFO terms lack DOID equivalents.'],
]
for s in hetio_sources:
    bold_para(f'{s[0]}. {s[1]}', '')
    doc.add_paragraph(f'Contribution: {s[2]}')
    doc.add_paragraph(f'Why included: {s[3]}')
    doc.add_paragraph(f'Limitation: {s[4]}')

doc.add_heading('Additional Parsers (2 sources)', level=2)

bold_para('23. HGNC Gene Families', '')
doc.add_paragraph(
    'Contribution: 1,934 GeneFamily nodes, 5,123 geneInFamily + 5,123 familyContainsGene edges. '
    'Why included: Only source for gene family membership. Useful for identifying related genes sharing '
    'drug targets or disease associations. Limitation: Not all genes belong to defined families.'
)

bold_para('24. ClinVar', '')
doc.add_paragraph(
    'Contribution: 4,488,042 Variant nodes, 2,267,095 hasVariant + 2,267,095 variantInGene edges with '
    'clinicalSignificance property. '
    'Why included: Authoritative source for clinical variant interpretations. Provides the Variant node '
    'backbone and variant-gene-disease linkages. '
    'Limitation: Largest single source (4.5M nodes). Only 0.12% of ClinVar variant-disease associations '
    'map to CardioKB Disease nodes because UMLS CUI overlap between ClinVar and Disease Ontology is only '
    '212 CUIs. This is expected and documented.'
)

doc.add_heading('How Sources Interconnect', level=2)
doc.add_paragraph(
    'The 24 sources form an interconnected web through shared node types. The Gene node (from NCBI Gene) is '
    'the central hub: STRING connects genes to genes, DrugBank/BindingDB connects drugs to genes, PubTator/'
    'OpenTargets connects genes to diseases, HPO connects genes to phenotypes, Reactome connects genes to '
    'pathways, CTD/LINCS L1000 connects drugs to genes via expression, and ClinVar connects variants to genes. '
    'The Drug node (from DrugBank/CTD) is the second hub: ClinicalTrials connects drugs to trials and diseases, '
    'DrugCentral provides pharmacologic classes and treatment edges, SIDER provides side effects. '
    'The Disease node (from Disease Ontology) connects everything through disease associations. '
    'This hub structure means a single query can traverse from any entity type to any other.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 3. GRAPH STRUCTURE
# ══════════════════════════════════════════════════════════
doc.add_heading('3. Graph Structure', level=1)

doc.add_heading('Scale', level=2)
doc.add_paragraph(
    'The current CardioKB graph (as of July 2026) contains 459,092 nodes across 17 distinct node types '
    'and 5,456,579 relationships across 28 relationship types. There are 22 data sources providing edges '
    '(each tagged with a source property), plus 3 ML prediction sources and 1 text-mined source (DrugBank_Indications). '
    'Every single relationship in the graph carries a source property identifying which database or model produced it.'
)

doc.add_heading('Node Types (17) — What Each Represents Biologically', level=2)
add_table(
    ['Node Type', 'Count', 'Source', 'Biological Meaning'],
    [
        ['Gene', '193,795', 'NCBI Gene', 'A segment of DNA that encodes a protein or functional RNA. Identified by official HGNC symbol (e.g., APOE, BRCA1). The central hub of the graph — most other entities connect through genes.'],
        ['Variant', '135,555', 'ClinVar', 'A specific genetic mutation or polymorphism (e.g., rs429358). Carries clinicalSignificance: Pathogenic, Benign, Likely Pathogenic, or VUS (Variant of Uncertain Significance).'],
        ['Drug', '32,849', 'DrugBank + CTD', 'Any chemical compound with pharmacological activity. Includes FDA-approved drugs, experimental compounds, and nutraceuticals. Identified by DrugBank ID.'],
        ['BiologicalProcess', '24,428', 'Gene Ontology', 'A recognized biological activity (e.g., "apoptotic process", "inflammatory response"). From the GO Biological Process ontology.'],
        ['ClinicalTrial', '21,578', 'ClinicalTrials.gov', 'A registered clinical study. Properties include phase (1-4), status (Recruiting, Completed, etc.), sponsor, and trial ID (NCT number).'],
        ['Phenotype', '19,389', 'HPO', 'An observable clinical feature (e.g., "cardiomegaly", "hypertension"). From the Human Phenotype Ontology. Connects genetic findings to what clinicians observe.'],
        ['MolecularFunction', '10,056', 'Gene Ontology', 'What a gene product does at the molecular level (e.g., "protein kinase activity", "DNA binding"). From the GO Molecular Function ontology.'],
        ['GeneFamily', '4,257', 'HGNC', 'A group of genes sharing a common ancestor or function (e.g., ABC transporters, receptor tyrosine kinases). Useful for identifying related drug targets.'],
        ['CellularComponent', '4,076', 'Gene Ontology', 'Where a gene product is located in the cell (e.g., "mitochondria", "plasma membrane"). From the GO Cellular Component ontology.'],
        ['Disease', '3,442', 'Disease Ontology', 'A recognized disease entity with a DOID identifier (e.g., DOID:1287 = cardiovascular system disease). Hierarchical — "atrial fibrillation" is a subtype of "cardiac arrhythmia."'],
        ['Pathway', '2,870', 'Reactome', 'A series of molecular interactions leading to a biological outcome (e.g., "Signaling by VEGF", "Cholesterol biosynthesis"). Curated by Reactome.'],
        ['PharmacologicClass', '2,359', 'DrugCentral', 'An FDA-defined drug class (e.g., "HMG-CoA Reductase Inhibitor" for statins, "Beta-Adrenergic Blocker" for beta-blockers).'],
        ['SideEffect', '2,227', 'SIDER', 'An adverse drug reaction (e.g., "headache", "nausea"). Linked to drugs via compoundCausesSideEffect edges. Important for safety profiling in repurposing.'],
        ['BodyPart', '1,400', 'Uberon', 'An anatomical structure (e.g., "heart", "liver", "aorta"). Connected to genes via tissue-specific expression data from Bgee.'],
        ['Symptom', '415', 'MeSH', 'A clinical symptom from NCBI Medical Subject Headings. Currently ~100% orphans (no active source provides symptom-disease edges).'],
        ['TranscriptionFactor', '367', 'DoRothEA', 'A protein that binds DNA to regulate gene expression. Connected to target genes with morScore (+1 = activates, -1 = represses) and confidence (A-E).'],
        ['DrugLabel', '29', 'ClinPGx', 'An FDA pharmacogenomic drug label annotation specifying how genetic variants affect drug response (e.g., "CYP2D6 poor metabolizers should use reduced dose").'],
    ],
    col_widths=[2.8, 1.5, 2.2, 10]
)

doc.add_heading('Relationship Types (27) — What Each Means Biologically', level=2)
add_table(
    ['Relationship', 'Direction', 'Source', 'Count', 'Properties', 'Biological Meaning'],
    [
        ['bodyPartOverexpressesGene', 'BodyPart→Gene', 'Bgee', '2,749,193', 'expressionScore', 'This gene is expressed at higher/lower levels in this tissue compared to baseline'],
        ['geneAssociatesWithDisease', 'Gene→Disease', 'PubTator + OpenTargets', '539,964', 'score (OT)', 'Literature or curated evidence linking this gene to this disease'],
        ['chemicalIncreasesExpression', 'Drug→Gene', 'CTD', '343,823', '—', 'This chemical increases the mRNA/protein expression of this gene'],
        ['chemicalDecreasesExpression', 'Drug→Gene', 'CTD', '328,726', '—', 'This chemical decreases the mRNA/protein expression of this gene'],
        ['geneAssociatesWithPhenotype', 'Gene→Phenotype', 'HPO', '270,265', '—', 'Mutations in this gene are associated with this observable clinical feature'],
        ['geneInteractsWithGene', 'Gene→Gene', 'STRING', '229,007', 'combinedScore', 'These two proteins physically interact or functionally cooperate (confidence 0-1000, filtered >700)'],
        ['geneInPathway', 'Gene→Pathway', 'Reactome', '137,116', 'evidenceCode', 'This gene participates in this biological pathway'],
        ['variantInGene', 'Variant→Gene', 'ClinVar', '135,393', '—', 'This genetic variant is located within this gene'],
        ['geneParticipatesInBP', 'Gene→BP', 'Gene Ontology', '122,117', '—', 'This gene product is involved in this biological process'],
        ['geneAssociatedWithCC', 'Gene→CC', 'Gene Ontology', '90,141', '—', 'This gene product is located in this cellular compartment'],
        ['geneHasMolecularFunction', 'Gene→MF', 'Gene Ontology', '76,612', '—', 'This gene product performs this molecular function'],
        ['compoundUpregulatesGene', 'Drug→Gene', 'LINCS L1000', '74,854', 'zScore', 'This drug increases expression of this gene (measured in cell lines)'],
        ['compoundCausesSideEffect', 'Drug→SideEffect', 'SIDER', '67,721', '—', 'This drug is known to cause this adverse reaction'],
        ['compoundDownregulatesGene', 'Drug→Gene', 'LINCS L1000', '64,661', 'zScore', 'This drug decreases expression of this gene (measured in cell lines)'],
        ['variantAssocWithDisease', 'Variant→Disease', 'ClinVar', '51,323', 'clinicalSignificance', 'This variant is clinically associated with this disease (Pathogenic/Benign/VUS)'],
        ['drugBindsGene', 'Drug→Gene', 'DrugBank', '29,363', 'interactionType', 'This drug physically binds to this gene product (protein target)'],
        ['geneInFamily', 'Gene→GeneFamily', 'HGNC', '27,022', '—', 'This gene belongs to this gene family'],
        ['compoundInPharmClass', 'Drug→PharmClass', 'DrugCentral', '25,687', '—', 'This drug belongs to this FDA pharmacologic class'],
        ['chemicalBindsGene', 'Drug→Gene', 'BindingDB', '22,735', '—', 'Experimentally measured binding between this chemical and this protein'],
        ['STUDIES_CONDITION', 'Trial→Disease', 'ClinicalTrials.gov', '20,667', '—', 'This clinical trial studies this disease condition'],
        ['tfInteractsWithGene', 'TF→Gene', 'DoRothEA', '15,082', 'morScore, confidence', 'This transcription factor regulates this gene (+1=activates, -1=represses)'],
        ['drugTreatsPhenotype', 'Drug→Phenotype', 'DrugBank_Indications', '10,955', '—', 'This drug treats this phenotype/condition (text-mined from DrugBank indication fields)'],
        ['drugTreatsDisease', 'Drug→Disease', 'CTD+CT+DC+DBI', '6,272', '—', 'This drug is used to treat this disease (curated + clinical trial + FDA + text-mined evidence)'],
        ['TESTS_INTERVENTION', 'Trial→Drug', 'ClinicalTrials.gov', '3,180', '—', 'This clinical trial tests this drug as an intervention'],
        ['diseaseIsSubtypeOf', 'Disease→Disease', 'Disease Ontology', '2,581', '—', 'This disease is a subtype of this broader disease category'],
        ['predictedTreatsDisease', 'Drug→Disease', 'ML Predictions', '1,500', 'confidence', 'ML model predicts this drug might treat this disease (NOT clinically validated)'],
        ['AFFECTS_RESPONSE_TO', 'Gene→Drug', 'ClinPGx', '74', '—', 'Genetic variation in this gene affects response to this drug'],
    ],
    col_widths=[3, 2, 1.8, 1.3, 1.8, 5.5]
)

doc.add_heading('Why Memgraph Over Neo4j', level=2)
doc.add_paragraph('CardioKB initially used Neo4j but migrated to Memgraph for several reasons:')
bullet('In-memory architecture: Memgraph stores the entire graph in RAM, providing sub-second query response for the interactive web interface. The full graph occupies ~14 GB in memory.')
bullet('Cypher compatibility: Memgraph uses the same Cypher query language as Neo4j, so all existing queries, the loader, and the Flask API required only minor syntax adjustments (e.g., replacing Neo4j-specific WHERE NOT (n)--() patterns).')
bullet('Docker-friendly: Memgraph\'s container image is lightweight (~250 MB) and starts faster than Neo4j, simplifying the Docker Compose deployment.')
bullet('No license restrictions: Memgraph\'s open-source edition has no node/relationship count limits, whereas Neo4j Community Edition has constraints that would affect a 459K-node graph.')

doc.add_heading('Orphan Nodes — What They Are and Why They Are Expected', level=2)
doc.add_paragraph(
    'An orphan node is a node with zero edges (no connections to any other node). Several node types have '
    'high orphan rates. These are expected and deliberate:'
)
add_table(
    ['Node Type', 'Orphan Rate', 'Why Expected'],
    [
        ['Symptom', '~100%', 'Only MEDLINE provides symptom edges (skipped). MeSH provides 966 nodes but no relationship mappings.'],
        ['DrugLabel', '~100%', 'ClinPGx labels whose target genes/drugs are not in the CVD-filtered graph.'],
        ['Gene', '~78%', 'Full NCBI catalog (193K genes) loaded so any gene referenced by other sources can be matched. Only ~43K have CVD-relevant edges.'],
        ['Disease', '~76%', 'Full Disease Ontology loaded (3,442 diseases) for subtype hierarchy traversal. Only ~813 have CVD-filtered edges.'],
    ],
    col_widths=[3, 2.5, 10.5]
)
doc.add_paragraph(
    'High orphan rates are a deliberate design choice: loading complete node catalogs ensures no edges are '
    'lost due to missing target nodes. The alternative (loading only nodes that have edges) would require '
    'pre-computing the edge set before loading nodes, adding complexity and preventing the pipeline from '
    'being order-independent.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 4. BASEAGENT PIPELINE
# ══════════════════════════════════════════════════════════
doc.add_heading('4. BaseAgent Pipeline — How the Graph Was Built', level=1)

doc.add_heading('What BaseAgent Is', level=2)
doc.add_paragraph(
    'BaseAgent is a multi-agent orchestration framework built by Binglan Li (Li Lab, Stanford University) '
    'that automates the construction of biomedical knowledge graphs. It was originally developed for AlzKB '
    "(Alzheimer's Knowledge Base) and provides a template-based system where each data source is integrated "
    'through a standardized pipeline of AI agents. The CardioKB-specific adaptations are on the cardiokb '
    'branch of the BaseAgent repository (~/Desktop/BaseAgent/cardiokb.ipynb).'
)

doc.add_heading('The Multi-Agent System', level=2)
doc.add_paragraph('BaseAgent uses four specialized agents working in sequence for each data source:')
bold_para('1. Ontology Agent: ', 'Analyzes the target data source and maps its entities to the knowledge graph schema. It determines which columns in the source data correspond to which node types, relationship types, and properties in the graph. For CardioKB, this produces one of the 86 ontology configs in src/ontology_configs.py.')
bold_para('2. Engineer Agent: ', 'Generates a parser (Python class extending BaseParser) that downloads raw data from the source, processes it, and converts it into standardized TSV format. Each parser handles the specific API or file format of its source (REST API, FTP download, XML parsing, TSV parsing, etc.).')
bold_para('3. Mapping Agent: ', 'Resolves identifier conflicts between the new source and existing graph data. For example, PubTator uses MeSH IDs for diseases but CardioKB uses DOID identifiers. The mapping agent generates the ID translation logic (MeSH-to-DOID, CUI-to-DOID, EFO-to-DOID, etc.).')
bold_para('4. Evaluator Agent: ', 'Validates the parsed output against expected schemas, checks for data quality issues (null values, duplicate IDs, malformed identifiers), and reports yield statistics (how many nodes/edges were produced, what percentage matched existing graph entities).')

doc.add_heading('The ista Pipeline and owl2.DataLoader', level=2)
doc.add_paragraph(
    'BaseAgent uses the ista Python package (Intelligent Schema-based Transformation Architecture) to '
    'manage ontology-driven data integration. The original template used ista\'s FlatFileDatabaseParser '
    'to load OWL2 ontology files, but this was extremely slow for large ontologies (10+ minutes for the '
    'Disease Ontology\'s 42,000 terms). The fix was switching to ista\'s owl2.DataLoader.execute() method, '
    'which processes the same file in under 10 milliseconds — a 60,000x speedup. However, ista\'s '
    'save_ontology() function only serializes schema elements (classes, properties), not data individuals '
    '(actual nodes and edges). This meant the RDF serialization approach had to be bypassed entirely, '
    'replaced with a direct TSV-to-Memgraph loading approach using 86 Python ontology configs.'
)

doc.add_heading('How Populate → Export → Load Works', level=2)
doc.add_paragraph('The pipeline orchestrator (src/main.py) coordinates five phases:')
bold_para('Phase 1 — Download: ', 'Each parser downloads raw data to data/raw/<source>/. Sources use various access methods: REST APIs (ClinicalTrials.gov, ClinPGx, DoRothEA, OpenTargets), FTP (NCBI Gene, PubTator, Bgee, ClinVar), XML files (DrugBank), TSV files (most Hetionet-derived sources). Downloaded data is cached so subsequent runs can skip this step with --skip-download.')
bold_para('Phase 2 — Parse: ', 'Each parser extracts nodes and edges into pandas DataFrames. This involves cleaning column names, filtering to CVD-relevant entries (for sources with AND-filters like PubTator and OpenTargets), resolving identifiers, and deduplicating entries.')
bold_para('Phase 3 — TSV Export: ', 'DataFrames are written to data/processed/<source>/*.tsv. Each TSV file has standardized column headers matching the ontology config (e.g., source_id, target_id, relationship_type, properties).')
bold_para('Phase 4 — Graph Load: ', 'memgraph_loader.py reads each TSV file paired with its ontology config from ontology_configs.py. It uses UNWIND-based Cypher batching (batch_size=1000) with MERGE to prevent duplicate nodes/edges. Every relationship gets r.source set from the config\'s source_label field (e.g., source: "DrugBank").')
bold_para('Phase 5 — Post-processing: ', 'scripts/compute_specificity.py calculates specificityScore = 1.0 / count(Disease neighbors) for every node and stores it as a node property. Disease nodes get 0.0, nodes with no disease connections get 1.0. Timestamp stored in a _Metadata node.')

doc.add_heading('Why It Matters: Automated Rebuild', level=2)
doc.add_paragraph(
    'Before BaseAgent, building a knowledge graph of this scale required months of manual work: writing '
    'individual parsers, debugging ID mismatches, validating output, and loading data. With BaseAgent, the '
    'entire CardioKB graph (459K nodes, 5.4M edges, 24 sources) can be rebuilt from scratch in approximately '
    '5 minutes. This makes the graph reproducible and updatable: when a source database releases new data, '
    're-running the pipeline incorporates the updates automatically.'
)

doc.add_heading('What SLURM Is and Why We Used HPC', level=2)
doc.add_paragraph(
    'SLURM (Simple Linux Utility for Resource Management) is a job scheduler used on High-Performance '
    'Computing (HPC) clusters. We used HPC for the ML embedding training because:'
)
bullet('CompGCN training requires 32 million parameters and ~46 minutes on a GPU. This would take hours on a CPU.')
bullet('RotatE training via PyKEEN requires ~10.3 hours on an L40S GPU with 64 negative samples per positive edge.')
bullet('Node2Vec needed SparseOTF mode on HPC to avoid out-of-memory errors on a 459K-node graph.')
doc.add_paragraph(
    'SLURM job scripts (hpc/*.slurm) request specific resources (GPU type, memory, time limit) and queue '
    'the training job on the cluster. Results (embeddings, models, evaluation reports) are saved to ml/data/ '
    'and transferred back to the local machine.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 5. KEY TECHNICAL CHALLENGES
# ══════════════════════════════════════════════════════════
doc.add_heading('5. Every Technical Challenge Solved', level=1)

doc.add_heading('5.1 FlatFileDatabaseParser → owl2.DataLoader (60,000x Speedup)', level=2)
bold_para('What it was: ', 'The original BaseAgent template used ista\'s FlatFileDatabaseParser to load OWL2 ontology files (like the Disease Ontology .owl file containing 42,000 disease terms). This parser reads every RDF triple sequentially, building an in-memory graph one statement at a time.')
bold_para('Why it broke: ', 'For the Disease Ontology file, this took over 10 minutes. The pipeline loads multiple ontology files (Disease Ontology, Gene Ontology, Uberon, HPO, MeSH), so the cumulative time made the pipeline take hours instead of minutes.')
bold_para('How fixed: ', 'Switched to ista\'s owl2.DataLoader.execute() method, which uses optimized batch RDF parsing. The same Disease Ontology file processed in under 10 milliseconds — a 60,000x speedup.')
bold_para('Why it matters: ', 'This single fix reduced the entire pipeline from multi-hour to ~5 minutes, making iterative development and automated rebuilds practical.')

doc.add_heading('5.2 STUDIES_CONDITION Synonym Resolution (1 → 20,667 edges)', level=2)
bold_para('What it was: ', 'The ClinicalTrials.gov parser initially produced only 1 STUDIES_CONDITION edge because it was doing exact string matching between trial condition names and Disease Ontology disease names.')
bold_para('Why it broke: ', 'Clinical trials use clinical terminology (e.g., "Heart Attack") while Disease Ontology uses formal names (e.g., "myocardial infarction"). These are synonyms for the same disease but don\'t match as strings.')
bold_para('How fixed: ', 'Implemented cascading synonym resolution using Disease Ontology\'s synonym fields. The parser now builds a lookup table mapping all synonyms (exact synonyms, related synonyms, narrow synonyms) to their canonical DOID. When matching trial conditions, it first tries exact disease name match, then synonym match, then partial string match.')
bold_para('Result: ', 'STUDIES_CONDITION edges increased from 1 to 20,667 (later growing to 27,866 with additional disease terms).')

doc.add_heading('5.3 CTD Expression Edge Fix (0 → 525,907 edges)', level=2)
bold_para('What it was: ', 'The initial CTD parser produced 0 chemicalIncreasesExpression and 0 chemicalDecreasesExpression edges despite CTD containing hundreds of thousands of chemical-gene expression relationships.')
bold_para('Why it broke: ', 'A column mapping bug in the parser was reading the wrong column indices from CTD\'s TSV format. Gene IDs were being read from the organism column and vice versa, so every lookup failed silently.')
bold_para('What chemId→commonName matching means: ', 'CTD identifies chemicals by CAS Registry Numbers and MeSH IDs, but CardioKB Drug nodes use DrugBank IDs and common names. The fix also implemented chemId-to-commonName resolution so CTD chemicals could be matched to existing Drug nodes.')
bold_para('Result: ', 'This yielded 525,907 expression edges (116,451 increases + 97,951 decreases in the final CVD-filtered build). This was the single largest edge count fix in the project.')

doc.add_heading('5.4 MemgraphExporter Streaming Fix (Silent Triple Loss)', level=2)
bold_para('What silent triple loss is: ', 'The BaseAgent TSVMemgraphExporter was silently dropping rows during export when source DataFrames contained null values or mixed-type columns (e.g., a column with both integers and strings). The exporter would skip entire rows without logging a warning if any cell failed type validation.')
bold_para('Why it happens on large RDF files: ', 'When parsing large datasets with hundreds of thousands of rows, some source files have inconsistent data types (a column that\'s mostly integers but has a few string values like "N/A"). Pandas loads these as object type, which the exporter couldn\'t handle.')
bold_para('How fixed: ', 'Added explicit null handling (replace NaN with empty string before export), mixed-type column coercion (force all values to string), and logging for every skipped row so drops are visible. This recovered thousands of previously lost edges across multiple sources.')

doc.add_heading('5.5 Mixed Endpoint Edge Types Fix', level=2)
bold_para('What this means: ', 'Several ClinicalTrials.gov records had endpoint types (primary, secondary) mixed into the edge type field, creating invalid relationship types like "STUDIES_CONDITION_primary" instead of the standard "STUDIES_CONDITION".')
bold_para('How fixed: ', 'Normalized all endpoint annotations into edge properties while keeping the relationship type clean. The endpoint type is preserved as metadata but doesn\'t pollute the relationship type namespace.')

doc.add_heading('5.6 Disease CUI Mapping Gap', level=2)
bold_para('What UMLS CUI is: ', 'The Unified Medical Language System (UMLS) assigns Concept Unique Identifiers (CUIs) to medical concepts. For example, CUI C0018801 = "Heart Failure". Many databases use CUIs as their disease identifier.')
bold_para('Why it matters: ', 'DrugCentral uses UMLS CUI identifiers for diseases, but CardioKB\'s Disease nodes use Disease Ontology (DOID) identifiers. Without CUI-to-DOID mapping, no DrugCentral edges could connect to Disease nodes.')
bold_para('How fixed: ', 'Implemented CUI-to-DOID mapping using Disease Ontology\'s xref (cross-reference) fields. Each DOID entry lists equivalent CUIs. Built a lookup table mapping CUIs to DOIDs. This increased DrugCentral drugTreatsDisease edges from near-zero to 245 (157 after deduplication with CTD). The mapping gap remains a known limitation: only 212 CUIs have DOID equivalents.')

doc.add_heading('5.7 drugTreatsDisease Edge Aggregation', level=2)
doc.add_paragraph(
    'The drugTreatsDisease relationship is the most important edge type for the ML drug repurposing pipeline. '
    'It comes from four complementary sources, each providing a different type of evidence:'
)
add_table(
    ['Source', 'Edges', 'Evidence Type', 'Why This Source'],
    [
        ['CTD', '2,757', 'Curated chemical-disease therapeutic relationships from scientific literature', 'Literature-curated, high confidence.'],
        ['DrugBank_Indications', '2,930', 'Text-mined from DrugBank XML indication free-text fields', 'Whole-word disease name matching against existing Disease nodes. Significantly increased coverage.'],
        ['ClinicalTrials.gov', '868', 'Extracted from Phase 3/4 trial intervention-condition pairs', 'Phase 3/4 trials are proxy evidence that a drug treats a disease.'],
        ['DrugCentral', '157', 'FDA-approved indications mapped via CUI-to-DOID', 'Regulatory evidence — FDA-approved indications.'],
        ['Total (deduplicated)', '6,272', 'Combined, deduplicated by (Drug, Disease) pair', 'The original 3,782 curated edges were used as ML training labels. The 2,930 DrugBank_Indications edges were added post-training.'],
    ],
    col_widths=[3, 1.5, 6, 5.5]
)

doc.add_heading('5.7.1 drugTreatsPhenotype — Phenotype Coverage Gap Fix', level=2)
doc.add_paragraph(
    'Many clinical conditions (e.g., tachycardia, edema, arrhythmia) exist only as Phenotype nodes (HPO), '
    'not Disease nodes (Disease Ontology). Since drugTreatsDisease connects Drug→Disease, these conditions '
    'had zero treatment edges. To fix this, the same DrugBank indication text-mining approach was applied '
    'against Phenotype node names, creating 10,955 drugTreatsPhenotype edges (Drug→Phenotype) from 2,646 drugs '
    'to 692 phenotypes. A blocklist filters out HPO modifier terms (Acute, Chronic, Severe, etc.) that would '
    'produce false positive matches. The NL2Cypher instructions use UNION ALL across both relationship types '
    'for treatment queries. Script: scripts/drugbank_indications.py.'
)

doc.add_heading('5.8 Named Docker Volume Persistence Issue', level=2)
bold_para('What it is: ', 'Docker containers can store data in volumes. An anonymous volume is created fresh each time a container starts and is deleted when the container is removed. A named volume persists independently of the container lifecycle.')
bold_para('Why it caused problems: ', 'Early in deployment, graph data was being lost on docker compose down because the Memgraph container was using an anonymous volume instead of a named volume. Every restart meant re-importing the entire 14 GB graph.')
bold_para('How fixed: ', 'Added an explicit named volume (memgraph-data) in docker-compose.yml mapped to /var/lib/memgraph inside the container. Also updated import_graph.sh to wait for Memgraph readiness (polling bolt port with mgconsole) instead of using a fixed sleep timer, preventing race conditions on slower machines.')

doc.add_heading('5.9 Node2Vec OOM Issues', level=2)
bold_para('What happened: ', 'Node2Vec requires building a transition probability matrix for random walks. On a 459K-node graph with 5.4M edges, the default approach exhausted available RAM.')
bold_para('SparseOTF vs PreComp vs DenseOTF: ', 'PecanPy (the Node2Vec implementation we used) offers three modes. PreComp pre-computes all transition probabilities — fast walks but requires storing the full matrix in memory (impossible for our graph). DenseOTF computes probabilities on-the-fly using dense matrices — lower memory but still too much for 459K nodes. SparseOTF computes on-the-fly using sparse matrices — the only mode that fits in memory for our graph size. We used SparseOTF mode, which completed successfully on HPC.')

doc.add_heading('5.10 Binglan\'s Edge Properties Cherry-Pick', level=2)
bold_para('What it fixed: ', 'The original BaseAgent TSV exporter did not propagate quantitative edge properties (combinedScore, morScore, expressionScore, etc.) from parsed DataFrames to the exported TSV files. Binglan Li contributed a fix on the cardiokb branch that cherry-picked these property columns through the export pipeline, ensuring that properties like STRING\'s combinedScore (0-1000), DoRothEA\'s morScore (-1/0/+1), and Bgee\'s expressionScore were preserved in the final graph. Without this fix, all 7 quantitative edge property types would have been lost.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 6. ML PIPELINE
# ══════════════════════════════════════════════════════════
doc.add_heading('6. ML Pipeline — Drug Repurposing Link Prediction', level=1)

doc.add_heading('6.1 Why Link Prediction', level=2)
bold_para('What link prediction is: ', 'Given a graph with known edges, link prediction asks: "Which edges are missing?" More precisely, given two nodes that are not currently connected by a specific relationship type, how likely is it that this relationship should exist? The model learns patterns from existing edges to score potential new ones.')
bold_para('Why drug repurposing specifically: ', 'Drug repurposing (finding new therapeutic uses for existing drugs) is faster and cheaper than de novo drug development because repurposed drugs have already passed safety testing. Bringing a repurposed drug to market costs ~$300M and takes ~6.5 years, compared to ~$2.6B and ~13 years for a new drug.')
bold_para('What the task is: ', 'Given the CardioKB graph with 3,782 known drugTreatsDisease edges (at time of training — now 6,272 after DrugBank_Indications enrichment), predict which of the remaining ~4.4 million possible Drug-Disease pairs are most likely to be genuine treatment relationships. The top predictions become hypotheses for experimental validation. Note: the ML models were trained on the original 3,782 curated edges; the 2,930 DrugBank_Indications edges were added post-training and do not affect ML results.')

doc.add_heading('6.2 Data Preparation', level=2)

bold_para('drugTreatsDisease breakdown: ', 'CTD provides 2,757 edges (curated from literature), ClinicalTrials.gov provides 868 edges (from Phase 3/4 trials where the drug is being tested as a treatment — this counts as proxy evidence because Phase 3/4 means the drug has already shown efficacy signals), and DrugCentral provides 157 edges (FDA-approved indications). After deduplication by (Drug, Disease) pair: 3,782 total positive edges.')

bold_para('Why Phase 3/4 counts as proxy treatment evidence: ', 'A drug in a Phase 3 or 4 clinical trial for a disease has already passed safety testing (Phase 1) and shown preliminary efficacy (Phase 2). While not yet FDA-approved, the existence of a large-scale efficacy trial is strong evidence of a therapeutic relationship. We only include trials where the drug is the intervention and the disease is the condition, not observational studies.')

bold_para('80/10/10 stratified split: ', '"Stratified" means all 25 edge types are split proportionally. If drugTreatsDisease is 0.07% of all edges, then ~0.07% of the train, validation, and test sets are drugTreatsDisease edges. This ensures the train set reflects the full graph topology, not just a biased sample. Split seed = 42 for reproducibility. Result: 4,349,555 train / 543,683 validation / 543,683 test edges. For drugTreatsDisease specifically: 3,026 train / ~310 validation / ~325 test.')

bold_para('Negative sampling: ', 'For each positive drugTreatsDisease edge (a drug that does treat a disease), we sample one negative: a random (Drug, Disease) pair that does not exist anywhere in the full graph (not just the train set). The 1:1 ratio keeps classes balanced. We sample against the FULL graph (not just train) to prevent a val/test positive from appearing as a train negative, which would be a false negative. Seeds: train=42, val=123, test=456, identical across all three methods for fair comparison.')

bold_para('Why data leakage matters and how we prevented it: ', 'Data leakage means the model sees test data during training, inflating performance. We prevent it four ways: (1) all embedding methods train ONLY on the train split — val/test edges are removed from the graph before training; (2) negative sampling checks the full graph, not just train; (3) XGBoost decoder early-stops on validation, never touching test; (4) test set is used only for final evaluation numbers.')

bold_para('Therapeutic filter: ', 'Not all 32,849 Drug nodes are therapeutically relevant. Many are metabolites, food compounds, or experimental chemicals with no treatment signal. We filter to drugs that have at least one "therapeutic signal" edge (drugTreatsDisease, drugBindsGene, chemicalIncreasesExpression, etc.). This reduces 32,849 to 9,735 therapeutic drugs. Without this filter, the prediction space would be dominated by irrelevant chemicals.')

doc.add_heading('6.3 Node2Vec (Dropped from Paper)', level=2)
bold_para('How random walks work: ', 'Node2Vec generates random paths through the graph, like a person randomly wandering through a city following streets. Starting from each node, it takes a series of steps (walk_length=80), at each step choosing a neighbor to visit. The choice is biased by two parameters: p (return parameter, set to 1.0) controls the likelihood of immediately revisiting the previous node, and q (in-out parameter, set to 0.5) controls whether the walk explores outward (BFS-like, discovering new neighborhoods) or stays local (DFS-like, exploring the immediate community). With q=0.5, walks are biased toward DFS — staying in the local neighborhood.')

bold_para('What Word2Vec Skip-gram does with them: ', 'After generating millions of random walks (num_walks=10 per node, so ~4.5 million walks), these walks are treated like sentences in a language model. Word2Vec Skip-gram learns to predict which nodes appear near each other in walks. Nodes that frequently co-occur in walks get similar 128-dimensional embedding vectors. The intuition: if two genes often appear in the same walk, they are structurally similar in the graph.')

bold_para('Why dropped from paper: ', 'Node2Vec treats ALL edge types identically. The random walks cannot distinguish between drugBindsGene, geneInteractsWithGene, and chemicalIncreasesExpression. Since CardioKB has 25 distinct relationship types with different biological meanings, this is a fundamental limitation. A drug binding a gene target is biologically very different from a drug increasing a gene\'s expression, but Node2Vec sees both as "Drug is connected to Gene." The paper focuses on RotatE and CompGCN, which are relation-aware methods. Node2Vec results are retained as a baseline reference.')

bold_para('Results: ', 'Best AUROC 0.9504 with XGBoost decoder. The OOM issues during training required using PecanPy\'s SparseOTF mode instead of PreComp (which pre-computes all transition probabilities and exceeds memory on a 459K-node graph).')

doc.add_heading('6.4 RotatE', level=2)
bold_para('What knowledge graph embeddings are: ', 'Knowledge graph embedding methods represent every entity and relation in the graph as a vector (list of numbers) in a continuous space. A triple (head, relation, tail) like (Aspirin, drugBindsGene, COX2) is valid if h + r ≈ t in the embedding space (TransE model), or if h * r ≈ t (RotatE model). The model learns these vectors by training on known triples and trying to distinguish real triples from fake ones.')

bold_para('How rotation in complex vector space works: ', 'RotatE models each relation as a rotation in complex number space. For a triple (head, relation, tail), the scoring function is ||h ∘ r - t||, where h and t are complex vectors and r is a unit-modulus complex vector (meaning |r| = 1, so it only rotates, never scales). Each of the 25 relation types learns its own rotation angle. This means drugBindsGene rotates embeddings differently than chemicalIncreasesExpression, so the model can distinguish between relationship types.')

bold_para('Implementation: ', 'PyKEEN framework. Embedding dim = 128 complex (256 real after concatenating real and imaginary parts). Training: 200 epochs, batch_size=4096, learning_rate=1e-4, 64 negative samples per positive, NSSALoss (Negative Sampling Self-Adversarial Loss with margin=9.0 and adversarial_temperature=1.0). Trained on L40S GPU on HPC for ~10.3 hours.')

bold_para('Native MRR: ', '0.1119 (PyKEEN filtered ranking). This is modest because ranking evaluates ALL 459K entities as candidates — the model must rank the correct tail entity above 459,091 others. The XGBoost decoder on the restricted Drug × Disease candidate space is the operationally relevant metric.')

bold_para('Why cosine decoder fails with RotatE: ', 'RotatE embeddings are optimized for the rotation scoring function (h ∘ r ≈ t), not for cosine similarity. Two entities can have very similar cosine angles but be far apart after the relation-specific rotation. Cosine AUROC is only 0.5299 (barely above random 0.5). The XGBoost decoder learns the nonlinear relationship between embedding features and treatment probability.')

bold_para('Results: ', 'AUROC 0.9652, AUPRC 0.9655, Hits@100 = 31.1%, Hits@200 = 60.0% with XGBoost decoder.')

doc.add_heading('6.5 CompGCN (Graph Neural Network)', level=2)
bold_para('What a GNN is: ', 'A Graph Neural Network (GNN) learns node embeddings by aggregating information from each node\'s neighborhood. For each node, it collects messages from its neighbors, transforms them through learned weight matrices, and combines them into a new embedding. This is done in layers: layer 1 sees 1-hop neighbors, layer 2 sees 2-hop neighbors, etc. CompGCN uses 2 layers, so each node\'s embedding captures its 2-hop neighborhood structure.')

bold_para('Difference between R-GCN and CompGCN: ', 'R-GCN (Relational Graph Convolutional Network) learns a separate weight matrix W_r for each relation type. With 25 relation types and 128-dimensional hidden layers, each R-GCN layer needs 25 × 128 × 128 = 409,600 parameters per layer. With only 3,782 drugTreatsDisease training edges, this risks severe overfitting — the model memorizes training data instead of learning generalizable patterns.')

bold_para('Why CompGCN over R-GCN: ', 'CompGCN instead learns one d-dimensional embedding per relation (25 × 128 = 3,200 relation parameters) plus one shared weight matrix (128 × 128 = 16,384). The relation embedding is composed with each neighbor\'s embedding before aggregation using a composition operator. We use subtraction: φ(e_neighbor, e_relation) = e_neighbor - e_relation. This gives dramatically fewer parameters (32.2M total including node embeddings vs. potentially overfitting with relation-specific matrices) while still capturing relation-specific semantics.')

bold_para('How splits work differently in GNNs: ', 'Unlike RotatE which only sees individual triples, CompGCN sees the ENTIRE graph structure during message passing. To prevent leakage, we remove all val/test edges from the graph before training. The GNN then only sees the train-split edges during message passing. At evaluation time, we use the learned node and relation embeddings (which were computed from train-only structure) to score val/test drug-disease pairs.')

bold_para('Training details: ', 'Pure PyTorch (custom implementation, not PyTorch Geometric). 2 layers, 128 hidden dim, subtraction composition, dropout=0.3, learning_rate=1e-3, Adam optimizer, gradient clipping max_norm=1.0. 200 max epochs with early stopping (patience=20, checked every 10 epochs). Best model at epoch 120. 32,244,096 total parameters. Trained on HPC GPU for ~46 minutes. Loss: Binary cross-entropy on positive edges + 1:1 negative sampling per epoch. Scoring: DistMult-style h * r * t element-wise product.')

bold_para('Results vs RotatE: ', 'CompGCN + XGBoost achieves AUROC 0.9717, improving over RotatE by +0.0065. However, ranking metrics (Hits@K) are nearly identical: Hits@100 = 30.5% vs 31.1%, Hits@200 = 60.6% vs 60.0%. This suggests both methods produce comparably useful candidate lists despite the AUROC difference.')

doc.add_heading('6.6 Decoders Compared', level=2)
bold_para('Cosine similarity: ', 'Measures the angle between two embedding vectors. Score = (emb_drug · emb_disease) / (|emb_drug| × |emb_disease|), ranging from -1 to 1. Works well when similar entities cluster together in embedding space (as in Node2Vec, where cosine AUROC = 0.7195). Fails for RotatE/CompGCN (AUROC ~0.5, random) because these embeddings are not optimized for angular similarity — they encode relational structure that cosine cannot capture.')

bold_para('XGBoost: ', 'An ensemble of decision trees (gradient-boosted trees). For each Drug-Disease pair, it receives a feature vector combining: (1) Hadamard product emb_drug * emb_disease (element-wise, captures dimensional interactions), (2) absolute difference |emb_drug - emb_disease| (captures dimensional distances), (3) cosine similarity (1 scalar), (4) L2 distance (1 scalar), (5) 6 structural features: shared neighbors count, Jaccard coefficient, Adamic-Adar index, log preferential attachment, log drug degree, log disease degree. Total features: 264 for 128-dim (Node2Vec, CompGCN), 520 for 256-dim (RotatE). Hyperparameters: n_estimators=300, max_depth=6, learning_rate=0.1, early_stopping_rounds=20 on validation set.')

bold_para('MLP (one hidden layer): ', 'A simple neural network with one hidden layer. Takes the same feature vector as XGBoost, passes it through a linear transformation + ReLU activation + dropout, then a final sigmoid output. Tried as an alternative nonlinear decoder. Performance is competitive with XGBoost but slightly lower (e.g., CompGCN MLP AUROC 0.9625 vs XGBoost 0.9717).')

bold_para('Why XGBoost won: ', 'XGBoost consistently outperforms both cosine and MLP across all three embedding methods. It excels because: (1) decision trees naturally handle the mix of embedding features (continuous, high-dimensional) and structural features (discrete, low-dimensional); (2) gradient boosting builds trees sequentially, with each tree correcting the previous one\'s errors; (3) it handles feature interactions without explicit engineering; (4) built-in regularization (max_depth, learning_rate) prevents overfitting on the relatively small positive set (3,026 training edges).')

bold_para('What the confidence score means: ', 'The XGBoost output is a probability between 0 and 1 representing how likely the model thinks a Drug-Disease pair should have a drugTreatsDisease edge, based on the embedding patterns of known treatment pairs. A score of 0.98 means the embedding geometry of this pair closely matches known drug-disease treatment relationships. It is NOT a clinical efficacy probability — it does not mean there is a 98% chance the drug works. It means the graph structure strongly suggests a therapeutic relationship exists.')

doc.add_heading('6.7 Evaluation Metrics', level=2)

bold_para('AUROC (Area Under ROC Curve): ', 'Measures the model\'s ability to discriminate between true drug-disease treatment pairs and random non-treatment pairs across all classification thresholds. In plain English: if you pick one real treatment pair and one random non-treatment pair, AUROC = 0.965 means the model assigns a higher score to the real pair 96.5% of the time. Limitation: our 1:1 random negatives are "easy" — most random Drug-Disease combos are clearly non-therapeutic (e.g., a veterinary antibiotic treating coronary artery disease), which inflates AUROC.')

bold_para('AUPRC (Area Under Precision-Recall Curve): ', 'Measures precision-recall balance across all thresholds. More informative than AUROC for imbalanced datasets because it focuses on how well the model performs on the POSITIVE class. A high AUPRC (0.9709 for CompGCN) means the model has few false positives among its top-ranked predictions — when it says a drug treats a disease, it\'s usually right. This is critical for drug repurposing where false positives waste expensive experimental validation resources.')

bold_para('MRR (Mean Reciprocal Rank): ', 'For each known treatment pair, rank all candidate diseases for that drug. MRR = average of 1/rank. If the true disease is ranked 1st → score 1.0; ranked 5th → score 0.2; ranked 100th → score 0.01. RotatE\'s native PyKEEN MRR is 0.1119, meaning the true disease is typically ranked around position 9 (1/0.1119 ≈ 9). This ranking is against ALL 459K entities — the XGBoost decoder on the restricted Drug × Disease space is the operationally relevant metric.')

bold_para('Hits@K: ', 'Fraction of true treatment pairs appearing in the top K predictions for each drug. Hits@100 = 30.5% means about 1 in 3 true treatments appears in the top 100 candidates for its drug. This maps directly to a practical screening scenario: if a researcher examines the top 100 predictions, they\'ll find ~31% of actual treatments. Hits@200 = 60.6% means examining the top 200 finds ~61% of true treatments.')

bold_para('Why rank-based metrics matter more for drug repurposing: ', 'Drug repurposing is a retrieval task, not a classification task. We don\'t need to classify every pair as treat/don\'t-treat — we need to surface the best candidates for experimental validation. Hits@K directly measures this: how many true treatments land in a feasibly small candidate list? Notably, Hits@K values are nearly identical across RotatE and CompGCN despite AUROC differences, suggesting both methods produce comparably useful candidate lists for practical screening.')

bold_para('Confusion matrix (CompGCN XGBoost, test set): ', '306 True Negatives (correctly identified as non-treatment), 252 True Positives (correctly identified known treatments), 16 False Positives (predicted treatment but actually not — drugs the model wrongly thinks treat a disease), 70 False Negatives (missed known treatments — real treatment pairs the model failed to identify). Precision = 252/(252+16) = 94.0%, Recall = 252/(252+70) = 78.3%.')

doc.add_heading('6.8 Feature Importance', level=2)
doc.add_paragraph('XGBoost feature importance reveals how much the model relies on embedding features vs. structural graph features:')
add_table(
    ['Method', 'Structural Feature %', 'Embedding Feature %', 'Top Structural Feature'],
    [
        ['Node2Vec', '~30% (L2 alone: 20.1%)', '~70%', 'L2 distance (0.201)'],
        ['RotatE', '0.6%', '21.8%', 'Disease degree (0.006)'],
        ['CompGCN', '4.3%', '50.8%', 'Disease degree (0.036)'],
    ],
    col_widths=[3, 4, 4, 5]
)
bold_para('Why structural features barely contribute for RotatE/CompGCN: ', 'Shared neighbors, Jaccard coefficient, and Adamic-Adar index contribute almost nothing (< 1%) for RotatE and CompGCN. This means the learned embeddings already capture the relevant graph structure — the model doesn\'t need explicit structural features because the embeddings encode that information implicitly. Node2Vec relies heavily on L2 distance (20.1%) and structural features (~30%) because its embeddings lack relation-type information, so structural features compensate.')

bold_para('What this tells us about the embeddings: ', 'CompGCN concentrates importance in specific embedding dimensions (hadamard_46 alone = 17.4% importance), suggesting it learns targeted relational representations — specific dimensions encode specific biological relationships. RotatE spreads importance more evenly across embedding dimensions. Both approaches work but represent different learning strategies.')

bold_para('What L2 distance means vs shared neighbors: ', 'L2 distance = Euclidean distance between two embedding vectors. A small L2 distance means the model places the drug and disease close together in embedding space. Shared neighbors = how many graph nodes are connected to both the drug and the disease. For Node2Vec, L2 is the dominant feature (20.1%) because its embeddings directly encode proximity. For RotatE/CompGCN, shared neighbors is negligible because the embeddings already encode this information through the GNN message passing or rotation learning.')

doc.add_heading('6.9 Results and Clinical Validation', level=2)

doc.add_heading('Full Results Table', level=3)
add_table(
    ['Method', 'Dim', 'Decoder', 'AUROC', 'AUPRC', 'Hits@100', 'Hits@200'],
    [
        ['Node2Vec', '128', 'Cosine', '0.7195', '0.7142', '25.2%', '45.0%'],
        ['Node2Vec', '128', 'XGBoost', '0.9504', '0.9579', '31.1%', '61.8%'],
        ['Node2Vec', '128', 'MLP', '0.9441', '0.9535', '30.8%', '61.2%'],
        ['RotatE', '256', 'Cosine', '0.5299', '0.5401', '19.3%', '32.3%'],
        ['RotatE', '256', 'XGBoost', '0.9652', '0.9655', '31.1%', '60.0%'],
        ['RotatE', '256', 'MLP', '0.9607', '0.9588', '30.7%', '60.9%'],
        ['CompGCN', '128', 'Cosine', '0.5058', '0.5041', '16.9%', '30.2%'],
        ['CompGCN', '128', 'XGBoost', '0.9717', '0.9709', '30.5%', '60.6%'],
        ['CompGCN', '128', 'MLP', '0.9625', '0.9625', '30.5%', '59.7%'],
    ],
    col_widths=[2.5, 1.2, 2, 1.8, 1.8, 2, 2]
)

doc.add_heading('Clinical Validation of Top 30 Predictions', level=3)
doc.add_paragraph(
    'For both RotatE and CompGCN, 4 out of 30 top predictions (13.3%) have supporting clinical trial '
    'evidence already in CardioKB. This validation rate is notable because these predictions are specifically '
    'for Drug-Disease pairs NOT in the training data — the model has never seen these treatment relationships.'
)
add_table(
    ['Method', 'Drug → Disease', 'Trial ID', 'Status'],
    [
        ['RotatE', 'ACE Inhibitors → heart disease', 'NCT00224809', 'Known trial'],
        ['RotatE', 'Atorvastatin → cardiomyopathy', 'NCT00317967', 'Known trial'],
        ['RotatE', 'Clopidogrel → cerebrovascular disease', 'NCT01823185, NCT02121288', 'Known trials'],
        ['RotatE', 'Colchicine → hypertension', 'NCT04916522', 'Known trial'],
        ['CompGCN', 'Levosimendan → coronary artery disease', 'NCT00130871', 'Known trial'],
        ['CompGCN', 'ACE Inhibitors → heart disease', 'NCT00224809', 'Known trial'],
        ['CompGCN', 'Adenosine → myocardial infarction', 'NCT00781404', 'Known trial'],
        ['CompGCN', 'Methylprednisolone → coronary artery disease', 'NCT07101367', 'Known trial'],
    ],
    col_widths=[2.5, 6, 4.5, 3]
)
doc.add_paragraph(
    'Many other top predictions are clinically plausible but lack trials in CardioKB (e.g., Clopidogrel '
    'for heart failure, statins for ischemia). The 13.3% validation rate against an independent data '
    'source (ClinicalTrials.gov) that was NOT used to generate the predictions provides evidence that the '
    'embeddings capture genuine pharmacological relationships.'
)

bold_para('Predictions stored in Memgraph: ', 'Top 500 predictions per method (confidence >= 0.5) are stored as predictedTreatsDisease edges with properties: confidence (XGBoost probability, float 0-1) and source ("Node2Vec_LinkPrediction", "RotatE_LinkPrediction", or "CompGCN_LinkPrediction"). Total: 1,500 prediction edges. Displayed in the web UI as cyan dashed edges with "not clinically validated" warning.')

doc.add_heading('6.10 Future ML Directions', level=2)
bullet('Attention-based GNNs (GAT, HGT): Learn attention weights over different neighbor types, potentially identifying which relationship types are most informative for drug repurposing.')
bullet('Deeper models: Current CompGCN uses 2 layers (2-hop neighborhood). Adding layers could capture longer-range dependencies but risks oversmoothing (all embeddings converging to the same vector).')
bullet('Edge-type-specific prediction heads: Instead of a single XGBoost decoder, train separate heads for different prediction tasks (drug repurposing, side effect prediction, gene-disease association).')
bullet('Temporal modeling: Incorporate clinical trial phase and status as temporal signals to predict which treatments will succeed in later phases.')
bullet('ASAREE/Persona experiments: Persona-based querying where different user types (clinician, molecular biologist, pharmacologist) see different subgraph views tailored to their research needs.')
bullet('Clinical validation by domain experts: Present top predictions to cardiologists for expert assessment of clinical plausibility beyond what trial database matching can verify.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 7. USER INTERFACE
# ══════════════════════════════════════════════════════════
doc.add_heading('7. User Interface — Every Feature Explained', level=1)
doc.add_paragraph(
    'CardioKB provides a single-page web application (interface/index.html) served by the Flask backend '
    'on port 5050. The interface uses vis.js for network visualization, Chart.js for analytics charts, '
    'and a custom dark/light theme system with CSS custom properties.'
)

doc.add_heading('7.1 Explore Tab', level=2)
doc.add_paragraph(
    'The Explore tab is the primary graph visualization interface. It lets users search for any disease, '
    'gene, or drug and visualize its immediate (1-hop) neighborhood as an interactive force-directed network graph.'
)

bold_para('Search bar and autocomplete: ', 'The search bar provides real-time autocomplete as the user types (triggers at 2+ characters). The backend queries Memgraph using STARTS WITH, falling back to CONTAINS if no results. The dropdown shows matching entities with color-coded type badges (blue=Gene, red=Disease, green=Drug). Users select from the dropdown for accurate matching. A "Max" input (default 200, max 1000) controls how many nodes are displayed. Quick-explore buttons appear when the graph is empty: "atrial fibrillation", "APOE", "Metoprolol", "coronary artery disease".')

bold_para('Graph visualization: ', 'The graph renders in a large container (70% viewport height, min 450px) using vis.js with ForceAtlas2Based physics (gravitationalConstant=-60, centralGravity=0.008, springLength=140). Nodes use persistent DataSets — never destroyed and recreated, only cleared and refilled. The searched entity gets a white border with glow effect. Disease nodes are diamonds; all others are circles. Node size reflects importance/degree.')

bold_para('Node type filter chips: ', 'Color-coded chips above the graph let users show/hide specific node types. All 17 types have chips, active by default. Clicking a chip toggles it (dimmed = hidden). Useful when a disease has hundreds of gene connections and the user wants to focus on drugs or pathways.')

bold_para('Edge type filter chips: ', 'A second row of chips filters by relationship type. Each chip represents one relationship type present in the current subgraph.')

bold_para('Node detail panel: ', 'Clicking any node opens a slide-in panel (320px wide) on the right showing: node name, type (color-coded badge), all properties (geneSymbol, drugId, diseaseName, specificityScore, etc.), and a Connected Nodes list (max 20) with their types and relationships. A "View All Connections" button loads the full connection list paginated from the database (50 per page, max 200 per page).')

bold_para('Node legend: ', 'Below the graph, a legend strip shows all node types with color dots.')

bold_para('Visual Guide: ', 'A collapsible panel explains node shapes (circle vs diamond), edge styles (solid gray = curated, dashed cyan = ML predicted, solid orange = disease hierarchy), and how to interact with the graph.')

bold_para('About panel: ', 'An expandable "About: How to Interpret This Graph" panel provides in-depth explanations of: what knowledge graphs are, disease-specificity scoring (formula and interpretation), ML prediction methodology (embedding → features → XGBoost → prediction), model performance table, important limitations, and data source overview.')

bold_para('Export options: ', 'Four buttons below the graph: Export CSV (edges as CSV with from/to/type/source columns), Export JSON (full node+edge data), Export PNG (canvas screenshot), Export PDF (browser print dialog).')

doc.add_heading('7.2 Query Tab', level=2)

bold_para('Natural language querying (CypherGPT): ', 'At the top of the Query tab, users type questions in plain English (e.g., "What drugs target genes associated with heart failure?"). The system translates to Cypher using an adaptation of CypherGPT/Eng2Cypher by Jay Moran (Center for AI Research). The pipeline: (1) introspects the live graph schema (node types, relationship types, properties) and caches it with 24-hour TTL, (2) builds a context-aware system prompt with CardioKB-specific instructions and examples, (3) calls Claude API (claude-sonnet-4-6, temperature=0.0), (4) sanitizes output (strips code fences, enforces case-insensitive matching with toLower()), (5) validates against schema using fuzzy matching (difflib.get_close_matches), auto-corrects typos, (6) handles cross-node-type medical conditions (Disease/Phenotype/SideEffect) via UNION ALL queries. Generated Cypher is shown in the editor for inspection, then auto-executed.')

bold_para('Query templates: ', '10 pre-built templates: Disease Subgraph, Gene Neighbors, Drug Targets, Clinical Trials, Pathway Genes, Top Connected Genes, Gene Expression, Drug-Side Effects, Drug Repurposing (with UNION for ML predictions), and Shared Genes Between Diseases.')

bold_para('Direct Cypher editor: ', 'Expandable via "Or write Cypher directly." Monospace textarea (SF Mono/Fira Code) with Ctrl+Enter to run. Read-only mode blocks writes (CREATE, DELETE, MERGE, SET, REMOVE, DROP). Save Query button bookmarks queries.')

bold_para('Multi-panel results (Neo4j Browser style): ', 'Each query appends a new collapsible result panel (not replacing previous ones) allowing side-by-side comparison. Each panel has: collapsible header (Cypher text, row count, duration), Table view tab (scrollable, sticky headers, max 500 rows), Graph view tab (vis.js network for returned nodes/edges), stats footer, close button. Clear All removes all panels.')

doc.add_heading('7.3 Edge Provenance Panel', level=2)
doc.add_paragraph(
    'Clicking any edge in the graph opens the "Why is this edge here?" panel showing:'
)
bullet('Relationship type (e.g., "geneAssociatesWithDisease")')
bullet('From → To nodes with type-colored labels')
bullet('Source database in a green badge (e.g., "PubTator", "DrugBank", "STRING")')
bullet('Evidence properties in table format: combinedScore (0-1000), expressionScore, morScore (-1/0/+1), confidence (A-E), evidenceCode, score (0-1), interactionType, clinicalSignificance, zScore')
bullet('For ML predictions: confidence percentage, method ("RotatE + XGBoost" or "CompGCN + XGBoost"), red warning "Not clinically validated"')
bullet('Provenance interpretation at bottom explaining the source database (e.g., "STRING combined scores range from 0-1000, >700 = high confidence")')

doc.add_heading('7.4 ML Predicted Edges', level=2)
doc.add_paragraph(
    'The "Show ML Predictions" toggle checkbox controls visibility of 1,500 predicted drug-disease edges. '
    'When enabled: cyan dashed edges appear connecting drugs to diseases. A Drug Repurposing Predictions '
    'panel expands below showing: method filter checkboxes (RotatE, CompGCN with AUROC scores), a sortable '
    'predictions table (Drug, Disease, Method, Confidence columns), CSV/JSON export buttons, and a '
    '"Methodology & Metrics" button linking to the About panel. An "Experimental" badge and disclaimer '
    'remind users these are computational predictions, not clinical recommendations.'
)

doc.add_heading('7.5 Sidebar Tools', level=2)
bold_para('System Check: ', 'Triggers a comprehensive SSE-streamed health check of all data sources, database connectivity, and data integrity. Results populate the Admin panel\'s Parser Status table and Health Checks section.')

bold_para('Extract Disease Subgraph: ', 'N-hop subgraph extraction with: disease name input (partial match), hops slider (1-3 with info popover explaining each level), "Build Subgraph" button, stats display (node/edge counts, types, sources), and Export JSON/CSV buttons. Queries existing data only — no new data fetched. The "Instant" badge indicates this. A "?" info button explains that 1 hop = direct connections (focused), 2 hops = connections of connections (shared pathways, comorbidity bridges), 3 hops = broad network for hypothesis generation (may be slow).')

bold_para('Theme toggle: ', 'Switches between dark theme (slate navy, #0f172a background, ideal for extended use) and light theme (white, #f1f5f9 background, better for screenshots and presentations).')

doc.add_heading('7.6 Admin Panel', level=2)
doc.add_paragraph(
    'Collapsible "Admin: Parser Status & Pipeline Log" panel at the bottom of the main content area:'
)
bullet('Parser Status table: Lists every parser with Working/Failed/Skipped badge and duration. Status determined dynamically by querying graph for expected nodes/relationships.')
bullet('Health Checks: Green (OK) or red (warning) indicators for node count validation, edge count validation, source label coverage (all 19 expected source labels), edge property coverage, database connectivity.')
bullet('Charts: Two Chart.js bar charts — node counts by type and relationship counts by type.')
bullet('ID Mapping Report: Table showing match rates for ID mappings between TSV relationship files and existing graph nodes. Identifies where edges are lost to unresolvable identifiers.')
bullet('Pipeline Log: Monospace panel with color-coded lines (green=success, red=error, blue=info) and timestamps.')

doc.add_heading('7.7 Visual Design', level=2)
bold_para('Welcome tour modal: ', 'Three-step modal on first visit: (1) Overview of CardioKB with live stats, (2) Explore & Query tabs walkthrough, (3) ML Drug Repurposing Predictions with limitations warning. Navigate with Next/Previous/Skip.')

bold_para('Tooltip system: ', 'Custom JS tooltip engine — hovering over any element with data-tip attribute shows a positioned bubble. Info buttons ("?") use a click-triggered popover system with richer HTML content (multi-paragraph explanations, color-coded tiers).')

bold_para('Responsive design: ', 'Below 768px viewport width, sidebar collapses to full-width header, detail panel becomes inline, and admin grid stacks vertically.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 8. DEPLOYMENT
# ══════════════════════════════════════════════════════════
doc.add_heading('8. Deployment', level=1)

doc.add_heading('Memgraph in Docker', level=2)
bold_para('What Memgraph is: ', 'Memgraph is an in-memory graph database that stores all data in RAM for fast query execution. It uses the Cypher query language (same as Neo4j) and communicates via the Bolt protocol on port 7687. The full CardioKB graph occupies ~14 GB in memory (uncompressed).')
bold_para('Why Docker: ', 'Docker packages Memgraph and its dependencies into a container that runs identically on any machine. No manual installation of Memgraph, Python, or dependencies needed. The docker-compose.yml defines two services: memgraph (the database on ports 7687/7444) and app (the Flask web app on port 5050). Data persists in a named Docker volume (memgraph-data) that survives container restarts.')

doc.add_heading('Flask App', level=2)
bold_para('What Flask is: ', 'Flask is a lightweight Python web framework. The CardioKB Flask app (src/api.py) serves the web interface (interface/index.html) and provides RESTful API endpoints for graph queries, autocomplete, NL-to-Cypher, health checks, and subgraph extraction. It connects to Memgraph via bolt://memgraph:7687 (Docker internal networking). Built from Python 3.11-slim Docker image with requirements.txt dependencies.')

doc.add_heading('mgconsole Import Method', level=2)
bold_para('Why better than Python loader: ', 'The graph data (459K nodes, 5.4M edges) could be loaded via the Python pipeline (which replays millions of MERGE statements), but this takes ~5 minutes and requires the full source data. Instead, we export the binary Memgraph data volume as a tar.gz (~300 MB compressed, ~14 GB uncompressed) and restore it directly. This is a binary-level copy of the entire database state — all indexes, constraints, and data — and takes seconds to restore. The import script (scripts/import_graph.sh) stops Memgraph, extracts the archive into the Docker volume, restarts Memgraph, and verifies node/relationship counts.')

doc.add_heading('How to Take a Memgraph Snapshot', level=2)
doc.add_paragraph('To export the current graph state for deployment to another machine:')
bullet('Run scripts/export_graph.sh — this stops Memgraph, tars /var/lib/memgraph from the Docker volume, and restarts Memgraph')
bullet('The script auto-detects the Docker Compose container and volume names from the project directory')
bullet('Output: data/export/memgraph-data.tar.gz (~300 MB)')
bullet('Transfer this file to the target machine, then run scripts/import_graph.sh <archive.tar.gz>')

doc.add_heading('Environment Variables', level=2)
add_table(
    ['Variable', 'Required?', 'Purpose'],
    [
        ['MEMGRAPH_PASSWORD', 'Yes', 'Graph database authentication password'],
        ['ADMIN_PASSWORD', 'Yes', 'Required for System Check and admin features in the web UI'],
        ['ANTHROPIC_API_KEY', 'Optional', 'Enables "Ask AI" natural language → Cypher in the Query tab. Web UI works fully without it.'],
        ['ANTHROPIC_FOUNDRY_API_KEY + BASE_URL', 'Optional', 'Azure AI Foundry (takes priority over ANTHROPIC_API_KEY when both set)'],
        ['DRUGBANK_USERNAME / PASSWORD', 'No', 'Pipeline only — not needed for deployment'],
    ],
    col_widths=[5, 2, 9]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 9. EVALUATION AND VALIDATION
# ══════════════════════════════════════════════════════════
doc.add_heading('9. Evaluation and Validation', level=1)

doc.add_heading('Automated Health Check System', level=2)
doc.add_paragraph(
    'CardioKB includes an automated evaluation system (src/admin_agent.py) that performs dynamic, '
    'graph-based parser status detection. It can be triggered from the web UI\'s "System Check" button '
    'or run programmatically. The system checks:'
)
bullet('Node count validation: Each of the 17 node types has a non-zero count matching expected ranges')
bullet('Edge count validation: Each of the 28 relationship types has expected counts')
bullet('Source label coverage: All 20 expected edge source labels are present (plus 3 ML sources)')
bullet('Property coverage: Edge properties (combinedScore, morScore, expressionScore, etc.) are populated at 100% coverage')
bullet('Orphan rate calculation: Percentage of nodes with zero edges per type')
bullet('Dangling edges: Edges pointing to non-existent nodes = 0 (verified)')

doc.add_heading('Node/Edge Counts by Type', level=2)
doc.add_paragraph(
    'The node and relationship count tables in Section 3 are the canonical reference. Key summary: '
    'Gene nodes dominate (42.2% of graph), followed by Variant (29.5%) and Drug (7.2%). '
    'bodyPartOverexpressesGene edges dominate relationships (50.5% of all edges, from Bgee).'
)

doc.add_heading('Property Coverage Gaps and What Was Fixed', level=2)
bullet('STRING combinedScore: 100% coverage (229,007 edges) — no gap')
bullet('DoRothEA morScore and confidence: 100% coverage (15,082 edges) — no gap')
bullet('Bgee expressionScore: 100% coverage (2,749,193 edges) — no gap')
bullet('LINCS L1000 zScore: 100% coverage (139,515 edges) — no gap')
bullet('ClinVar clinicalSignificance: loaded on Variant nodes as property — no gap')
bullet('ML confidence: 100% coverage on predictedTreatsDisease edges (1,500 edges) — no gap')
bullet('BindingDB: No binding affinity values loaded (only edge existence) — known gap')
bullet('CTD: No PubMed IDs loaded as edge properties — known gap')

doc.add_heading('Known Gaps', level=2)
bullet('ClinVar variant-disease yield: Only 0.12% of ClinVar variant-disease associations map to CardioKB diseases (212 UMLS CUI overlap). Expected given CVD-focused disease filter.')
bullet('Symptom orphan rate: ~100% — no active source provides symptom-disease edges. MeSH provides nodes only.')
bullet('MEDLINE: 0 edges in current build (configs skipped). Legacy source with minimal contribution.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 10. FUTURE DIRECTIONS
# ══════════════════════════════════════════════════════════
doc.add_heading('10. Future Directions', level=1)

doc.add_heading('CompGCN Results Analysis for Binglan', level=2)
doc.add_paragraph(
    'The CompGCN results (AUROC 0.9717, best overall) demonstrate that GNN-based approaches outperform '
    'both shallow (Node2Vec) and translational (RotatE) embedding methods on this graph. The explainability '
    'analysis (ml/explainability_analysis.py) has produced paper-ready tables comparing top predictions, '
    'feature importance, and clinical validation rates. These results are ready for discussion with Binglan '
    'about integration into the BaseAgent paper or a standalone CardioKB methods paper.'
)

doc.add_heading('ASAREE/Persona Experiments', level=2)
bullet('Evidence quality scoring: Automatically assess gene-disease association strength based on the number and type of supporting edges (literature-mined vs. curated vs. experimental)')
bullet('Persona-based querying: Different user personas (clinician, molecular biologist, pharmacologist) see different subgraph views and node rankings tailored to their research needs')
bullet('Automated hypothesis generation: Generate and rank novel biological hypotheses based on graph topology patterns')

doc.add_heading('CypherGPT Natural Language Layer (Implemented)', level=2)
doc.add_paragraph(
    'The Query tab now includes an AI-powered natural language interface, adapted from CypherGPT/Eng2Cypher '
    'by Jay Moran (Center for AI Research). Schema is cached with 24-hour TTL in config/nl2cypher/. '
    'CardioKB-specific instructions and examples in config/nl2cypher/instructions.md and examples.md guide '
    'the LLM. The system handles case-insensitive matching, fuzzy property/label correction, and cross-node-type '
    'medical condition queries via UNION ALL.'
)

doc.add_heading('drugTreatsDisease Enrichment (DrugBank Indications) — Implemented', level=2)
doc.add_paragraph(
    'The original 3,782 drugTreatsDisease edges came from 3 sources (CTD, ClinicalTrials.gov, DrugCentral). '
    'DrugBank XML files contain free-text indication fields that were not captured as structured edges. '
    'A text-mining script (scripts/drugbank_indications.py) was implemented to extract treatment relationships '
    'by matching Disease and Phenotype node names against indication text using whole-word regex matching. '
    'This added 2,930 new drugTreatsDisease edges (total: 6,272) and created a new drugTreatsPhenotype '
    'relationship type with 10,955 edges from 2,646 drugs to 692 phenotypes. The phenotype edges address '
    'conditions like tachycardia, arrhythmia, and edema that exist only as HPO Phenotype nodes. A blocklist '
    'filters out HPO modifier terms (Acute, Chronic, Severe, etc.) to prevent false positives. These edges '
    'were added post-ML-training and do not affect the ML pipeline results.'
)

doc.add_heading('diseaseIsSubtypeOf Hierarchy Edges', level=2)
doc.add_paragraph(
    'Disease Ontology provides 6,447 diseaseIsSubtypeOf edges creating a full disease taxonomy. Currently '
    '2,581 are loaded. Expanding coverage would enable hierarchical disease queries (e.g., "find all drugs '
    'treating any subtype of cardiovascular system disease") and improve the ML pipeline by providing '
    'disease similarity signals through the graph structure.'
)

doc.add_heading('CardioKB Paper Outline', level=2)
doc.add_paragraph(
    'A methods paper could cover: (1) the automated BaseAgent pipeline and its advantages over manual KG '
    'construction, (2) the 22-source deduplication audit methodology, (3) the RotatE vs CompGCN drug '
    'repurposing comparison with clinical validation, (4) the web interface as a tool for CVD researchers. '
    'The explainability analysis has produced paper-ready LaTeX tables and figures.'
)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'CardioKB_Complete_Study_Guide.docx')
doc.save(output_path)
print(f'Saved to {output_path}')
